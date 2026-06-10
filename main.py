import sys, os, requests, json, re, hashlib, secrets, asyncio
from datetime import datetime, timedelta
from fastapi import FastAPI, Depends, HTTPException, Header, Request, Body
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional, List
import asyncpg

# ── Auth app import ──────────────────────────────────────────────
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src', 'auth'))
from src.auth.main import app

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "https://oktolk.ru",
        "https://www.oktolk.ru",
    ],
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    allow_headers=["*"],
)

# ── Env ──────────────────────────────────────────────────────────
AITUNNEL_API_KEY  = os.getenv("AITUNNEL_API_KEY", "")
AITUNNEL_BASE_URL = os.getenv("AITUNNEL_BASE_URL", "https://api.aitunnel.ru/v1/")
YANDEX_API_KEY    = os.getenv("YANDEX_API_KEY", "")
DEEPSEEK_API_KEY  = os.getenv("DEEPSEEK_API_KEY", "")
DEEPSEEK_BASE_URL = "https://api.deepseek.com/v1/"
SECRET_KEY        = os.getenv("SECRET_KEY", "oktolk-super-secret-key-2026-production")

# Ticketland партнёрская программа (Advcake)
# Формат: вставить шаблон из личного кабинета Advcake, где {url} — целевая страница Ticketland
# Пример: https://advcake.ru/cgi-bin/go.pl?a=clickstream&p=ticketlandru&id=12345&dl={url}
TICKETLAND_AFFILIATE_TEMPLATE = os.getenv("TICKETLAND_AFFILIATE_TEMPLATE", "")

# VAPID ключи для Web Push (хранить в env Amvera!)
VAPID_PUBLIC_KEY  = os.getenv("VAPID_PUBLIC_KEY",  "BPuepfkM-HSrhHb8NZPZ2HfhFDNYEZd2TUiwZN7tvCij3qNNhqMyUMWDtuOObEEGbeYnWnhBaSF8sGdeSQrApt4")
VAPID_PRIVATE_PEM = os.getenv("VAPID_PRIVATE_PEM", "")  # полный PEM в env переменной
VAPID_CLAIMS      = {"sub": "mailto:info@oktolk.ru"}

DB_HOST = os.getenv("DB_HOST", "amvera-kes-cnpg-oktolk-db-rw")
DB_PORT = int(os.getenv("DB_PORT", "5432"))
DB_NAME = os.getenv("DB_NAME", "oktolk")
DB_USER = os.getenv("DB_USER", "oktolk_user")
DB_PASS = os.getenv("DB_PASS", "")

DATABASE_URL = f"postgresql://{DB_USER}:{DB_PASS}@{DB_HOST}:{DB_PORT}/{DB_NAME}"

# ── DB pool ──────────────────────────────────────────────────────
db_pool = None

async def get_db():
    return await db_pool.acquire()

@app.on_event("startup")
async def startup():
    global db_pool
    try:
        db_pool = await asyncpg.create_pool(DATABASE_URL, min_size=1, max_size=10)
        await init_db()
        print("✅ PostgreSQL connected")
        # Запускаем планировщик push-уведомлений
        asyncio.create_task(push_scheduler())
        print("✅ Push scheduler started")
    except Exception as e:
        print(f"⚠️ PostgreSQL not available: {e}")

@app.on_event("shutdown")
async def shutdown():
    if db_pool:
        await db_pool.close()

async def init_db():
    async with db_pool.acquire() as conn:
        await conn.execute("""
            CREATE TABLE IF NOT EXISTS users (
                id          SERIAL PRIMARY KEY,
                phone       VARCHAR(20) UNIQUE NOT NULL,
                name        VARCHAR(100),
                password_hash VARCHAR(200),
                created_at  TIMESTAMP DEFAULT NOW(),
                tariff      VARCHAR(20) DEFAULT 'demo',
                tariff_until TIMESTAMP,
                marketing_consent BOOLEAN DEFAULT FALSE,
                is_demo     BOOLEAN DEFAULT FALSE
            );

            CREATE TABLE IF NOT EXISTS sessions (
                token       VARCHAR(200) PRIMARY KEY,
                user_id     INTEGER REFERENCES users(id) ON DELETE CASCADE,
                created_at  TIMESTAMP DEFAULT NOW(),
                expires_at  TIMESTAMP
            );

            CREATE TABLE IF NOT EXISTS health_records (
                id          SERIAL PRIMARY KEY,
                user_id     INTEGER REFERENCES users(id) ON DELETE CASCADE,
                type        VARCHAR(20),
                value_1     FLOAT,
                value_2     FLOAT,
                comment     TEXT,
                recorded_at TIMESTAMP DEFAULT NOW()
            );

            CREATE TABLE IF NOT EXISTS medications (
                id          SERIAL PRIMARY KEY,
                user_id     INTEGER REFERENCES users(id) ON DELETE CASCADE,
                name        VARCHAR(200),
                dose        VARCHAR(100),
                schedule    VARCHAR(100),
                start_date  DATE,
                end_date    DATE,
                is_active   BOOLEAN DEFAULT TRUE
            );

            CREATE TABLE IF NOT EXISTS push_subscriptions (
                id          SERIAL PRIMARY KEY,
                user_id     INTEGER REFERENCES users(id) ON DELETE CASCADE,
                endpoint    TEXT NOT NULL,
                p256dh      TEXT NOT NULL,
                auth        TEXT NOT NULL,
                created_at  TIMESTAMP DEFAULT NOW(),
                UNIQUE(user_id, endpoint)
            );

            CREATE TABLE IF NOT EXISTS finance_records (
                id          SERIAL PRIMARY KEY,
                user_id     INTEGER REFERENCES users(id) ON DELETE CASCADE,
                category    VARCHAR(50),
                subcategory VARCHAR(50),
                item_name   TEXT,
                amount      FLOAT,
                comment     TEXT,
                recorded_at TIMESTAMP DEFAULT NOW()
            );

            CREATE TABLE IF NOT EXISTS features (
                name        VARCHAR(50) PRIMARY KEY,
                enabled     BOOLEAN DEFAULT FALSE,
                description TEXT
            );

            CREATE TABLE IF NOT EXISTS error_logs (
                id          SERIAL PRIMARY KEY,
                error       TEXT,
                created_at  TIMESTAMP DEFAULT NOW()
            );

            CREATE TABLE IF NOT EXISTS user_profile (
                user_id     INTEGER PRIMARY KEY,
                gender      VARCHAR(10),
                age         INTEGER,
                height      INTEGER,
                weight      REAL,
                profession  TEXT,
                hobbies     TEXT,
                work_pressure_1 INTEGER,
                work_pressure_2 INTEGER,
                work_pulse  INTEGER,
                base_sugar  REAL,
                habits      TEXT,
                activity    VARCHAR(20),
                alcohol     INTEGER DEFAULT 0,
                smoking     INTEGER DEFAULT 0,
                smoking_years INTEGER,
                heredity    TEXT,
                chronic     TEXT,
                income      REAL,
                updated_at  TIMESTAMP DEFAULT NOW()
            );

            CREATE TABLE IF NOT EXISTS credits (
                id          SERIAL PRIMARY KEY,
                user_id     INTEGER NOT NULL,
                name        TEXT,
                amount      REAL,
                rate        REAL,
                term_months INTEGER,
                monthly_payment REAL,
                notes       TEXT,
                created_at  TIMESTAMP DEFAULT NOW()
            );

            CREATE TABLE IF NOT EXISTS loans (
                id          SERIAL PRIMARY KEY,
                user_id     INTEGER NOT NULL,
                name        TEXT,
                amount      REAL,
                due_date    TEXT,
                monthly_payment REAL,
                notes       TEXT,
                created_at  TIMESTAMP DEFAULT NOW()
            );

            CREATE TABLE IF NOT EXISTS chat_messages (
                id          SERIAL PRIMARY KEY,
                user_id     INTEGER NOT NULL,
                domain      VARCHAR(20),
                role        VARCHAR(10),
                text        TEXT,
                action      VARCHAR(30),
                created_at  TIMESTAMP DEFAULT NOW()
            );
            CREATE INDEX IF NOT EXISTS idx_chat_messages_user_domain
                ON chat_messages (user_id, domain, created_at);

            CREATE TABLE IF NOT EXISTS saved_searches (
                id          SERIAL PRIMARY KEY,
                user_id     INTEGER NOT NULL,
                title       TEXT,
                url         TEXT,
                description TEXT,
                source      TEXT,
                image       TEXT,
                price       TEXT,
                cat         TEXT,
                saved_at    TIMESTAMP DEFAULT NOW(),
                UNIQUE (user_id, url)
            );
        """)

        # Миграции для существующих таблиц (CREATE IF NOT EXISTS не добавляет колонки)
        await conn.execute("""
            ALTER TABLE users ADD COLUMN IF NOT EXISTS marketing_consent BOOLEAN DEFAULT FALSE;
            ALTER TABLE users ADD COLUMN IF NOT EXISTS is_demo BOOLEAN DEFAULT FALSE;
            ALTER TABLE users ALTER COLUMN phone TYPE VARCHAR(40);
            ALTER TABLE users ADD COLUMN IF NOT EXISTS city VARCHAR(100) DEFAULT 'Москва';
            ALTER TABLE users ADD COLUMN IF NOT EXISTS timezone VARCHAR(50) DEFAULT 'Europe/Moscow';
            ALTER TABLE medications ADD COLUMN IF NOT EXISTS user_timezone VARCHAR(50) DEFAULT 'Europe/Moscow';
            CREATE TABLE IF NOT EXISTS notifications (
                id SERIAL PRIMARY KEY,
                user_id INTEGER REFERENCES users(id) ON DELETE CASCADE,
                type VARCHAR(20) NOT NULL,
                title VARCHAR(200) NOT NULL,
                body TEXT,
                is_read BOOLEAN DEFAULT FALSE,
                created_at TIMESTAMP DEFAULT NOW()
            );
            CREATE INDEX IF NOT EXISTS idx_notif_user ON notifications(user_id, created_at DESC);
            ALTER TABLE users ADD COLUMN IF NOT EXISTS tokens_used INTEGER DEFAULT 0;
            ALTER TABLE users ADD COLUMN IF NOT EXISTS tokens_limit INTEGER DEFAULT 50000;
            ALTER TABLE users ADD COLUMN IF NOT EXISTS is_admin BOOLEAN DEFAULT FALSE;
            CREATE TABLE IF NOT EXISTS site_visits (
                id          SERIAL PRIMARY KEY,
                path        VARCHAR(120),
                visited_at  TIMESTAMP DEFAULT NOW()
            );
            CREATE INDEX IF NOT EXISTS idx_visits_date ON site_visits(visited_at);
            CREATE TABLE IF NOT EXISTS app_errors (
                id          SERIAL PRIMARY KEY,
                source      VARCHAR(80),
                message     TEXT,
                created_at  TIMESTAMP DEFAULT NOW()
            );
            CREATE INDEX IF NOT EXISTS idx_errors_date ON app_errors(created_at DESC);
            ALTER TABLE users ADD COLUMN IF NOT EXISTS news_topics TEXT;
            ALTER TABLE user_profile ADD COLUMN IF NOT EXISTS alcohol INTEGER DEFAULT 0;
            ALTER TABLE user_profile ADD COLUMN IF NOT EXISTS smoking INTEGER DEFAULT 0;
            ALTER TABLE user_profile ADD COLUMN IF NOT EXISTS smoking_years INTEGER;
            ALTER TABLE user_profile ADD COLUMN IF NOT EXISTS marital_status TEXT;
            ALTER TABLE user_profile ADD COLUMN IF NOT EXISTS children TEXT;
            ALTER TABLE user_profile ADD COLUMN IF NOT EXISTS family_notes TEXT;
            ALTER TABLE loans ADD COLUMN IF NOT EXISTS monthly_payment REAL;
            ALTER TABLE finance_records ADD COLUMN IF NOT EXISTS subcategory VARCHAR(50);
            ALTER TABLE finance_records ADD COLUMN IF NOT EXISTS item_name TEXT;

            -- Колонки для лекарств с уведомлениями
            ALTER TABLE medications ADD COLUMN IF NOT EXISTS times_per_day INTEGER DEFAULT 1;
            ALTER TABLE medications ADD COLUMN IF NOT EXISTS times_json TEXT DEFAULT '[]';
            ALTER TABLE medications ADD COLUMN IF NOT EXISTS food_rule VARCHAR(50) DEFAULT 'any';
            ALTER TABLE medications ADD COLUMN IF NOT EXISTS food_minutes INTEGER;
            ALTER TABLE medications ADD COLUMN IF NOT EXISTS notify BOOLEAN DEFAULT FALSE;
            ALTER TABLE medications ADD COLUMN IF NOT EXISTS taken_today TEXT DEFAULT '[]';
            ALTER TABLE medications ADD COLUMN IF NOT EXISTS taken_date DATE;
        """)

        # Insert default feature flags
        await conn.execute("""
            INSERT INTO features (name, enabled, description) VALUES
                ('health_module',  FALSE, 'Раздел Моё здоровье'),
                ('finance_module', FALSE, 'Раздел Мои финансы'),
                ('telegram_bot',   FALSE, 'Напоминания через Telegram'),
                ('max_bot',        FALSE, 'Напоминания через Макс'),
                ('subscription',   FALSE, 'Платная подписка ЮKassa'),
                ('new_nav',        FALSE, 'Новая навигация v2')
            ON CONFLICT (name) DO NOTHING;
        """)

# ── Auth helpers ─────────────────────────────────────────────────
def hash_password(password: str) -> str:
    return hashlib.sha256((password + SECRET_KEY).encode()).hexdigest()

def generate_token() -> str:
    return secrets.token_urlsafe(32)

async def get_current_user(authorization: str = Header(None)):
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Не авторизован")
    token = authorization.split(" ")[1]
    if not db_pool:
        raise HTTPException(status_code=503, detail="БД недоступна")
    async with db_pool.acquire() as conn:
        session = await conn.fetchrow(
            "SELECT user_id FROM sessions WHERE token=$1 AND expires_at > NOW()", token
        )
        if not session:
            raise HTTPException(status_code=401, detail="Токен истёк или неверный")
        user = await conn.fetchrow("SELECT * FROM users WHERE id=$1", session["user_id"])
        return dict(user)


async def get_admin_user(authorization: str = Header(None)):
    """Только для администратора (владельца проекта)."""
    user = await get_current_user(authorization)
    if not user.get("is_admin"):
        raise HTTPException(status_code=403, detail="Доступ только для администратора")
    return user

# ── Models ───────────────────────────────────────────────────────
class RegisterRequest(BaseModel):
    phone: str
    name: str
    password: str
    marketing_consent: bool = False

class LoginRequest(BaseModel):
    phone: str
    password: str

class ChatRequest(BaseModel):
    message: str
    history: Optional[List] = []
    system: Optional[str] = None
    mode: Optional[str] = None

class AnalyzeRequest(BaseModel):
    text: str
    context: Optional[str] = None

class HealthRecordRequest(BaseModel):
    type: str        # pressure / pulse / sugar / note
    value_1: Optional[float] = None
    value_2: Optional[float] = None
    comment: Optional[str] = None

class MedicationRequest(BaseModel):
    name: str
    dose: str
    schedule: str
    start_date: Optional[str] = None
    end_date: Optional[str] = None

class FinanceRecordRequest(BaseModel):
    category: str   # pharmacy / shop / utility / credit / other
    amount: float
    comment: Optional[str] = None

# ── AI helper ────────────────────────────────────────────────────
SYSTEM_PROMPT = """Ты помощник OkTolk. Говори просто, без сложных слов.
Каждый шаг с новой строки, нумеруй: 1. 2. 3.
Максимум 3-4 шага в ответе.
НЕ используй ** ## --- и другую разметку.
Если спрашивают про мошенников: СТОП! Это мошенники! Не отвечайте им!
Заканчивай: Если не получилось - напишите мне снова!"""

HEALTH_PROMPT = """Ты — грамотный помощник по здоровью в приложении OkTolk.
Отвечай как знающий человек, который умеет объяснять сложное простым языком — живо, по-человечески, без болтовни.

Правила:
— Коротко и по делу: 2-5 предложений для простых вопросов.
— Объясняй понятно, как образованный собеседник, а не сухая инструкция.
— Если перечисляешь — каждый пункт с новой строки, начинай с «— ».
— Не ставь диагнозы, не назначай лекарства и дозировки.
— При тревожных симптомах мягко советуй обратиться к врачу.
— Структурируй ответ для удобного чтения: короткие абзацы с пустой строкой между смысловыми блоками; списки — каждый пункт с новой строки через «— »; важное можно выделить **жирным**.
— Опирайся на проверенную медицинскую информацию.

Если предоставлены данные пользователя — обязательно учитывай их:
— Рабочее давление — это его личная норма, сравнивай измерения с ней, а не только с общей нормой.
— Учитывай возраст, вредные привычки, физическую активность при оценке показателей.
— Если пользователь курит или употребляет алкоголь — можешь мягко упомянуть связь с показателями, без нравоучений.
— Профессию и хобби используй чтобы ответ был ближе к его жизни (например, «для водителя важно следить за давлением»).
— Не пересказывай все данные пользователя обратно — просто используй их в ответе."""

FINANCE_CHAT_PROMPT = """Ты — финансовый помощник пользователя OkTolk. За плечами 20 лет в управлении семейными финансами: вёл бюджеты сотен семей, помогал выбираться из долгов, оптимизировать расходы, копить на цели. Знаешь как устроены кредиты в РФ, льготы, налоговые вычеты, рефинансирование, способы дополнительного заработка для разных возрастов.

ТВОЙ ТОН: разговариваешь как опытный друг — без сухости, без воды, без банковского жаргона. Не «оптимизируйте структуру расходов», а «слушай, у тебя на кафе уходит как на ЖКУ — давай разберёмся».

КАК ТЫ ДУМАЕШЬ ПЕРЕД ОТВЕТОМ (про себя, не показывай):
1. Что человек реально хочет узнать — цифру, совет, успокоение?
2. Какие у меня есть данные о нём (доход, расходы, кредиты)? Использую их.
3. Если данных мало для точного ответа — говорю что есть и что нужно уточнить.
4. Есть ли красный флаг? (долг > 50% дохода, расходы > дохода, импульсивные траты)

ПРАВИЛА ОТВЕТА:
- Сразу к сути. Цифры конкретные, в рублях, округлённые («примерно 12 тысяч», не «12 487,33 руб»).
- Если знаешь доход — считай проценты («это четверть твоего дохода»).
- Если знаешь кредитную нагрузку — упоминай её только когда это к месту.
- Совет всегда конкретный («перенеси платёж на 5-е число — будет проще планировать»), а не общий («больше экономьте»).
- Если есть способ заработать в его ситуации — назови (вычет НДФЛ, продажа ненужного на Авито, льготы по возрасту, самозанятость).
- Заканчивай вопросом или предложением действия — чтобы разговор продолжался.

ЧТО НЕ ДЕЛАЕШЬ:
- Не советуешь инвестиции и конкретные банки (это к финансовому консультанту).
- Не лечишь моралью («надо меньше тратить на удовольствия»).
- Не пугаешь — даже плохие новости подаёшь с планом действий.

СТРУКТУРА ОТВЕТА (RichText красиво её покажет):
- Короткие абзацы с пустой строкой между смыслами.
- Списки — каждый пункт с новой строки через «— ».
- Важное число или вывод — выделяй **жирным**.
- Объём: 3-7 предложений на простой вопрос, до 12 на сложный."""

SCAM_PROMPT = """Ты — помощник по защите от мошенников в приложении OkTolk.
Пользователь спрашивает про подозрительный звонок, сообщение, ссылку или ситуацию.

Правила:
- Сразу по делу: есть ли признаки мошенничества и что делать.
- Если ситуация явно опасна — начни с «СТОП! Похоже на мошенников.»
- Дай 2-3 конкретных шага простым языком.
- Никогда не советуй сообщать коды из СМС, переводить деньги «для проверки», устанавливать программы по просьбе звонящих.
- Структурируй ответ для удобного чтения: короткие абзацы с пустой строкой между смысловыми блоками; списки — каждый пункт с новой строки через «— »; важное можно выделить **жирным**.
- Если нужно проверить конкретную ссылку или скриншот — подскажи открыть раздел «Стоп Мошенник»."""

NEWS_CHAT_PROMPT = """Ты — помощник OkTolk. Пользователь спрашивает про новости, законы, льготы или события в стране.
Ответь коротко и понятно простым языком (2-4 предложения). Структурируй ответ для удобного чтения: короткие абзацы с пустой строкой между смысловыми блоками; списки — каждый пункт с новой строки через «— »; важное можно выделить **жирным**.
Если для точного ответа нужны свежие новости — подскажи открыть раздел «Новости», где лента подбирается по темам пользователя."""

SEARCH_CHAT_PROMPT = """Ты — помощник OkTolk. Пользователь хочет что-то найти: товар, услугу, место или мероприятие.
Ответь коротко и по делу простым языком. Структурируй ответ для удобного чтения: короткие абзацы с пустой строкой между смысловыми блоками; списки — каждый пункт с новой строки через «— »; важное можно выделить **жирным**.
Если нужен реальный поиск с актуальными вариантами — подскажи открыть раздел «Поиск»."""

SYSTEM_PROMPTS = {
    "health": HEALTH_PROMPT,
    "finance": FINANCE_CHAT_PROMPT,
    "scam": SCAM_PROMPT,
    "news": NEWS_CHAT_PROMPT,
    "search": SEARCH_CHAT_PROMPT,
}

# Специализированные эксперты по категориям финансов — подключаются когда mode='finance:<category>'
FINANCE_UTILITY_EXPERT = """Ты — эксперт по ЖКХ с 15-летним опытом работы в управляющих компаниях и расчётно-кассовых центрах. Знаешь устройство квитанции до запятой: содержание жилья, ХВС/ГВС, водоотведение, электроэнергия по двух/трёхтарифному, отопление, газ, ОДН/КР на СОИ, капремонт, домофон. Знаешь как считаются нормативы, тарифы 2025-2026 по регионам РФ, как делается перерасчёт, как читать ОДПУ/ИПУ.

Что ты умеешь подсказывать пользователю:
- Какая статья платежа реально может быть завышена (типичные места: ОДН на электричество, отопление при разнице с соседями, ХВС без счётчика).
- Куда обратиться: УК, ГЖИ (государственная жилищная инспекция), прокуратура — что делать в какой ситуации.
- Льготы: субсидия на оплату ЖКУ если платёж >22% дохода семьи (в ряде регионов 10-15%), льготы ветеранам, инвалидам, многодетным, малоимущим. Компенсация 50% на капремонт для одиноких пенсионеров от 70 лет, 100% — от 80.
- Перерасчёт за отсутствие (>5 дней) — какие документы нужны.
- Установка счётчиков — когда окупается, когда нет.
- Сезонные особенности: летом меньше отопления, осенью обычно пересчитывают тариф.

Когда отвечаешь:
- Сначала разбери конкретные цифры пользователя — за что больше всего платит, что выглядит подозрительно для его типа жилья.
- Дай 1-2 конкретных действия (не «обратитесь в УК», а «напиши заявление в УК — образец есть на ГИС ЖКХ — на перерасчёт за…»).
- Если видишь право на льготу — назови её и куда подавать.

Тон: опытный сосед, который сам через всё прошёл. «Слушай», «смотри по цифрам», «знаешь что я бы сделал».
Структура: короткие абзацы с пустой строкой, списки через «— », важные числа и названия — **жирным**. 5-10 предложений."""

FINANCE_TRANSPORT_EXPERT = """Ты — эксперт по транспортным расходам с опытом логиста и автомобилиста. Знаешь экономику владения автомобилем (топливо, ТО, страховка, налог, износ), стоимость такси vs общественный транспорт vs каршеринг в разных городах РФ, тарифы на проездные (Тройка, Подорожник, Ярко и др.), льготы пенсионерам и школьникам.

Что подсказываешь:
- Если основная статья — такси: посчитай сколько выходит проездной на месяц, окупается ли при текущем числе поездок.
- Если бензин: посмотри расход на 100 км, средняя цена топлива в регионе, есть ли смысл в карте лояльности АЗС, разница 92/95 для конкретной машины.
- Если парковка: платные зоны vs перехватывающие, льготы для жителей района, абонемент.
- ТО/ремонт: при крупной сумме — стоит ли держать машину или продать (грубо: если содержание > 30% стоимости машины в год — продавай).
- Льготы на проезд: пенсионеры (бесплатный проезд в Москве, скидки в регионах), школьники, многодетные.

Когда отвечаешь:
- Сначала факт по его цифрам, потом конкретная оптимизация с расчётом в рублях.
- Совет конкретный: не «пользуйтесь меньше такси», а «у тебя 18 поездок такси в месяц по 400₽ — это 7200₽. Месячный проездной 2570₽. Экономия 4,5 тысячи».

Тон: «слушай», «смотри», «давай посчитаем».
Структура: абзацы с пустой строкой, списки через «— », числа **жирным**. 5-9 предложений."""

FINANCE_PHARMACY_EXPERT = """Ты — эксперт по затратам на медицину и лекарства. 10 лет работал провизором, знаешь дженерики, систему льготных лекарств, налоговый вычет на лечение, ОМС/ДМС, диспансеризацию.

Что подсказываешь:
- Дженерики: к дорогому препарату — название действующего вещества и 1-2 более дешёвых аналога (без советов по лечению, только информация о существовании аналогов; «спроси врача можно ли заменить»).
- Налоговый вычет 13% за лечение и лекарства по рецепту — реально вернуть. Лимит 150 000 руб/год расходов. Документы: справка из клиники для налоговой, чеки, рецепты по форме 107-1/у.
- Льготные категории: пенсионеры по ряду заболеваний, инвалиды, беременные, дети до 3 лет — список лекарств бесплатно/со скидкой 50%.
- Программа «Лекарства от давления, инфаркта и инсульта» — для перенёсших ССЗ в течение 2 лет бесплатные препараты.
- ДМС: если работает — что входит, как использовать.
- Бесплатная диспансеризация по ОМС — 1 раз в 3 года до 40 лет, ежегодно после 40. Включает анализы, которые иначе стоят 5-10 тыс.

Когда отвечаешь:
- Посмотри что покупает регулярно — есть ли смысл искать аналог, льготу.
- Назови сумму потенциального вычета НДФЛ (13% от трат на медицину).
- Если видишь хронические покупки одного препарата — упомяни возможность льготы по диагнозу.

НЕ ставишь диагнозы, НЕ заменяешь препараты, НЕ рекомендуешь схемы лечения. Только финансовая сторона.

Тон: «слушай», «знаешь что», «давай посмотрим что можно сэкономить».
Структура: абзацы, списки через «— », суммы и названия — **жирным**. 5-10 предложений."""

FINANCE_SHOP_EXPERT = """Ты — эксперт по семейному бюджету и оптимизации покупок. Знаешь как устроены сетевые магазины, программы лояльности, акции, маркетплейсы, и как реально (а не на словах) экономить на ежедневных тратах.

Что подсказываешь:
- Если основная масса — продукты: средний чек, частота, есть ли смысл в больших закупках раз в неделю vs ежедневные. Сезонность овощей-фруктов. Карты лояльности (Пятёрочка X5, Магнит, Перекрёсток) — реальная скидка 3-7%.
- Маркетплейсы (Wildberries, Ozon, Яндекс.Маркет) — кэшбек 1-15%, спец-цены подписчикам Премиума.
- Кэшбек-карты — Тинькофф, Альфа, Сбер: 5% на продукты или категорию месяца = реально 2-5 тыс/мес назад.
- Если одежда/техника — упомяни Авито за б/у в хорошем состоянии (особенно техника, инструменты), сезонные распродажи (Чёрная пятница, Январские, межсезонье).
- Семейные подходы: общий список перед походом → меньше импульсивных покупок, готовая еда на неделю → дешевле и здоровее.

Когда отвечаешь:
- Посмотри топ-товары и частоту — что покупает регулярно.
- Дай конкретный расчёт экономии в рублях, не общими словами.
- Если есть способ заработать (продать ненужное на Авито, кэшбек) — назови сумму.

Тон: «смотри по цифрам», «знаешь что у тебя выйдет», как друг который сам всё это делает.
Структура: абзацы с пустой строкой, списки через «— », суммы **жирным**. 5-9 предложений."""

FINANCE_CREDIT_EXPERT = """Ты — эксперт по потребительскому кредитованию и долгам с 15-летним опытом. Знаешь рынок кредитов РФ (ставки, банки, продукты), рефинансирование, реструктуризацию, кредитные каникулы (ФЗ 106-ФЗ и 348-ФЗ), банкротство физлиц, как работают коллекторы и закон 230-ФЗ.

КЛЮЧЕВЫЕ ОРИЕНТИРЫ:
- Безопасная кредитная нагрузка — до 30% от дохода. 30-50% — настороже. >50% — критическая зона, риск долговой ямы.
- Свободные деньги (доход − обяз. расходы − платежи) должны быть положительными, иначе долги растут.

ЧТО ПОДСКАЗЫВАЕШЬ:
- Если нагрузка <30% — спокойствие, но «копейка рубль бережёт»: можно ли досрочно гасить (% уменьшается).
- 30-50% — конкретные шаги: какой кредит закрывать первым (метод снежного кома или метод лавины: снежный — самый маленький долг закрывают первым для психологического эффекта; лавина — самый дорогой по ставке, выгоднее по деньгам).
- >50% — экстренный план: переговоры с банком о реструктуризации, рефинансирование под меньший % (если кредитная история ОК), кредитные каникулы (трудная жизненная ситуация — потеря дохода >30%), при невозможности — банкротство (>500 тыс долга, выгодно только если совсем безвыходно).
- Никогда не советуй брать новый кредит чтобы закрыть старый «просто так» — только осознанное рефинансирование.
- Стоит ли брать новый кредит: смотри свободные деньги, цель, можно ли подождать или копить.

Способы выйти из долгов:
- Дополнительный заработок: подработка, фриланс, продажа ненужного, репетиторство, самозанятость.
- Возврат налогов: НДФЛ-вычеты (имущественный 260 тыс с покупки жилья, лечение, обучение).
- Сокращение крупных статей: пересмотр ОСАГО/КАСКО, отказ от ненужных подписок.

Тон: серьёзный, но не пугающий. «Слушай, ситуация такая», «давай по шагам». Если нагрузка критическая — мягко но прямо, и сразу с планом действий.
Структура: абзацы, нумерованные шаги, суммы и %% — **жирным**. 6-12 предложений."""

FINANCE_LEISURE_EXPERT = """Ты — эксперт по дискреционным тратам (досуг, развлечения, кафе, подписки) и психологии трат. Знаешь как работают подписочные сервисы (Яндекс Плюс, Сбер Прайм, IVI, Netflix-аналоги), кэшбек на досуг, как ловят на «копилке мелких трат».

Что подсказываешь:
- Подписки — частая утечка: посмотри на дублирующиеся (3 онлайн-кинотеатра одновременно), забытые автопродления. Назови сумму в год если оставить только нужное.
- Кафе: посчитай средний чек × частоту = в месяц. Не моралью, а фактом: «11 тысяч за месяц». Альтернативы: домашний кофе с термокружкой, бизнес-ланч вместо вечера, готовые кофейные капсулы.
- Кино/выставки: льготные дни (понедельник-четверг в большинстве сетей), Пушкинская карта для 14-22 лет, скидки пенсионерам.
- Спорт: годовой абонемент в районный ФОК часто 2-4 тыс/год vs модные сети 3-6 тыс/мес.
- Психология: правило 24 часов на нерациональные покупки, «потратил 5000 = заработай 5000 в подработке мысленно».

Когда отвечаешь:
- Не лечишь моралью. Удовольствия — норма, важно осознанность.
- Конкретный расчёт: «если из 4 подписок оставишь 1 — это **3200₽ в месяц** или **38 тыс в год**».
- Если видишь импульсивные траты по паттерну — назови мягко.

Тон: «слушай», «прикинь», как друг без морализаторства.
Структура: абзацы, списки через «— », суммы **жирным**. 5-9 предложений."""

# Постфикс ко всем экспертным промптам: правила работы с актуальными данными
_EXPERT_RESEARCH_RULES = """

КАК ИСПОЛЬЗОВАТЬ АКТУАЛЬНЫЕ ДАННЫЕ:
Если в твой системный контекст подгрузили блок «АКТУАЛЬНЫЕ ДАННЫЕ ИЗ ИНТЕРНЕТА» — это свежая информация из проверенных источников. Опирайся на неё, а не на общие знания. В конце ответа кратко укажи источники в формате: «По данным gosuslugi.ru, cbr.ru». Если достоверность 'low' — добавь честную оговорку «лучше уточнить у [название источника] или по горячей линии».
Если актуальных данных нет, а вопрос требует точных цифр (тариф, ставка, лимит) — честно скажи: «точную цифру лучше уточнить у [источник]», и дай общую логику ответа."""

# Карта: категория → специализированный промпт
FINANCE_CATEGORY_PROMPTS = {
    "utility": FINANCE_UTILITY_EXPERT + _EXPERT_RESEARCH_RULES,
    "transport": FINANCE_TRANSPORT_EXPERT + _EXPERT_RESEARCH_RULES,
    "pharmacy": FINANCE_PHARMACY_EXPERT + _EXPERT_RESEARCH_RULES,
    "shop": FINANCE_SHOP_EXPERT + _EXPERT_RESEARCH_RULES,
    "credit": FINANCE_CREDIT_EXPERT + _EXPERT_RESEARCH_RULES,
    "leisure": FINANCE_LEISURE_EXPERT + _EXPERT_RESEARCH_RULES,
}

# Агенты-аналитики — работают с реальными данными пользователя
FINANCE_ANALYST_PROMPT = """Ты — финансовый аналитик-друг пользователя OkTolk. 20 лет помогаешь людям разобраться куда уходят деньги и где их можно вернуть в семейный кошелёк. Видишь паттерны там, где человек видит хаос.

Перед тобой РЕАЛЬНЫЕ траты пользователя из его базы — не выдумывай и не используй средние по стране. Опирайся только на эти цифры.

КАК ТЫ ЧИТАЕШЬ ДАННЫЕ (про себя):
1. Где основная масса денег — назови 1-2 категории с конкретными суммами.
2. Есть ли резкое — категория сильно больше других, разовая крупная трата, частые мелкие, которые суммарно выходят дорого.
3. Сравни с прошлым периодом если данные есть («стало больше на треть»).
4. Если есть доход — посчитай долю расходов от дохода в %.
5. Какой паттерн виден — много мелких трат на кафе, частые такси вместо метро, дублирующиеся подписки, импульсивные покупки.

ЧТО ДАЁШЬ ПОЛЬЗОВАТЕЛЮ:
- **Главный вывод** — на что уходит больше всего, с цифрой и % от дохода.
- **Что заметил** — паттерн или красный флаг (импульсивные траты, рост категории, превышение бюджета).
- **Дельный совет** — конкретный, под его данные. Не «экономьте на кафе», а «если из 8 кофе в неделю оставить 3 — за месяц вернёшь 4 тысячи». Если видишь способ заработать (продать ненужное, налоговый вычет, льгота по возрасту) — назови.
- **Вопрос в конце** — чтобы продолжить разбор: «Хочешь посмотреть транспорт?», «Сравнить с прошлым месяцем?», «Разобрать кафе подробнее?».

ТОН: друг, который реально разобрался. «Слушай», «смотри что получается», «дам совет». Никакой сухости. Никаких морализаторских ноток.

СТРУКТУРА (RichText покажет красиво):
- Короткие абзацы с пустой строкой между ними.
- Списки — каждый пункт с новой строки через «— ».
- Главные цифры — **жирным**.
- Объём 5-9 предложений. Не растекайся."""

HEALTH_ANALYST_PROMPT = """Ты — помощник по здоровью OkTolk. Перед тобой реальные показатели пользователя.

Твоя задача:
1. Оцени динамику показателя — растёт, стабилен или снижается.
2. Сравни с нормой, а если задано рабочее давление пользователя — сравнивай и с ним.
3. Отметь если есть тревожные значения, но мягко.
4. В КОНЦЕ задай один уточняющий вопрос ИЛИ предложи что ещё посмотреть — например: «Хотите посмотреть пульс?», «Показать за более долгий период?».

Правила: не ставь диагнозы, не назначай лекарства. Тёплый простой язык. Структурируй ответ для удобного чтения: короткие абзацы с пустой строкой между смысловыми блоками; списки — каждый пункт с новой строки через «— »; важное можно выделить **жирным**.
При тревожных показателях мягко советуй обратиться к врачу. Коротко: 4-7 предложений."""

# Ключевые слова намерения «проанализировать данные пользователя»
_ANALYSIS_KW = ("проанализир", "анализ", "разбор", "статистик", "отчёт", "отчет",
                "динамик", "потрат", "трач", "уход", "ушло", "больше всего",
                "мои расход", "мои трат", "куда ден", "на что ден",
                "как показат", "что с давлен", "что с весом", "что с сахар",
                "что с пульс", "как давлен", "как вес", "как сахар", "показател")

def call_ai(messages, model="deepseek-chat"):
    # Основной: DeepSeek напрямую (дешевле в разы)
    if DEEPSEEK_API_KEY:
        try:
            headers = {"Authorization": f"Bearer {DEEPSEEK_API_KEY}", "Content-Type": "application/json"}
            data = {"model": "deepseek-chat", "messages": messages, "max_tokens": 800, "temperature": 0.7}
            response = requests.post("https://api.deepseek.com/v1/chat/completions", headers=headers, json=data, timeout=30)
            if response.status_code == 200:
                j = response.json()
                u = j.get("usage", {})
                print(f"[AI] DeepSeek ok | in={u.get('prompt_tokens','?')} out={u.get('completion_tokens','?')} tokens")
                return j["choices"][0]["message"]["content"]
            else:
                print(f"[AI] DeepSeek HTTP {response.status_code} → fallback AItunnel")
        except Exception as e:
            print(f"[AI] DeepSeek error: {e} → fallback AItunnel")
    # Fallback: AItunnel
    if AITUNNEL_API_KEY:
        try:
            headers = {"Authorization": f"Bearer {AITUNNEL_API_KEY}", "Content-Type": "application/json"}
            data = {"model": "deepseek-v4-flash", "messages": messages, "max_tokens": 800, "temperature": 0.7}
            response = requests.post(f"{AITUNNEL_BASE_URL}chat/completions", headers=headers, json=data, timeout=30)
            if response.status_code == 200:
                print(f"[AI] AItunnel fallback ok (DeepSeek недоступен — тратится баланс AItunnel)")
                return response.json()["choices"][0]["message"]["content"]
            else:
                print(f"[AI] AItunnel HTTP {response.status_code}")
        except Exception as e:
            print(f"[AI] AItunnel error: {e}")
    raise Exception("AI недоступен: нет рабочих API ключей")

def _ds_headers():
    return {"Authorization": f"Bearer {DEEPSEEK_API_KEY}", "Content-Type": "application/json"}

def _ds_post(messages, system=None, max_tokens=800):
    """Прямой вызов DeepSeek API с fallback на AItunnel"""
    msgs = []
    if system:
        msgs.append({"role": "system", "content": system})
    msgs.extend(messages)
    # DeepSeek прямой
    if DEEPSEEK_API_KEY:
        try:
            r = requests.post(
                "https://api.deepseek.com/v1/chat/completions",
                headers=_ds_headers(),
                json={"model": "deepseek-chat", "messages": msgs, "max_tokens": max_tokens, "temperature": 0.7},
                timeout=30
            )
            if r.status_code == 200:
                j = r.json()
                u = j.get("usage", {})
                print(f"[AI/ds_post] DeepSeek ok | in={u.get('prompt_tokens','?')} out={u.get('completion_tokens','?')} tokens")
                return j["choices"][0]["message"]["content"].strip()
            else:
                print(f"[AI/ds_post] DeepSeek HTTP {r.status_code} → fallback AItunnel")
        except Exception as e:
            print(f"[AI/ds_post] DeepSeek error: {e} → fallback AItunnel")
    # Fallback AItunnel
    if AITUNNEL_API_KEY:
        print("[AI/ds_post] AItunnel fallback (DeepSeek недоступен)")
        r = requests.post(
            f"{AITUNNEL_BASE_URL}chat/completions",
            headers={"Authorization": f"Bearer {AITUNNEL_API_KEY}", "Content-Type": "application/json"},
            json={"model": "deepseek-v4-flash", "messages": msgs, "max_tokens": max_tokens, "temperature": 0.7},
            timeout=30
        )
        return r.json()["choices"][0]["message"]["content"].strip()
    raise Exception("AI недоступен")

# ── Лимиты токенов ───────────────────────────────────────────────
async def check_token_limit(user_id) -> bool:
    """True если пользователь в пределах лимита. При ошибке/отсутствии user_id — не блокируем."""
    if not user_id or not db_pool:
        return True
    try:
        async with db_pool.acquire() as conn:
            row = await conn.fetchrow(
                "SELECT tokens_used, tokens_limit FROM users WHERE id=$1", user_id
            )
        if not row:
            return True
        used = row["tokens_used"] or 0
        limit = row["tokens_limit"] or 50000
        return used < limit
    except Exception as e:
        print(f"[tokens] check error: {e}")
        return True

async def add_tokens(user_id, prompt_text: str, reply_text: str):
    """Приблизительный инкремент (1 токен ≈ 2 символа кириллицы, с запасом)."""
    if not user_id or not db_pool:
        return
    try:
        est = (len(prompt_text or "") + len(reply_text or "")) // 2
        if est <= 0:
            return
        async with db_pool.acquire() as conn:
            await conn.execute(
                "UPDATE users SET tokens_used = COALESCE(tokens_used,0) + $1 WHERE id=$2",
                est, user_id
            )
    except Exception as e:
        print(f"[tokens] add error: {e}")


async def save_chat_message(user_id, domain, role, text, action=None):
    """Сохраняет сообщение в единую историю диалогов (для дублирования в разделах)."""
    if not db_pool or not user_id:
        return
    try:
        async with db_pool.acquire() as conn:
            await conn.execute(
                "INSERT INTO chat_messages (user_id, domain, role, text, action) VALUES ($1,$2,$3,$4,$5)",
                user_id, domain, role, (text or "")[:4000], action
            )
    except Exception as e:
        print(f"[chat_history] save error: {e}")

# ── API v1 ───────────────────────────────────────────────────────

# Auth
@app.post("/api/v1/auth/register")
async def register(req: RegisterRequest):
    if not db_pool:
        raise HTTPException(status_code=503, detail="БД недоступна")
    async with db_pool.acquire() as conn:
        existing = await conn.fetchrow("SELECT id FROM users WHERE phone=$1", req.phone)
        if existing:
            raise HTTPException(status_code=400, detail="Телефон уже зарегистрирован")
        user_id = await conn.fetchval(
            "INSERT INTO users (phone, name, password_hash, marketing_consent) VALUES ($1,$2,$3,$4) RETURNING id",
            req.phone, req.name, hash_password(req.password), req.marketing_consent
        )
        token = generate_token()
        expires = datetime.now() + timedelta(days=30)
        await conn.execute(
            "INSERT INTO sessions (token, user_id, expires_at) VALUES ($1,$2,$3)",
            token, user_id, expires
        )
        return {"token": token, "user_id": user_id, "name": req.name, "tariff": "demo"}

@app.post("/api/v1/auth/login")
async def login(req: LoginRequest):
    if not db_pool:
        raise HTTPException(status_code=503, detail="БД недоступна")
    async with db_pool.acquire() as conn:
        user = await conn.fetchrow(
            "SELECT * FROM users WHERE phone=$1 AND password_hash=$2",
            req.phone, hash_password(req.password)
        )
        if not user:
            raise HTTPException(status_code=401, detail="Неверный телефон или пароль")
        token = generate_token()
        expires = datetime.now() + timedelta(days=30)
        await conn.execute(
            "INSERT INTO sessions (token, user_id, expires_at) VALUES ($1,$2,$3)",
            token, user["id"], expires
        )
        return {"token": token, "user_id": user["id"], "name": user["name"], "tariff": user["tariff"]}

@app.get("/api/v1/auth/me")
async def me(user=Depends(get_current_user)):
    return {"id": user["id"], "name": user["name"], "phone": user["phone"], "tariff": user["tariff"]}

@app.post("/api/v1/auth/demo")
async def auth_demo():
    """Создать анонимного демо-пользователя без персональных данных и выдать токен"""
    if not db_pool:
        raise HTTPException(status_code=503, detail="БД недоступна")

    # Гарантируем наличие колонки is_demo и расширяем phone (миграция)
    try:
        async with db_pool.acquire() as conn:
            await conn.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS is_demo BOOLEAN DEFAULT FALSE")
            await conn.execute("ALTER TABLE users ALTER COLUMN phone TYPE VARCHAR(40)")
    except Exception:
        pass

    # Создаём пользователя в чистом соединении (телефон <= 20 символов на всякий случай)
    demo_phone = "d_" + secrets.token_hex(6)
    async with db_pool.acquire() as conn:
        user_id = await conn.fetchval(
            "INSERT INTO users (phone, name, password_hash, tariff, is_demo) VALUES ($1,$2,$3,$4,$5) RETURNING id",
            demo_phone, "Гость", "demo_no_password", "demo", True
        )
        token = generate_token()
        expires = datetime.now() + timedelta(days=90)
        await conn.execute(
            "INSERT INTO sessions (token, user_id, expires_at) VALUES ($1,$2,$3)",
            token, user_id, expires
        )
    return {"token": token, "user_id": user_id, "name": "Гость", "tariff": "demo", "is_demo": True}

# Feature Flags
@app.get("/api/v1/features")
async def get_features():
    if not db_pool:
        return {"features": []}
    async with db_pool.acquire() as conn:
        rows = await conn.fetch("SELECT name FROM features WHERE enabled=TRUE")
        return {"features": [r["name"] for r in rows]}

async def build_profile_context(user_id: int) -> str:
    """Строит текстовый контекст профиля пользователя для AI"""
    if not db_pool:
        return ""
    try:
        async with db_pool.acquire() as conn:
            p = await conn.fetchrow("SELECT * FROM user_profile WHERE user_id=$1", user_id)
            credits = await conn.fetch(
                "SELECT name, amount, monthly_payment, rate FROM credits WHERE user_id=$1", user_id)
            loans = await conn.fetch(
                "SELECT name, amount, due_date, monthly_payment FROM loans WHERE user_id=$1", user_id)

        parts = []

        # Личные данные (если профиль заполнен)
        if p:
            if p.get("age"):
                parts.append(f"Возраст: {p['age']} лет")
            if p.get("gender"):
                parts.append(f"Пол: {'мужской' if p['gender'] == 'м' else 'женский'}")
            if p.get("height") and p.get("weight"):
                parts.append(f"Рост: {p['height']} см, вес: {p['weight']} кг")
            elif p.get("height"):
                parts.append(f"Рост: {p['height']} см")
            elif p.get("weight"):
                parts.append(f"Вес: {p['weight']} кг")
            if p.get("profession"):
                parts.append(f"Профессия: {p['profession']}")
            if p.get("hobbies"):
                parts.append(f"Хобби и интересы: {p['hobbies']}")
            if p.get("work_pressure_1") and p.get("work_pressure_2"):
                parts.append(f"Рабочее давление: {p['work_pressure_1']}/{p['work_pressure_2']} мм рт.ст.")
            if p.get("work_pulse"):
                parts.append(f"Обычный пульс: {p['work_pulse']} уд/мин")
            if p.get("base_sugar"):
                parts.append(f"Базовый сахар: {p['base_sugar']} ммоль/л")
            alcohol_labels = ["не употребляет алкоголь", "употребляет раз в месяц или реже",
                              "употребляет 2-4 раза в месяц", "употребляет 2-3 раза в неделю",
                              "употребляет 4+ раза в неделю"]
            smoke_labels = ["не курит", "курит до 5 сигарет в день", "курит около 10 сигарет в день",
                            "курит около 15 сигарет в день", "курит 20+ сигарет в день"]
            alc = p.get("alcohol", 0) or 0
            smk = p.get("smoking", 0) or 0
            if alc > 0:
                parts.append(f"Алкоголь: {alcohol_labels[min(alc, 4)]}")
            if smk > 0:
                label = smoke_labels[min(smk, 4)]
                stazh = p.get("smoking_years")
                parts.append(f"Курение: {label}" + (f", стаж {stazh} лет" if stazh else ""))
            if alc == 0 and smk == 0:
                parts.append("Вредные привычки: нет")
            activity_map = {"низкая": "низкая физическая активность",
                            "средняя": "средняя физическая активность",
                            "высокая": "высокая физическая активность"}
            if p.get("activity"):
                parts.append(f"Физическая активность: {activity_map.get(p['activity'], p['activity'])}")
            if p.get("heredity"):
                parts.append(f"Наследственность: {p['heredity']}")
            if p.get("chronic"):
                parts.append(f"Хронические заболевания: {p['chronic']}")
            if p.get("income"):
                income = int(p['income'])
                parts.append(f"Ежемесячный доход: {income} руб.")
            # Семья
            if p.get("marital_status"):
                ms_map = {"single": "не женат/не замужем", "married": "в браке",
                          "divorced": "в разводе", "widowed": "вдовец/вдова",
                          "partner": "есть партнёр"}
                parts.append(f"Семейное положение: {ms_map.get(p['marital_status'], p['marital_status'])}")
            if p.get("children"):
                try:
                    kids = json.loads(p["children"]) if isinstance(p["children"], str) else p["children"]
                    if kids:
                        descs = []
                        for k in kids:
                            g = {"м": "мальчик", "ж": "девочка"}.get(k.get("gender", ""), "ребёнок")
                            age = k.get("age")
                            descs.append(f"{g} {age} лет" if age is not None else g)
                        parts.append(f"Дети ({len(kids)}): " + ", ".join(descs))
                except Exception:
                    pass
            if p.get("family_notes"):
                parts.append(f"О семье: {p['family_notes']}")

        # Кредиты — грузятся всегда, независимо от профиля
        if credits:
            total_payment = sum(r["monthly_payment"] or 0 for r in credits)
            credit_list = "; ".join(
                f"{r['name'] or 'Кредит'}: {int(r['amount'] or 0)} ₽, платёж {int(r['monthly_payment'] or 0)} ₽/мес"
                for r in credits
            )
            parts.append(f"Кредиты: {credit_list}. Суммарный платёж: {int(total_payment)} руб./мес.")
            # Кредитная нагрузка от дохода
            if p and p.get("income") and p["income"] > 0:
                load_pct = round(total_payment / p["income"] * 100)
                parts.append(f"Кредитная нагрузка: {load_pct}% от дохода")

        # Займы — грузятся всегда
        if loans:
            loan_list = "; ".join(
                f"{r['name'] or 'Займ'}: {int(r['amount'] or 0)} ₽" +
                (f" до {r['due_date']}" if r['due_date'] else "") +
                (f", платёж {int(r['monthly_payment'])} ₽/мес" if r['monthly_payment'] else "")
                for r in loans
            )
            parts.append(f"Займы: {loan_list}")

        if not parts:
            return ""
        return "\n\nДанные пользователя (используй для персонального ответа):\n" + \
               "\n".join(f"— {x}" for x in parts)
    except Exception:
        return ""

# Chat (с опциональной авторизацией — берём профиль если токен есть)
# Ключевые слова для мгновенной маршрутизации мошенничества (без вызова AI)
_SCAM_KW = ("мошенник", "обман", "развод на деньг", "подозрительн", "фишинг", "скам", "взлом")

async def classify_domain(message: str) -> str:
    """Оркестратор: определяет раздел запроса.
    Возвращает один из: health|finance|scam|news|search|general"""
    ql = (message or "").lower()
    # Быстрый shortcut для явных случаев — экономим вызов AI
    if any(kw in ql for kw in _SCAM_KW):
        return "scam"
    if not AITUNNEL_API_KEY:
        return "general"
    prompt = (
        "Ты маршрутизатор приложения OkTolk. Определи тему сообщения пользователя.\n"
        "Темы:\n"
        "health — здоровье, давление, пульс, сахар, вес, лекарства, самочувствие, симптомы\n"
        "finance — деньги, расходы, доходы, кредиты, бюджет, цены, экономия\n"
        "scam — мошенники, подозрительные звонки/смс/ссылки, обман\n"
        "news — новости, что нового, законы, льготы, события в стране\n"
        "search — найти товар/услугу/место, куда сходить, где купить, мероприятия\n"
        "general — всё остальное, общие вопросы, беседа\n\n"
        "Ответь ТОЛЬКО JSON без пояснений: {\"domain\":\"...\"}\n"
        f"Сообщение: {message}"
    )

    def _call():
        headers = {"Authorization": f"Bearer {AITUNNEL_API_KEY}", "Content-Type": "application/json"}
        data = {"model": "gemini-2.5-flash-lite",
                "messages": [{"role": "user", "content": prompt}], "max_tokens": 30}
        r = requests.post(f"{AITUNNEL_BASE_URL}chat/completions", headers=headers, json=data, timeout=10)
        return r.json()["choices"][0]["message"]["content"]

    try:
        raw = await asyncio.to_thread(_call)
        m = re.search(r'"domain"\s*:\s*"(\w+)"', raw)
        domain = m.group(1) if m else "general"
        valid = {"health", "finance", "scam", "news", "search", "general"}
        result = domain if domain in valid else "general"
        print(f"[orchestrator] '{message[:40]}' → {result}")
        return result
    except Exception as e:
        print(f"[orchestrator] error: {e} → general")
        return "general"


# Ключевые слова намерения записать данные
_RECORD_KW = ("запиши", "запомни", "добавь", "сохрани", "внес", "занеси",
              "зафиксируй", "записать", "отметь", "фиксируй", "учти расход", "плюсуй")

# Косвенные траты (только finance) — глаголы покупки. Срабатывают если есть сумма.
_SPEND_KW = ("купил", "купила", "куплю", "покупк", "потрат", "заплат", "оплат",
             "истрат", "обошлся", "обошёлся", "стоил", "ушло на", "ушло ", "взял за",
             "приобрёл", "приобрел", "расход")

# Ключевые слова info/research-запроса → подключаем Tavily
_RESEARCH_KW = ("закон", "льгот", "пособи", "выплат", "субсиди", "пенси",
                "как оформить", "как получить", "как подать", "документ",
                "новост", "что нового", "актуальн", "изменени", "правила",
                "положено ли", "имею право", "сколько стоит", "когда вступ")

async def extract_record(message: str, domain: str) -> dict:
    """Извлекает запись(и) из сообщения.
    Возвращает {ok:True, records:[...], label:"..."} или {ok:False}."""
    if not AITUNNEL_API_KEY or domain not in ("health", "finance"):
        return {"ok": False}

    if domain == "health":
        prompt = (
            "Пользователь хочет записать показатель(и) здоровья. Извлеки данные.\n"
            "Типы: pressure (давление, value_1=верхнее value_2=нижнее), "
            "pulse (пульс, value_1), sugar (сахар ммоль/л, value_1), weight (вес кг, value_1).\n"
            "Если несколько показателей — верни массив.\n"
            "Верни ТОЛЬКО JSON-массив: [{\"type\":\"pressure|pulse|sugar|weight\",\"value_1\":число,\"value_2\":число или null}]\n"
            "Если это не показатель здоровья — верни [].\n"
            f"Сообщение: {message}"
        )
    else:  # finance
        prompt = (
            "Пользователь хочет записать расход(ы). Извлеки ВСЕ расходы из сообщения.\n"
            "Категории (category):\n"
            "shop — магазины, продукты, одежда, техника, бытовое\n"
            "pharmacy — аптека, лекарства, врачи, анализы\n"
            "utility — ЖКУ, коммуналка, свет, вода, интернет\n"
            "credit — кредиты, займы, рассрочка\n"
            "transport — такси, автобус, метро, бензин, парковка, И БИЛЕТЫ на самолёт/поезд/автобус, авиабилеты, ж/д\n"
            "leisure — кафе, рестораны, кино, развлечения, подписки, спорт\n"
            "other — всё прочее\n"
            "ВАЖНО: если в сообщении несколько расходов — верни МАССИВ из нескольких объектов, каждый отдельно.\n"
            "Верни ТОЛЬКО JSON-массив: [{\"category\":\"код\",\"amount\":число,\"comment\":\"1-2 слова\"}]\n"
            "Если суммы нет — верни [].\n"
            f"Сообщение: {message}"
        )

    def _call():
        headers = {"Authorization": f"Bearer {AITUNNEL_API_KEY}", "Content-Type": "application/json"}
        data = {"model": "gemini-2.5-flash-lite",
                "messages": [{"role": "user", "content": prompt}], "max_tokens": 300}
        r = requests.post(f"{AITUNNEL_BASE_URL}chat/completions", headers=headers, json=data, timeout=15)
        return r.json()["choices"][0]["message"]["content"]

    try:
        raw = await asyncio.to_thread(_call)
        m = re.search(r'\[.*\]|\{.*\}', raw, re.DOTALL)
        if not m:
            return {"ok": False}
        arr = json.loads(m.group())
        if isinstance(arr, dict):
            arr = [arr]
        if not arr:
            return {"ok": False}

        if domain == "health":
            records, parts = [], []
            for d in arr:
                t = d.get("type")
                v1 = d.get("value_1")
                if not t or v1 is None:
                    continue
                try:
                    v1 = float(v1)
                    v2 = float(d["value_2"]) if d.get("value_2") is not None else None
                except (ValueError, TypeError):
                    continue
                records.append({"type": t, "value_1": v1, "value_2": v2})
                lbl = {
                    "pressure": f"давление {int(v1)}/{int(v2)}" if v2 else f"давление {int(v1)}",
                    "pulse": f"пульс {int(v1)}", "sugar": f"сахар {v1}", "weight": f"вес {v1} кг",
                }.get(t, f"{t} {v1}")
                parts.append(lbl)
            if not records:
                return {"ok": False}
            print(f"[extract_record] health → {'; '.join(parts)}")
            return {"ok": True, "records": records, "label": "; ".join(parts)}

        else:  # finance
            cat_ru = {"shop": "покупки", "pharmacy": "аптека", "utility": "ЖКУ",
                      "credit": "кредит", "transport": "транспорт", "leisure": "досуг", "other": "прочее"}
            records, parts = [], []
            for it in arr:
                amt = it.get("amount")
                if amt is None:
                    continue
                try:
                    amt = float(amt)
                except (ValueError, TypeError):
                    continue
                cat = it.get("category", "other") or "other"
                cm = it.get("comment", "") or ""
                records.append({"category": cat, "amount": amt, "comment": cm})
                parts.append(f"{int(amt)} ₽ {cat_ru.get(cat, cat)}" + (f" ({cm})" if cm else ""))
            if not records:
                return {"ok": False}
            print(f"[extract_record] finance → {'; '.join(parts)}")
            return {"ok": True, "records": records, "label": "; ".join(parts)}
    except Exception as e:
        print(f"[extract_record] error: {e}")
        return {"ok": False}


async def research_answer(query: str) -> str:
    """Tavily research → понятный ответ простым языком. Для info-вопросов (законы, правила, как оформить)."""
    if not TAVILY_API_KEY:
        return ""
    try:
        data = await asyncio.to_thread(tavily_search_adv, query, "general", None, 5, True)
        answer = (data.get("answer") or "").strip()
        items = data.get("results", []) or []
        if answer and items:
            top_url = items[0].get("url", "")
            if top_url:
                deep = await asyncio.to_thread(tavily_extract, top_url, 3000)
                if deep:
                    try:
                        enrich = await asyncio.to_thread(call_ai, [{"role": "user", "content": (
                            "Ответь на вопрос простым русским языком, коротко и по делу (3-6 предложений), "
                            "на основе краткого ответа и текста источника. Без markdown-разметки.\n\n"
                            f"ВОПРОС: {query}\n\nКраткий ответ: {answer}\n\n"
                            f"Текст источника:\n{deep[:2500]}"
                        )}])
                        if enrich and len(enrich.strip()) > 20:
                            answer = enrich.strip()
                    except Exception as e:
                        print(f"[research] enrich error: {e}")
        if answer and items:
            src = items[0].get("url", "")
            dom = src.split("/")[2].replace("www.", "") if src.startswith("http") else ""
            if dom:
                answer += f"\n\nИсточник: {dom}"
        print(f"[research] '{query[:40]}' → {'ответ ' + str(len(answer)) + ' симв.' if answer else 'пусто'}")
        return answer or ""
    except Exception as e:
        print(f"[research] error: {e}")
        return ""


async def finance_data_context(user_id) -> str:
    """Собирает реальные траты пользователя (30 дней + 7 дней) для агента-аналитика."""
    if not db_pool or not user_id:
        return ""
    cat_ru = {"shop": "Покупки/продукты", "pharmacy": "Аптека/медицина", "utility": "ЖКУ",
              "credit": "Кредиты", "transport": "Транспорт", "leisure": "Досуг", "other": "Прочее"}
    try:
        async with db_pool.acquire() as conn:
            month = await conn.fetch(
                "SELECT category, SUM(amount) t, COUNT(*) c FROM finance_records "
                "WHERE user_id=$1 AND recorded_at >= NOW() - INTERVAL '30 days' "
                "GROUP BY category ORDER BY t DESC", user_id)
            week = await conn.fetch(
                "SELECT category, SUM(amount) t FROM finance_records "
                "WHERE user_id=$1 AND recorded_at >= NOW() - INTERVAL '7 days' "
                "GROUP BY category ORDER BY t DESC", user_id)
        if not month:
            return "NO_DATA"
        total_m = sum(float(r["t"]) for r in month)
        lines = [f"Расходы за 30 дней (всего {int(total_m)} ₽):"]
        for r in month:
            lines.append(f"— {cat_ru.get(r['category'], r['category'])}: {int(r['t'])} ₽ ({r['c']} опер.)")
        if week:
            total_w = sum(float(r["t"]) for r in week)
            lines.append(f"\nЗа последние 7 дней (всего {int(total_w)} ₽):")
            for r in week:
                lines.append(f"— {cat_ru.get(r['category'], r['category'])}: {int(r['t'])} ₽")
        return "\n".join(lines)
    except Exception as e:
        print(f"[analyze finance] error: {e}")
        return ""


async def finance_category_data(user_id, category: str) -> str:
    """Собирает подробные данные по конкретной категории для специализированного эксперта."""
    if not db_pool or not user_id or not category:
        return ""
    cat_ru = {"shop": "Покупки/продукты", "pharmacy": "Аптека/медицина", "utility": "ЖКУ",
              "credit": "Кредиты", "transport": "Транспорт", "leisure": "Досуг", "other": "Прочее"}
    subcat_ru = {
        "taxi": "такси", "public": "общественный транспорт", "fuel": "топливо",
        "parking": "парковка", "service": "ТО/ремонт", "other_transport": "прочий транспорт",
        "drugs": "лекарства", "doctor": "врачи", "tests": "анализы", "hospital": "больница",
        "groceries": "продукты", "clothes": "одежда", "household": "бытовое",
        "cafe": "кафе/рестораны", "cinema": "кино/развлечения", "subscription": "подписки", "sport": "спорт",
    }
    try:
        async with db_pool.acquire() as conn:
            month_total = await conn.fetchval(
                "SELECT COALESCE(SUM(amount),0) FROM finance_records WHERE user_id=$1 AND category=$2 "
                "AND recorded_at >= date_trunc('month', NOW())", user_id, category) or 0
            prev_total = await conn.fetchval(
                "SELECT COALESCE(SUM(amount),0) FROM finance_records WHERE user_id=$1 AND category=$2 "
                "AND recorded_at >= date_trunc('month', NOW()) - interval '1 month' "
                "AND recorded_at < date_trunc('month', NOW())", user_id, category) or 0
            count = await conn.fetchval(
                "SELECT COUNT(*) FROM finance_records WHERE user_id=$1 AND category=$2 "
                "AND recorded_at >= date_trunc('month', NOW())", user_id, category) or 0
            subcats = await conn.fetch(
                "SELECT subcategory, SUM(amount) t, COUNT(*) c FROM finance_records "
                "WHERE user_id=$1 AND category=$2 AND recorded_at >= date_trunc('month', NOW()) "
                "AND subcategory IS NOT NULL GROUP BY subcategory ORDER BY t DESC", user_id, category)
            items = await conn.fetch(
                "SELECT item_name, COUNT(*) c, SUM(amount) t FROM finance_records "
                "WHERE user_id=$1 AND category=$2 AND item_name IS NOT NULL "
                "AND recorded_at >= date_trunc('month', NOW()) GROUP BY item_name "
                "ORDER BY c DESC, t DESC LIMIT 5", user_id, category)
            prof = await conn.fetchrow("SELECT income FROM user_profile WHERE user_id=$1", user_id)
            income = float(prof["income"]) if prof and prof["income"] else 0
        if not month_total and not count:
            return ""
        m_total = float(month_total)
        lines = [f"Категория: {cat_ru.get(category, category)}"]
        lines.append(f"Текущий месяц: {int(m_total)} ₽" +
                     (f", {round(m_total / income * 100)}% от дохода ({int(income)} ₽)" if income > 0 else ""))
        lines.append(f"Прошлый месяц: {int(prev_total)} ₽" +
                     (f" (изменение {'+' if m_total > prev_total else ''}{int(m_total - prev_total)} ₽)" if prev_total else ""))
        lines.append(f"Операций за месяц: {count}, средний чек: {int(m_total / count) if count else 0} ₽")
        if subcats:
            lines.append("Разбивка по подкатегориям:")
            for r in subcats:
                lines.append(f"— {subcat_ru.get(r['subcategory'], r['subcategory'])}: {int(r['t'])} ₽ ({r['c']} опер.)")
        if items:
            lines.append("Часто покупаемые позиции:")
            for r in items:
                lines.append(f"— {r['item_name']}: {r['c']} раз на {int(r['t'])} ₽")
        return "\n".join(lines)
    except Exception as e:
        print(f"[finance_category_data] error: {e}")
        return ""


# ═════════════════════════════════════════════════════════════════
# Expert Research — экспертный поиск с проверкой достоверности
# ═════════════════════════════════════════════════════════════════

# Whitelist доверенных доменов по темам (для проверки достоверности)
TRUSTED_DOMAINS = {
    "utility": [
        "gosuslugi.ru", "dom.gosuslugi.ru", "mintrudrf.ru", "minstroyrf.ru",
        "gji.mos.ru", "fas.gov.ru", "consultant.ru", "garant.ru",
    ],
    "credit": [
        "cbr.ru", "nalog.gov.ru", "government.ru", "banki.ru",
        "consultant.ru", "garant.ru", "duma.gov.ru", "pravo.gov.ru",
    ],
    "pharmacy": [
        "minzdrav.gov.ru", "rosminzdrav.ru", "grls.rosminzdrav.ru",
        "rspp.ru", "nalog.gov.ru", "consultant.ru",
    ],
    "transport": [
        "mintrans.gov.ru", "gibdd.ru", "transport.mos.ru",
        "rzd.ru", "gosuslugi.ru", "consultant.ru",
    ],
    "shop": [
        "rospotrebnadzor.ru", "fas.gov.ru", "nalog.gov.ru",
        "gosuslugi.ru", "consultant.ru",
    ],
    "leisure": [
        "kultura.gov.ru", "gosuslugi.ru", "culture.ru", "pushkinskayakarta.ru",
    ],
    "general": [
        "gosuslugi.ru", "government.ru", "pravo.gov.ru", "consultant.ru",
        "garant.ru", "cbr.ru", "nalog.gov.ru",
    ],
}

# Триггеры запроса актуальных внешних данных
_FRESH_DATA_KW = (
    "актуальн", "текущ", "свеж", "новый", "новые", "новая", "сейчас",
    "сегодня", "в 2025", "в 2026", "в этом году", "последн", "обновл",
    "тариф", "ставк", "норматив", "льгот", "пособи", "субсиди", "вычет",
    "закон", "постановлени", "поправк", "изменени",
)

# Простой in-memory кеш (час) — экономим лимит Tavily 1000/мес
_research_cache = {}  # key → (timestamp, result)
_RESEARCH_CACHE_TTL = 3600  # 1 час
_research_stats = {"calls": 0, "cache_hits": 0}


def _needs_external_data(message: str, category: str = None) -> bool:
    """Решает: нужны ли свежие внешние данные для качественного ответа."""
    ql = (message or "").lower()
    # Явные триггеры в запросе
    if any(kw in ql for kw in _FRESH_DATA_KW):
        return True
    # Числа года → точно про актуальное
    if any(y in ql for y in ("2025", "2026", "2027")):
        return True
    # В категориях с быстро меняющимися фактами — research чаще
    fast_changing = {"utility", "credit", "pharmacy"}
    if category in fast_changing and ("?" in ql or any(s in ql for s in ("как", "сколько", "какой", "можно ли", "положен"))):
        return True
    return False


async def expert_research(query: str, topic: str = "general") -> dict:
    """Экспертный поиск через Tavily с проверкой достоверности по whitelist.
    Возвращает {ok, summary, sources, confidence} или {ok:False}.

    confidence:
      'high'   — ≥2 доверенных источника подтверждают
      'medium' — 1 доверенный источник
      'low'    — недоверенные источники или противоречия
    """
    if not TAVILY_API_KEY or not query:
        return {"ok": False}

    # Кеш
    cache_key = f"{topic}::{query.lower().strip()[:200]}"
    now = int(__import__("time").time())
    cached = _research_cache.get(cache_key)
    if cached and now - cached[0] < _RESEARCH_CACHE_TTL:
        _research_stats["cache_hits"] += 1
        print(f"[expert_research] cache hit '{query[:40]}'")
        return cached[1]

    _research_stats["calls"] += 1
    print(f"[expert_research] '{query[:60]}' topic={topic} (calls={_research_stats['calls']}, hits={_research_stats['cache_hits']})")

    try:
        # 1. Поиск с приоритетом доверенных доменов
        trusted = TRUSTED_DOMAINS.get(topic, TRUSTED_DOMAINS["general"])
        data = await asyncio.to_thread(
            tavily_search_adv, query, "general", None, 8, True, trusted
        )
        items = data.get("results", []) or []
        answer = (data.get("answer") or "").strip()

        # 2. Если в доверенных пусто — пробуем без фильтра
        used_trusted = True
        if not items:
            print(f"[expert_research] нет в доверенных, ищу шире")
            used_trusted = False
            data = await asyncio.to_thread(
                tavily_search_adv, query, "general", None, 5, True
            )
            items = data.get("results", []) or []
            answer = (data.get("answer") or "").strip()

        if not items and not answer:
            result = {"ok": False}
            _research_cache[cache_key] = (now, result)
            return result

        # 3. Достоверность: считаем сколько результатов из доверенных
        def _dom(url):
            try:
                return url.split("/")[2].replace("www.", "")
            except Exception:
                return ""

        trusted_items = [it for it in items if any(td in _dom(it.get("url", "")) for td in TRUSTED_DOMAINS.get(topic, []))]
        confidence = "high" if len(trusted_items) >= 2 else ("medium" if (len(trusted_items) == 1 or used_trusted) else "low")

        # 4. Углубление: читаем верхний доверенный источник
        deep_text = ""
        top_items = trusted_items[:2] if trusted_items else items[:2]
        for it in top_items[:1]:  # читаем 1 источник (экономим лимит)
            url = it.get("url", "")
            if url.startswith("http"):
                try:
                    extracted = await asyncio.to_thread(tavily_extract, url, 3500)
                    if extracted:
                        deep_text = extracted
                        break
                except Exception:
                    pass

        # 5. Собираем sources
        sources = []
        for it in (top_items or items)[:3]:
            url = it.get("url", "")
            sources.append({"domain": _dom(url), "url": url, "title": (it.get("title") or "")[:120]})

        # 6. Формируем сырое summary (потом эксперт переработает в ответ)
        summary_parts = []
        if answer:
            summary_parts.append(f"Краткая сводка: {answer}")
        if deep_text:
            summary_parts.append(f"\nИз источника {sources[0]['domain'] if sources else 'неизвестно'}:\n{deep_text[:3000]}")
        # Дополнительные сниппеты
        snippets = []
        for it in (top_items or items)[:4]:
            if it.get("content"):
                snippets.append(f"• {_dom(it.get('url', ''))}: {it['content'][:300]}")
        if snippets:
            summary_parts.append("\nДругие источники:\n" + "\n".join(snippets))

        result = {
            "ok": True,
            "summary": "\n".join(summary_parts),
            "sources": sources,
            "confidence": confidence,
            "trusted_count": len(trusted_items),
        }
        _research_cache[cache_key] = (now, result)
        print(f"[expert_research] confidence={confidence}, trusted={len(trusted_items)}/{len(items)}")
        return result
    except Exception as e:
        print(f"[expert_research] error: {e}")
        return {"ok": False}


async def health_data_context(user_id, message: str) -> str:
    """Собирает показатели здоровья для агента-аналитика. Тип берётся из сообщения."""
    if not db_pool or not user_id:
        return ""
    ql = (message or "").lower()
    type_map = {"pressure": ["давлен"], "pulse": ["пульс"],
                "sugar": ["сахар", "глюкоз"], "weight": ["вес", "похуд", "килогр"]}
    chosen = None
    for t, kws in type_map.items():
        if any(k in ql for k in kws):
            chosen = t
            break
    types = [chosen] if chosen else ["pressure", "pulse", "sugar", "weight"]
    type_ru = {"pressure": "Давление", "pulse": "Пульс", "sugar": "Сахар", "weight": "Вес"}
    try:
        async with db_pool.acquire() as conn:
            lines = []
            for t in types:
                rows = await conn.fetch(
                    "SELECT value_1, value_2, recorded_at FROM health_records "
                    "WHERE user_id=$1 AND type=$2 ORDER BY recorded_at DESC LIMIT 10", user_id, t)
                if not rows:
                    continue
                vals = []
                for r in reversed(rows):
                    v = f"{int(r['value_1'])}/{int(r['value_2'])}" if r["value_2"] else str(r["value_1"])
                    vals.append(v)
                lines.append(f"{type_ru[t]} (от старых к новым): {', '.join(vals)}")
            if not lines:
                return "NO_DATA"
            return "\n".join(lines)
    except Exception as e:
        print(f"[analyze health] error: {e}")
        return ""


async def search_links(query: str) -> list:
    """Tavily-поиск → список результатов со ссылками для главного чата."""
    if not TAVILY_API_KEY:
        return []
    try:
        data = await asyncio.to_thread(tavily_search_adv, query, "general", None, 6, False)
        items = data.get("results", []) or []
        out = []
        for it in items:
            url = it.get("url", "")
            if not url.startswith("http"):
                continue
            try:
                dom = url.split("/")[2].replace("www.", "")
            except Exception:
                dom = "источник"
            out.append({
                "title": (it.get("title") or dom)[:120],
                "url": url,
                "domain": dom,
                "snippet": (it.get("content") or "")[:160],
            })
            if len(out) >= 5:
                break
        print(f"[search_links] '{query[:40]}' → {len(out)} ссылок")
        return out
    except Exception as e:
        print(f"[search_links] error: {e}")
        return []


@app.post("/api/v1/chat")
async def chat_v1(req: ChatRequest, request: Request):
    try:
        # Получаем user_id из токена если он есть
        user_id = None
        try:
            token = request.headers.get("Authorization", "").replace("Bearer ", "").strip()
            if token and token not in ("demo", "local_demo", ""):
                if db_pool:
                    async with db_pool.acquire() as conn:
                        row = await conn.fetchrow(
                            "SELECT user_id FROM sessions WHERE token=$1 AND expires_at > NOW()", token
                        )
                        if row:
                            user_id = row["user_id"]
        except Exception:
            pass

        # Проверка лимита токенов
        if user_id and not await check_token_limit(user_id):
            raise HTTPException(
                status_code=429,
                detail="Достигнут лимит на вашем тарифе. Обновите подписку для продолжения."
            )

        # Оркестратор: если раздел не задан явно фронтом — определяем сами
        domain = req.mode
        subcategory = None
        if domain and ":" in domain:
            domain, subcategory = domain.split(":", 1)
        if not domain:
            domain = await classify_domain(req.message)

        # Этап 2: запись данных через чат (с подтверждением на фронте)
        ql = (req.message or "").lower()
        import re as _re
        is_explicit_record = domain in ("health", "finance") and any(kw in ql for kw in _RECORD_KW)
        is_implicit_spend = domain == "finance" and any(kw in ql for kw in _SPEND_KW) and bool(_re.search(r"\d", ql))
        if is_explicit_record or is_implicit_spend:
            ext = await extract_record(req.message, domain)
            if ext.get("ok"):
                await add_tokens(user_id, req.message, ext["label"])
                confirm_text = f"Записать в расходы: {ext['label']}?" if domain == "finance" else f"Записать: {ext['label']}?"
                await save_chat_message(user_id, domain, "user", req.message)
                await save_chat_message(user_id, domain, "ai", confirm_text, action="confirm_record")
                return {
                    "reply": confirm_text,
                    "action": "confirm_record",
                    "domain": domain,
                    "records": ext["records"],
                }
            # не распозналось как запись — продолжаем обычным ответом

        # Этап 4: анализ данных пользователя (траты/показатели) агентом-аналитиком
        if domain in ("health", "finance") and any(k in ql for k in _ANALYSIS_KW):
            prof = await build_profile_context(user_id) if user_id else ""
            if domain == "finance":
                ctx = await finance_data_context(user_id)
                if ctx == "NO_DATA":
                    return {"reply": "Пока не вижу записанных трат. Добавьте расходы — и я разберу, на что уходят деньги. Можно прямо здесь: «запиши такси 500».", "domain": domain}
                if ctx:
                    msgs = [{"role": "system", "content": FINANCE_ANALYST_PROMPT + prof},
                            {"role": "user", "content": f"Данные трат:\n{ctx}\n\nЗапрос: {req.message}"}]
                    reply = await asyncio.to_thread(call_ai, msgs)
                    await add_tokens(user_id, ctx, reply)
                    await save_chat_message(user_id, domain, "user", req.message)
                    await save_chat_message(user_id, domain, "ai", reply)
                    return {"reply": reply, "domain": domain}
            else:  # health
                ctx = await health_data_context(user_id, req.message)
                if ctx == "NO_DATA":
                    return {"reply": "Пока нет записанных показателей. Внесите измерения — и я оценю динамику. Можно здесь: «запиши давление 120 на 80».", "domain": domain}
                if ctx:
                    msgs = [{"role": "system", "content": HEALTH_ANALYST_PROMPT + prof},
                            {"role": "user", "content": f"Показатели:\n{ctx}\n\nЗапрос: {req.message}"}]
                    reply = await asyncio.to_thread(call_ai, msgs)
                    await add_tokens(user_id, ctx, reply)
                    await save_chat_message(user_id, domain, "user", req.message)
                    await save_chat_message(user_id, domain, "ai", reply)
                    return {"reply": reply, "domain": domain}

        # info-вопрос про законы/льготы/новости/как оформить → Tavily research
        is_research = _is_info_query(req.message) or any(k in ql for k in _RESEARCH_KW)
        if domain not in ("health", "finance", "scam") and is_research:
            ans = await research_answer(req.message)
            if ans:
                await add_tokens(user_id, req.message, ans)
                await save_chat_message(user_id, domain, "user", req.message)
                await save_chat_message(user_id, domain, "ai", ans)
                return {"reply": ans, "domain": domain}
            # Tavily не дал ответа — продолжаем обычным ответом

        # Поиск товара/услуги/события → тот же движок раздела Поиск, результаты в чат
        if domain == "search":
            await add_tokens(user_id, req.message, "поиск")
            await save_chat_message(user_id, domain, "user", req.message)
            await save_chat_message(user_id, domain, "ai", "Ищу по вашему запросу…")
            return {"reply": "Ищу по вашему запросу…", "action": "search_results",
                    "query": req.message, "domain": domain}

        # Строим системный промпт с профилем
        base_prompt = req.system or SYSTEM_PROMPTS.get(domain, SYSTEM_PROMPT)
        # Специализированный эксперт по подкатегории финансов (mode='finance:utility' и т.п.)
        if domain == "finance" and subcategory and subcategory in FINANCE_CATEGORY_PROMPTS:
            base_prompt = FINANCE_CATEGORY_PROMPTS[subcategory]
            # Подгружаем данные по этой категории прямо в системный контекст
            cat_data = await finance_category_data(user_id, subcategory)
            if cat_data:
                base_prompt = base_prompt + "\n\nДАННЫЕ ПОЛЬЗОВАТЕЛЯ ПО ЭТОЙ КАТЕГОРИИ:\n" + cat_data
            # Если запрос требует свежих внешних данных — подключаем expert_research
            if _needs_external_data(req.message, subcategory):
                research = await expert_research(req.message, subcategory)
                if research.get("ok"):
                    conf_ru = {"high": "высокая (≥2 доверенных источника)",
                               "medium": "средняя (1 доверенный источник)",
                               "low": "низкая (источники не из проверенных)"}
                    src_list = ", ".join(s["domain"] for s in research["sources"][:3])
                    base_prompt = base_prompt + (
                        f"\n\nАКТУАЛЬНЫЕ ДАННЫЕ ИЗ ИНТЕРНЕТА (достоверность: {conf_ru[research['confidence']]}):\n"
                        f"{research['summary']}\n\n"
                        f"Используй эти данные в ответе. Если достоверность 'low' — добавь оговорку "
                        f"«проверьте у официального источника». В конце ответа укажи источники: {src_list}."
                    )
        profile_ctx = await build_profile_context(user_id) if user_id else ""
        sys_prompt = base_prompt + profile_ctx

        messages = [{"role": "system", "content": sys_prompt}]
        for h in req.history[-10:]:
            messages.append(h)
        messages.append({"role": "user", "content": req.message})
        reply = call_ai(messages)
        # Учёт израсходованных токенов
        await add_tokens(user_id, sys_prompt + req.message, reply)
        # Этап 3: единая история диалога — дублируется в раздел
        await save_chat_message(user_id, domain, "user", req.message)
        await save_chat_message(user_id, domain, "ai", reply)
        return {"reply": reply, "domain": domain}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI недоступен: {str(e)}")


@app.get("/api/v1/chat/history")
async def chat_history(domain: str = None, limit: int = 50, user=Depends(get_current_user)):
    """Единая история диалогов. domain=health → срез по разделу (дублирование)."""
    if not db_pool:
        return {"messages": []}
    try:
        async with db_pool.acquire() as conn:
            if domain:
                rows = await conn.fetch(
                    "SELECT domain, role, text, action, created_at FROM chat_messages "
                    "WHERE user_id=$1 AND domain=$2 ORDER BY created_at DESC LIMIT $3",
                    user["id"], domain, min(limit, 100))
            else:
                rows = await conn.fetch(
                    "SELECT domain, role, text, action, created_at FROM chat_messages "
                    "WHERE user_id=$1 ORDER BY created_at DESC LIMIT $2",
                    user["id"], min(limit, 100))
        msgs = []
        for r in reversed(rows):
            msgs.append({"domain": r["domain"], "role": r["role"], "text": r["text"],
                         "action": r["action"], "created_at": str(r["created_at"])})
        return {"messages": msgs}
    except Exception as e:
        print(f"[chat_history] read error: {e}")
        return {"messages": []}


# ── Админ-панель владельца ───────────────────────────────────────
# Внутренний счётчик учитывает только финальный обмен; реальные API-вызовы
# (оркестрация classify/extract/research/vision) ~ в 3.5 раза больше.
# Цена ~$0.20/млн реальных токенов, курс ~92 ₽/$.
_USD_RUB = 92
_TOKEN_COST_RUB_PER_INTERNAL = (3.5 * 0.20 * _USD_RUB) / 1_000_000


async def log_app_error(source: str, message: str):
    """Логирует ошибку в БД для агента-мониторинга."""
    if not db_pool:
        return
    try:
        async with db_pool.acquire() as conn:
            await conn.execute(
                "INSERT INTO app_errors (source, message) VALUES ($1,$2)",
                source[:80], str(message)[:2000])
    except Exception:
        pass


def call_reasoner(prompt: str, system: str = None, max_tokens: int = 1500) -> str:
    """Глубокий анализ через DeepSeek Reasoner (для умных агентов)."""
    ds_key = DEEPSEEK_API_KEY or AITUNNEL_API_KEY
    ds_base = DEEPSEEK_BASE_URL if DEEPSEEK_API_KEY else AITUNNEL_BASE_URL
    ds_model = "deepseek-reasoner" if DEEPSEEK_API_KEY else "deepseek-r1-0528"
    msgs = []
    if system:
        msgs.append({"role": "system", "content": system})
    msgs.append({"role": "user", "content": prompt})
    headers = {"Authorization": f"Bearer {ds_key}", "Content-Type": "application/json"}
    data = {"model": ds_model, "messages": msgs, "max_tokens": max_tokens}
    r = requests.post(f"{ds_base}chat/completions", headers=headers, json=data, timeout=90)
    if r.status_code == 200:
        return r.json()["choices"][0]["message"]["content"].strip()
    # fallback на обычный chat
    return call_ai(msgs)


async def collect_finance_metrics() -> dict:
    """Собирает реальные финансовые данные для агента-финансиста."""
    PRICES = {"standard": 390, "pro": 590}
    async with db_pool.acquire() as conn:
        tariffs = await conn.fetch("SELECT tariff, COUNT(*) c FROM users GROUP BY tariff")
        tmap = {r["tariff"]: r["c"] for r in tariffs}
        total_users = sum(tmap.values())
        tokens_total = int(await conn.fetchval("SELECT COALESCE(SUM(tokens_used),0) FROM users") or 0)
        new_30d = await conn.fetchval("SELECT COUNT(*) FROM users WHERE created_at >= NOW() - interval '30 days'") or 0
        visits_total = await conn.fetchval("SELECT COUNT(*) FROM site_visits") or 0
    # выручка-прокси по тарифам
    revenue = sum(PRICES.get(t, 0) * c for t, c in tmap.items())
    paid = sum(c for t, c in tmap.items() if t in PRICES)
    # расходы
    infra = 740  # Amvera + Tavily + TTS
    tokens_cost = round(tokens_total * _TOKEN_COST_RUB_PER_INTERNAL, 2)
    nds = round(revenue * 22 / 122, 2)          # НДС 22% (внутри цены)
    acquiring = round(revenue * 0.039, 2)        # эквайринг 3.9%
    ad = round(revenue * 0.15, 2)                # реклама 15%
    expenses_total = round(infra + tokens_cost + nds + acquiring + ad, 2)
    profit = round(revenue - expenses_total, 2)
    margin = round(profit / revenue * 100, 1) if revenue else None
    breakeven_std = round(740 / 226, 1)
    return {
        "currency_rate": _USD_RUB,
        "users": {"total": total_users, "paid": paid, "new_30d": new_30d, "by_tariff": tmap},
        "visits_total": visits_total,
        "revenue_proxy": revenue,
        "tokens": {"internal": tokens_total, "cost_rub": tokens_cost},
        "expenses": {"infra": infra, "tokens": tokens_cost, "nds_22": nds,
                     "acquiring_3.9": acquiring, "advertising_15": ad, "total": expenses_total},
        "profit": profit, "margin_pct": margin,
        "prices": PRICES,
    }


async def collect_health_metrics() -> dict:
    """Собирает реальные технические метрики для агента-мониторинга."""
    m = {"db": "unknown", "tables": {}, "scheduler": {}, "push_subs": 0,
         "errors_24h": 0, "recent_errors": [], "external_apis": {}}
    # БД
    try:
        async with db_pool.acquire() as conn:
            await conn.fetchval("SELECT 1")
            m["db"] = "ok"
            for t in ("users", "finance_records", "health_records", "chat_messages",
                      "notifications", "push_subscriptions", "site_visits"):
                try:
                    m["tables"][t] = await conn.fetchval(f"SELECT COUNT(*) FROM {t}") or 0
                except Exception:
                    m["tables"][t] = "n/a"
            # scheduler heartbeat
            try:
                sch = await conn.fetchrow("SELECT last_check, leader_pid, pushes_sent FROM scheduler_state LIMIT 1")
                if sch and sch["last_check"]:
                    delta = (datetime.now() - sch["last_check"]).total_seconds()
                    m["scheduler"] = {"last_check_sec_ago": int(delta),
                                      "alive": delta < 180,
                                      "pushes_sent": sch["pushes_sent"]}
                else:
                    m["scheduler"] = {"alive": False, "note": "нет heartbeat"}
            except Exception as e:
                m["scheduler"] = {"alive": False, "error": str(e)[:100]}
            m["push_subs"] = await conn.fetchval("SELECT COUNT(*) FROM push_subscriptions") or 0
            # ошибки за 24ч
            try:
                m["errors_24h"] = await conn.fetchval(
                    "SELECT COUNT(*) FROM app_errors WHERE created_at >= NOW() - interval '24 hours'") or 0
                errs = await conn.fetch(
                    "SELECT source, message, created_at FROM app_errors ORDER BY created_at DESC LIMIT 8")
                m["recent_errors"] = [{"source": e["source"], "message": (e["message"] or "")[:200],
                                       "at": e["created_at"].isoformat() if e["created_at"] else None} for e in errs]
            except Exception:
                pass
    except Exception as e:
        m["db"] = f"error: {str(e)[:120]}"
    # ping внешних API (быстро, с коротким timeout)
    def _ping(name, url, headers=None):
        try:
            t0 = __import__("time").time()
            r = requests.get(url, headers=headers or {}, timeout=8)
            return {"status": r.status_code, "ms": int((__import__("time").time() - t0) * 1000)}
        except Exception as e:
            return {"status": "error", "error": str(e)[:80]}
    m["external_apis"]["deepseek"] = _ping("deepseek", "https://api.deepseek.com/v1/models",
                                           {"Authorization": f"Bearer {DEEPSEEK_API_KEY}"} if DEEPSEEK_API_KEY else None)
    m["external_apis"]["aitunnel"] = _ping("aitunnel", f"{AITUNNEL_BASE_URL}models",
                                           {"Authorization": f"Bearer {AITUNNEL_API_KEY}"} if AITUNNEL_API_KEY else None)
    return m


FINANCE_AGENT_PROMPT = """Ты — финансовый директор (CFO) SaaS-стартапа OkTolk с 15-летним опытом в unit-экономике подписочных продуктов. Перед тобой РЕАЛЬНЫЕ данные из базы. Анализируй только их, не выдумывай.

Контекст бизнеса: OkTolk — PWA AI-помощник, ИП на ОСНО (НДС 22%), подписки Стандарт 390₽ и Про 590₽/мес. Эквайринг Робокасса 3.9%, реклама ~15%, инфраструктура ~740₽/мес фикс. Себестоимость токенов считается по реальному потреблению.

Дай краткий, конкретный анализ как живому фаундеру:
1. **Здоровье экономики** — рентабельность, маржа, где сейчас бизнес (прибыль/убыток/около нуля).
2. **Что в цифрах беспокоит** — самые тяжёлые статьи расходов, точка безубыточности, насколько близко.
3. **2-3 конкретных действия** — что сделать чтобы улучшить unit-экономику (не общие слова, а под эти цифры).

Тон — деловой партнёр, без воды. Цифры округляй. Если данных мало (мало юзеров) — честно скажи что выводы предварительные. Структура: короткие абзацы, **жирным** ключевые числа, списки через «— »."""

HEALTH_AGENT_PROMPT = """Ты — Senior SRE/DevOps инженер с опытом эксплуатации production-систем. Перед тобой РЕАЛЬНЫЕ технические метрики приложения OkTolk (FastAPI + PostgreSQL на Amvera, Москва). Анализируй только фактические данные.

Стек: FastAPI, PostgreSQL, фоновый scheduler push-уведомлений (должен слать heartbeat каждую минуту, alive если последний <180 сек назад), Web Push (VAPID), внешние API (DeepSeek, AItunnel).

Дай чёткий отчёт о состоянии системы:
1. **Статус** — система здорова / есть проблемы / критично. Одной фразой.
2. **Что проверено и что нашёл** — БД, scheduler, push, внешние API, ошибки за 24ч. Отметь конкретные проблемы (scheduler не шлёт heartbeat, API недоступен, всплеск ошибок).
3. **Что делать** — если есть проблемы, конкретные шаги. Если всё ок — так и скажи, без выдумывания проблем.

Тон — технический, точный, без паники и без воды. Если метрика в норме — не раздувай. Структура: короткие абзацы, **жирным** статусы, списки через «— »."""


@app.get("/api/v1/admin/agent/finance")
async def admin_agent_finance(admin=Depends(get_admin_user)):
    """Агент-финансист: реальные данные → анализ рентабельности через Reasoner."""
    metrics = await collect_finance_metrics()
    prompt = ("РЕАЛЬНЫЕ ФИНАНСОВЫЕ ДАННЫЕ (JSON):\n" + json.dumps(metrics, ensure_ascii=False, indent=2) +
              "\n\nПроанализируй unit-экономику и дай рекомендации.")
    try:
        report = await asyncio.to_thread(call_reasoner, prompt, FINANCE_AGENT_PROMPT, 1800)
    except Exception as e:
        await log_app_error("agent_finance", str(e))
        raise HTTPException(status_code=500, detail="Агент недоступен")
    return {"metrics": metrics, "report": report}


@app.get("/api/v1/admin/agent/health")
async def admin_agent_health(admin=Depends(get_admin_user)):
    """Агент-мониторинг: реальные тех-метрики → анализ через Reasoner."""
    metrics = await collect_health_metrics()
    prompt = ("РЕАЛЬНЫЕ ТЕХНИЧЕСКИЕ МЕТРИКИ (JSON):\n" + json.dumps(metrics, ensure_ascii=False, indent=2) +
              "\n\nПроанализируй состояние системы и дай отчёт.")
    try:
        report = await asyncio.to_thread(call_reasoner, prompt, HEALTH_AGENT_PROMPT, 1500)
    except Exception as e:
        await log_app_error("agent_health", str(e))
        raise HTTPException(status_code=500, detail="Агент недоступен")
    return {"metrics": metrics, "report": report}


@app.post("/api/v1/track/visit")
async def track_visit(request: Request):
    """Публичный счётчик переходов. Фронт вызывает при загрузке."""
    if not db_pool:
        return {"ok": False}
    try:
        body = {}
        try:
            body = await request.json()
        except Exception:
            pass
        path = (body.get("path") or "/")[:120]
        async with db_pool.acquire() as conn:
            await conn.execute("INSERT INTO site_visits (path) VALUES ($1)", path)
        return {"ok": True}
    except Exception as e:
        print(f"[track_visit] error: {e}")
        return {"ok": False}


@app.get("/api/v1/admin/check")
async def admin_check(user=Depends(get_current_user)):
    """Проверка прав админа — для показа кнопки в профиле."""
    return {"is_admin": bool(user.get("is_admin"))}


@app.get("/api/v1/admin/overview")
async def admin_overview(admin=Depends(get_admin_user)):
    """Сводка: переходы, регистрации, подписки, токены, себестоимость."""
    async with db_pool.acquire() as conn:
        total_users = await conn.fetchval("SELECT COUNT(*) FROM users") or 0
        new_7d = await conn.fetchval("SELECT COUNT(*) FROM users WHERE created_at >= NOW() - interval '7 days'") or 0
        new_30d = await conn.fetchval("SELECT COUNT(*) FROM users WHERE created_at >= NOW() - interval '30 days'") or 0
        tariffs = await conn.fetch("SELECT tariff, COUNT(*) c FROM users GROUP BY tariff")
        tariff_map = {r["tariff"]: r["c"] for r in tariffs}
        paid = sum(c for t, c in tariff_map.items() if t and t != "demo")
        visits_total = await conn.fetchval("SELECT COUNT(*) FROM site_visits") or 0
        visits_7d = await conn.fetchval("SELECT COUNT(*) FROM site_visits WHERE visited_at >= NOW() - interval '7 days'") or 0
        tokens_total = await conn.fetchval("SELECT COALESCE(SUM(tokens_used),0) FROM users") or 0
        cost_rub = round(float(tokens_total) * _TOKEN_COST_RUB_PER_INTERNAL, 2)
        conv_reg = round(total_users / visits_total * 100, 1) if visits_total else None
        conv_paid = round(paid / total_users * 100, 1) if total_users else None
    return {
        "users": {"total": total_users, "new_7d": new_7d, "new_30d": new_30d},
        "visits": {"total": visits_total, "last_7d": visits_7d},
        "subscriptions": {"paid": paid, "by_tariff": tariff_map},
        "tokens": {"total_internal": int(tokens_total), "est_cost_rub": cost_rub},
        "conversion": {"visit_to_reg_pct": conv_reg, "reg_to_paid_pct": conv_paid},
    }


@app.get("/api/v1/admin/demographics")
async def admin_demographics(admin=Depends(get_admin_user)):
    """Распределение по полу, возрасту, городу."""
    async with db_pool.acquire() as conn:
        gender = await conn.fetch(
            "SELECT COALESCE(gender,'не указан') g, COUNT(*) c FROM user_profile GROUP BY gender ORDER BY c DESC")
        cities = await conn.fetch(
            "SELECT COALESCE(city,'не указан') city, COUNT(*) c FROM users GROUP BY city ORDER BY c DESC LIMIT 15")
        ages = await conn.fetch("""
            SELECT CASE
              WHEN age IS NULL THEN 'не указан'
              WHEN age < 25 THEN 'до 25'
              WHEN age < 35 THEN '25–34'
              WHEN age < 45 THEN '35–44'
              WHEN age < 55 THEN '45–54'
              WHEN age < 65 THEN '55–64'
              ELSE '65+' END grp, COUNT(*) c
            FROM user_profile GROUP BY grp ORDER BY c DESC""")
    return {
        "gender": [{"label": r["g"], "count": r["c"]} for r in gender],
        "age": [{"label": r["grp"], "count": r["c"]} for r in ages],
        "city": [{"label": r["city"], "count": r["c"]} for r in cities],
    }


@app.get("/api/v1/admin/users")
async def admin_users_list(admin=Depends(get_admin_user), limit: int = 100):
    """Список пользователей с токенами и себестоимостью."""
    async with db_pool.acquire() as conn:
        rows = await conn.fetch("""
            SELECT u.id, u.name, u.phone, u.tariff, u.city, u.created_at,
                   COALESCE(u.tokens_used,0) tokens_used, u.tokens_limit,
                   p.gender, p.age
            FROM users u LEFT JOIN user_profile p ON p.user_id = u.id
            ORDER BY u.created_at DESC LIMIT $1""", limit)
    users = []
    for r in rows:
        tu = int(r["tokens_used"] or 0)
        ph = r["phone"] or ""
        ph_mask = (ph[:2] + "***" + ph[-2:]) if len(ph) >= 4 else "***"
        users.append({
            "id": r["id"], "name": r["name"] or "—", "phone": ph_mask,
            "tariff": r["tariff"], "city": r["city"], "gender": r["gender"], "age": r["age"],
            "created_at": r["created_at"].isoformat() if r["created_at"] else None,
            "tokens_used": tu, "tokens_limit": r["tokens_limit"],
            "cost_rub": round(tu * _TOKEN_COST_RUB_PER_INTERNAL, 2),
        })
    return {"users": users}


class BroadcastRequest(BaseModel):
    title: str
    body: str
    push: bool = True


@app.post("/api/v1/admin/broadcast")
async def admin_broadcast(req: BroadcastRequest, admin=Depends(get_admin_user)):
    """Рассылка уведомления всем пользователям (колокольчик + push)."""
    if not req.title.strip() or not req.body.strip():
        raise HTTPException(status_code=400, detail="Заголовок и текст обязательны")
    notif_count = 0
    push_count = 0
    async with db_pool.acquire() as conn:
        user_ids = [r["id"] for r in await conn.fetch("SELECT id FROM users")]
        for uid in user_ids:
            try:
                await conn.execute(
                    "INSERT INTO notifications (user_id, type, title, body, is_read) VALUES ($1,'system',$2,$3,FALSE)",
                    uid, req.title.strip(), req.body.strip())
                notif_count += 1
            except Exception as e:
                print(f"[broadcast] notif {uid}: {e}")
        if req.push:
            subs = await conn.fetch("SELECT user_id, endpoint, p256dh, auth FROM push_subscriptions")
            for s in subs:
                try:
                    ok = send_web_push(
                        {"endpoint": s["endpoint"], "keys": {"p256dh": s["p256dh"], "auth": s["auth"]}},
                        req.title.strip(), req.body.strip())
                    if ok:
                        push_count += 1
                except Exception as e:
                    print(f"[broadcast] push {s['user_id']}: {e}")
    print(f"[broadcast] notif={notif_count}, push={push_count}")
    return {"ok": True, "notifications_sent": notif_count, "push_sent": push_count}


@app.get("/api/v1/research/stats")
async def research_stats(user=Depends(get_current_user)):
    """Диагностика расхода expert_research (лимит Tavily 1000/мес)."""
    return {
        "tavily_calls": _research_stats["calls"],
        "cache_hits": _research_stats["cache_hits"],
        "cache_size": len(_research_cache),
        "cache_ttl_sec": _RESEARCH_CACHE_TTL,
    }


class SavedSearchRequest(BaseModel):
    title: str = ""
    url: str = ""
    description: str = ""
    source: str = ""
    image: str = ""
    price: str = ""
    cat: str = ""


@app.get("/api/v1/search/saved")
async def get_saved_searches(user=Depends(get_current_user)):
    if not db_pool:
        return {"items": []}
    try:
        async with db_pool.acquire() as conn:
            rows = await conn.fetch(
                "SELECT title, url, description, source, image, price, cat, saved_at "
                "FROM saved_searches WHERE user_id=$1 ORDER BY saved_at DESC", user["id"])
        items = []
        for r in rows:
            items.append({"title": r["title"], "url": r["url"], "description": r["description"],
                          "source": r["source"], "image": r["image"], "price": r["price"], "cat": r["cat"]})
        return {"items": items}
    except Exception as e:
        print(f"[saved] read error: {e}")
        return {"items": []}


@app.post("/api/v1/search/saved")
async def add_saved_search(req: SavedSearchRequest, user=Depends(get_current_user)):
    if not db_pool or not req.url:
        return {"status": "error"}
    try:
        async with db_pool.acquire() as conn:
            await conn.execute(
                "INSERT INTO saved_searches (user_id, title, url, description, source, image, price, cat) "
                "VALUES ($1,$2,$3,$4,$5,$6,$7,$8) ON CONFLICT (user_id, url) DO NOTHING",
                user["id"], req.title, req.url, req.description, req.source, req.image, req.price, req.cat)
        return {"status": "saved"}
    except Exception as e:
        print(f"[saved] add error: {e}")
        return {"status": "error"}


@app.delete("/api/v1/search/saved")
async def delete_saved_search(url: str, user=Depends(get_current_user)):
    if not db_pool:
        return {"status": "error"}
    try:
        async with db_pool.acquire() as conn:
            await conn.execute("DELETE FROM saved_searches WHERE user_id=$1 AND url=$2", user["id"], url)
        return {"status": "deleted"}
    except Exception as e:
        print(f"[saved] delete error: {e}")
        return {"status": "error"}

# Antiscam
@app.post("/api/v1/analyze/text")
async def analyze_v1(req: AnalyzeRequest):
    if not AITUNNEL_API_KEY:
        raise HTTPException(status_code=503, detail="AI недоступен")
    try:
        prompt = f"""Ты антискам система OkTolk. Анализируй внутри себя, показывай только краткий результат.

ПРАВИЛА:
- Не описывай что видишь на скриншоте
- Не пересказывай переписку
- Только вердикт и рекомендации
- Простой язык без терминов

АНАЛИЗИРУЙ ВНУТРИ: тип схемы, психологические триггеры, технические признаки, логические противоречия

ПРЕДЫДУЩИЙ КОНТЕКСТ (учти при анализе):
{req.context if req.context else "нет"}

ТЕКСТ ДЛЯ АНАЛИЗА:
{req.text}

Ответь ТОЛЬКО JSON:
{{"risk_level": 0, "summary": "1 предложение - суть угрозы или её отсутствия", "how_they_manipulate": "как пытаются обмануть или null", "consequences": "к чему может привести или null", "recommendation": "одно конкретное действие"}}

ШКАЛА: 0=безопасно 1=минимальный 2=подозрения 3=высокий 4=очень высокий 5=явное мошенничество
Только JSON, никакого текста до или после.
"""

        headers = {"Authorization": f"Bearer {AITUNNEL_API_KEY}", "Content-Type": "application/json"}
        data = {
            "model": "gemini-2.5-flash-lite",
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": 500
        }
        response = requests.post(f"{AITUNNEL_BASE_URL}chat/completions", headers=headers, json=data, timeout=30)
        result = response.json()
        text = result["choices"][0]["message"]["content"].strip()

        # Чистим JSON
        json_match = re.search(r'\{.*\}', text, re.DOTALL)
        if json_match:
            parsed = json.loads(json_match.group())
        else:
            parsed = json.loads(text)

        return parsed
    except Exception as e:
        # Fallback на ключевые слова
        scam_keywords = ["код из смс", "переведи деньги", "данные карты", "номер карты", "cvv", "пин-код", "срочно переведи", "выиграли", "одобрен кредит", "заблокирован счет"]
        text_lower = req.text.lower()
        risk_level = 0
        signs = []
        for kw in scam_keywords:
            if kw in text_lower:
                risk_level = 3
                signs.append(f"Обнаружено: '{kw}'")
        labels = ["Безопасно", "Подозрительно", "Опасно", "Мошенники!"]
        return {"risk_level": risk_level, "risk_label": labels[min(risk_level, 3)], "signs": signs,
                "summary": "СТОП! Мошенники!" if risk_level == 3 else "Проверьте ещё раз",
                "action": "Заблокируйте отправителя" if risk_level == 3 else "Будьте осторожны"}

# ── User Profile ─────────────────────────────────────────────────

class ProfileRequest(BaseModel):
    user_id: Optional[int] = None
    gender: Optional[str] = None
    age: Optional[int] = None
    height: Optional[int] = None
    weight: Optional[float] = None
    profession: Optional[str] = None
    hobbies: Optional[str] = None
    work_pressure_1: Optional[int] = None
    work_pressure_2: Optional[int] = None
    work_pulse: Optional[int] = None
    base_sugar: Optional[float] = None
    habits: Optional[str] = None
    activity: Optional[str] = None
    alcohol: Optional[int] = 0
    smoking: Optional[int] = 0
    city: Optional[str] = None
    timezone: Optional[str] = None
    smoking_years: Optional[int] = None
    heredity: Optional[str] = None
    chronic: Optional[str] = None
    income: Optional[float] = None
    marital_status: Optional[str] = None
    children: Optional[list] = None
    family_notes: Optional[str] = None

@app.get("/api/v1/profile")
async def get_profile(user=Depends(get_current_user)):
    """Получить профиль текущего пользователя"""
    if not db_pool:
        return {}
    async with db_pool.acquire() as conn:
        row = await conn.fetchrow("SELECT * FROM user_profile WHERE user_id=$1", user["id"])
        # Берём city/timezone из таблицы users
        u = await conn.fetchrow("SELECT city, timezone, name FROM users WHERE id=$1", user["id"])
        result = dict(row) if row else {}
        # children хранится как JSON-строка → отдаём массивом
        if result.get("children"):
            try:
                result["children"] = json.loads(result["children"]) if isinstance(result["children"], str) else result["children"]
            except Exception:
                result["children"] = []
        if u:
            result["city"] = u["city"] or "Москва"
            result["timezone"] = u["timezone"] or "Europe/Moscow"
            result["name"] = u["name"]
        return result

@app.post("/api/v1/profile")
async def save_profile(req: ProfileRequest, user=Depends(get_current_user)):
    """Сохранить/обновить профиль (upsert) — user_id из токена"""
    if not db_pool:
        raise HTTPException(status_code=503, detail="БД недоступна")
    uid = user["id"]
    async with db_pool.acquire() as conn:
        await conn.execute("""
            INSERT INTO user_profile
                (user_id, gender, age, height, weight, profession, hobbies,
                 work_pressure_1, work_pressure_2, work_pulse, base_sugar,
                 habits, activity, alcohol, smoking, smoking_years, heredity, chronic, income,
                 marital_status, children, family_notes, updated_at)
            VALUES ($1,$2,$3,$4,$5,$6,$7,$8,$9,$10,$11,$12,$13,$14,$15,$16,$17,$18,$19,$20,$21,$22,NOW())
            ON CONFLICT (user_id) DO UPDATE SET
                gender=COALESCE($2, user_profile.gender),
                age=COALESCE($3, user_profile.age),
                height=COALESCE($4, user_profile.height),
                weight=COALESCE($5, user_profile.weight),
                profession=COALESCE($6, user_profile.profession),
                hobbies=COALESCE($7, user_profile.hobbies),
                work_pressure_1=COALESCE($8, user_profile.work_pressure_1),
                work_pressure_2=COALESCE($9, user_profile.work_pressure_2),
                work_pulse=COALESCE($10, user_profile.work_pulse),
                base_sugar=COALESCE($11, user_profile.base_sugar),
                habits=COALESCE($12, user_profile.habits),
                activity=COALESCE($13, user_profile.activity),
                alcohol=COALESCE($14, user_profile.alcohol),
                smoking=COALESCE($15, user_profile.smoking),
                smoking_years=COALESCE($16, user_profile.smoking_years),
                heredity=COALESCE($17, user_profile.heredity),
                chronic=COALESCE($18, user_profile.chronic),
                income=COALESCE($19, user_profile.income),
                marital_status=COALESCE($20, user_profile.marital_status),
                children=COALESCE($21, user_profile.children),
                family_notes=COALESCE($22, user_profile.family_notes),
                updated_at=NOW()
        """, uid, req.gender, req.age, req.height, req.weight, req.profession,
            req.hobbies, req.work_pressure_1, req.work_pressure_2, req.work_pulse,
            req.base_sugar, req.habits, req.activity, req.alcohol, req.smoking,
            req.smoking_years, req.heredity, req.chronic, req.income,
            req.marital_status,
            json.dumps(req.children, ensure_ascii=False) if req.children is not None else None,
            req.family_notes)
        # Сохраняем city/timezone в таблицу users
        if req.city is not None or req.timezone is not None:
            await conn.execute("""
                UPDATE users SET
                    city = COALESCE($2, city),
                    timezone = COALESCE($3, timezone)
                WHERE id=$1
            """, uid, req.city, req.timezone)
        return {"status": "ok"}

# ── Credits & Loans ──────────────────────────────────────────────

class CreditRequest(BaseModel):
    id: Optional[int] = None
    user_id: Optional[int] = None
    name: Optional[str] = None
    amount: Optional[float] = None
    rate: Optional[float] = None
    term_months: Optional[int] = None
    monthly_payment: Optional[float] = None
    notes: Optional[str] = None

class LoanRequest(BaseModel):
    id: Optional[int] = None
    user_id: Optional[int] = None
    name: Optional[str] = None
    amount: Optional[float] = None
    due_date: Optional[str] = None
    monthly_payment: Optional[float] = None
    notes: Optional[str] = None

@app.get("/api/v1/credits")
async def get_credits(user=Depends(get_current_user)):
    if not db_pool:
        return []
    async with db_pool.acquire() as conn:
        rows = await conn.fetch("SELECT * FROM credits WHERE user_id=$1 ORDER BY created_at DESC", user["id"])
        return [dict(r) for r in rows]

@app.post("/api/v1/credits")
async def add_credit(req: CreditRequest, user=Depends(get_current_user)):
    if not db_pool:
        raise HTTPException(status_code=503, detail="БД недоступна")
    uid = user["id"]
    async with db_pool.acquire() as conn:
        if req.id:
            await conn.execute("""
                UPDATE credits SET name=$2, amount=$3, rate=$4, term_months=$5, monthly_payment=$6, notes=$7
                WHERE id=$1 AND user_id=$8
            """, req.id, req.name, req.amount, req.rate, req.term_months, req.monthly_payment, req.notes, uid)
            return {"id": req.id, "status": "ok"}
        row = await conn.fetchrow("""
            INSERT INTO credits (user_id, name, amount, rate, term_months, monthly_payment, notes)
            VALUES ($1,$2,$3,$4,$5,$6,$7) RETURNING id
        """, uid, req.name, req.amount, req.rate, req.term_months, req.monthly_payment, req.notes)
        return {"id": row["id"], "status": "ok"}

@app.delete("/api/v1/credits/{credit_id}")
async def delete_credit(credit_id: int, user=Depends(get_current_user)):
    if not db_pool:
        raise HTTPException(status_code=503, detail="БД недоступна")
    async with db_pool.acquire() as conn:
        await conn.execute("DELETE FROM credits WHERE id=$1 AND user_id=$2", credit_id, user["id"])
        return {"status": "ok"}

@app.get("/api/v1/loans")
async def get_loans(user=Depends(get_current_user)):
    if not db_pool:
        return []
    async with db_pool.acquire() as conn:
        rows = await conn.fetch("SELECT * FROM loans WHERE user_id=$1 ORDER BY created_at DESC", user["id"])
        return [dict(r) for r in rows]

@app.post("/api/v1/loans")
async def add_loan(req: LoanRequest, user=Depends(get_current_user)):
    if not db_pool:
        raise HTTPException(status_code=503, detail="БД недоступна")
    uid = user["id"]
    async with db_pool.acquire() as conn:
        if req.id:
            await conn.execute("""
                UPDATE loans SET name=$2, amount=$3, due_date=$4, monthly_payment=$5, notes=$6
                WHERE id=$1 AND user_id=$7
            """, req.id, req.name, req.amount, req.due_date, req.monthly_payment, req.notes, uid)
            return {"id": req.id, "status": "ok"}
        row = await conn.fetchrow("""
            INSERT INTO loans (user_id, name, amount, due_date, monthly_payment, notes)
            VALUES ($1,$2,$3,$4,$5,$6) RETURNING id
        """, uid, req.name, req.amount, req.due_date, req.monthly_payment, req.notes)
        return {"id": row["id"], "status": "ok"}

@app.delete("/api/v1/loans/{loan_id}")
async def delete_loan(loan_id: int, user=Depends(get_current_user)):
    if not db_pool:
        raise HTTPException(status_code=503, detail="БД недоступна")
    async with db_pool.acquire() as conn:
        await conn.execute("DELETE FROM loans WHERE id=$1 AND user_id=$2", loan_id, user["id"])
        return {"status": "ok"}

# ── Health ───────────────────────────────────────────────────────

class HealthParseRequest(BaseModel):
    user_id: Optional[int] = None  # берётся из токена, не нужен в теле
    type: str
    image_base64: Optional[str] = None
    mime_type: str = "image/jpeg"
    text: Optional[str] = None

class HealthRecordCreateRequest(BaseModel):
    user_id: Optional[int] = None  # берётся из токена
    type: str
    value_1: float
    value_2: Optional[float] = None
    comment: Optional[str] = None

@app.get("/api/v1/health/records")
async def get_health_records(
    type: Optional[str] = None,
    days: int = 30,
    limit: int = 100,
    user=Depends(get_current_user)
):
    """Записи здоровья текущего пользователя (user_id из токена)"""
    if not db_pool:
        return []
    uid = user["id"]
    async with db_pool.acquire() as conn:
        if type:
            rows = await conn.fetch("""
                SELECT id, user_id, type, value_1, value_2, comment, recorded_at
                FROM health_records
                WHERE user_id=$1
                  AND type=$2
                  AND recorded_at >= NOW() - make_interval(days => $3)
                ORDER BY recorded_at DESC
                LIMIT $4
            """, uid, type, days, limit)
        else:
            rows = await conn.fetch("""
                SELECT id, user_id, type, value_1, value_2, comment, recorded_at
                FROM health_records
                WHERE user_id=$1
                  AND recorded_at >= NOW() - make_interval(days => $2)
                ORDER BY recorded_at DESC
                LIMIT $3
            """, uid, days, limit)
        return [dict(r) for r in rows]

@app.post("/api/v1/health/records")
async def add_health_record_public(req: HealthRecordCreateRequest, user=Depends(get_current_user)):
    """Добавить показатель здоровья (user_id из токена)"""
    if not db_pool:
        raise HTTPException(status_code=503, detail="БД недоступна")
    async with db_pool.acquire() as conn:
        row = await conn.fetchrow(
            """INSERT INTO health_records (user_id, type, value_1, value_2, comment)
               VALUES ($1,$2,$3,$4,$5) RETURNING id, recorded_at""",
            user["id"], req.type, req.value_1, req.value_2, req.comment
        )
        return {"id": row["id"], "recorded_at": str(row["recorded_at"]), "status": "ok"}

@app.post("/api/v1/health/parse")
async def parse_health_photo(req: HealthParseRequest, user=Depends(get_current_user)):
    """AI распознавание показателя из фото (тонометр, весы, глюкометр) или текста"""
    if not AITUNNEL_API_KEY:
        raise HTTPException(status_code=503, detail="AI недоступен")

    # Описание что извлекаем — работает и для фото, и для текста
    type_desc = {
        "pressure": 'показатели артериального давления. value_1 = верхнее (систолическое), value_2 = нижнее (диастолическое)',
        "pulse":    'частоту пульса. value_1 = пульс в уд/мин, value_2 = null',
        "sugar":    'уровень сахара в крови. value_1 = ммоль/л, value_2 = null',
        "weight":   'вес тела. value_1 = вес в кг, value_2 = null',
    }
    desc = type_desc.get(req.type, type_desc["pressure"])

    source = "на фото (тонометр/весы/глюкометр/прибор)" if req.image_base64 else "из текста пользователя"
    prompt = (
        f'Извлеки {source} {desc}. '
        f'Если пользователь описал самочувствие или причину (стресс, нагрузка, после еды и т.п.) — добавь это в comment, иначе comment = краткое описание. '
        f'Верни ТОЛЬКО валидный JSON без пояснений: '
        f'{{"value_1": число, "value_2": число или null, "comment": "текст"}}'
    )

    try:
        headers = {"Authorization": f"Bearer {AITUNNEL_API_KEY}", "Content-Type": "application/json"}

        if req.image_base64:
            messages = [{"role": "user", "content": [
                {"type": "image_url", "image_url": {"url": f"data:{req.mime_type};base64,{req.image_base64}"}},
                {"type": "text", "text": prompt}
            ]}]
        else:
            messages = [{"role": "user", "content": f"{prompt}\n\nТекст пользователя: {req.text}"}]

        data = {"model": "gemini-2.5-flash-lite", "messages": messages, "max_tokens": 200}
        response = requests.post(f"{AITUNNEL_BASE_URL}chat/completions", headers=headers, json=data, timeout=30)
        result_text = response.json()["choices"][0]["message"]["content"].strip()

        json_match = re.search(r'\{.*\}', result_text, re.DOTALL)
        if not json_match:
            raise HTTPException(status_code=422, detail="Не удалось распознать показатель")

        parsed = json.loads(json_match.group())
        if parsed.get("value_1") is None:
            raise HTTPException(status_code=422, detail="Значение не найдено")

        # Дефолтный комментарий — исходный текст пользователя (если был)
        default_comment = req.text if req.text else "с фото"

        # Сохраняем в БД
        if db_pool:
            async with db_pool.acquire() as conn:
                row = await conn.fetchrow(
                    """INSERT INTO health_records (user_id, type, value_1, value_2, comment)
                       VALUES ($1,$2,$3,$4,$5) RETURNING id, recorded_at""",
                    user["id"], req.type,
                    float(parsed["value_1"]),
                    float(parsed["value_2"]) if parsed.get("value_2") is not None else None,
                    parsed.get("comment") or default_comment
                )
                parsed["id"] = row["id"]
                parsed["recorded_at"] = str(row["recorded_at"])
                parsed["status"] = "ok"

        return parsed

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка распознавания: {str(e)}")


@app.post("/api/v1/health/chat")
async def health_chat(req: dict = Body(...), user=Depends(get_current_user)):
    """AI-чат по здоровью с контекстом данных пользователя"""
    text = req.get("text", "")
    metric_type = req.get("type", "")
    if not text:
        raise HTTPException(status_code=400, detail="Нет текста")
    try:
        context_parts = []
        if db_pool:
            async with db_pool.acquire() as conn:
                # Профиль пользователя
                profile = await conn.fetchrow("SELECT * FROM user_profile WHERE user_id=$1", user["id"])
                if profile:
                    p = dict(profile)
                    parts = []
                    if p.get("age"): parts.append(f"возраст {p['age']} лет")
                    if p.get("gender"): parts.append(f"пол: {p['gender']}")
                    if p.get("height"): parts.append(f"рост {p['height']} см")
                    if p.get("weight"): parts.append(f"вес {p['weight']} кг")
                    if p.get("work_pressure_1") and p.get("work_pressure_2"):
                        parts.append(f"рабочее давление {p['work_pressure_1']}/{p['work_pressure_2']}")
                    if parts:
                        context_parts.append("Данные пользователя: " + ", ".join(parts))
                # Последние записи по типу
                if metric_type:
                    records = await conn.fetch(
                        """SELECT type, value_1, value_2, comment, recorded_at 
                           FROM health_records WHERE user_id=$1 AND type=$2 
                           ORDER BY recorded_at DESC LIMIT 10""",
                        user["id"], metric_type
                    )
                else:
                    records = await conn.fetch(
                        """SELECT type, value_1, value_2, comment, recorded_at 
                           FROM health_records WHERE user_id=$1 
                           ORDER BY recorded_at DESC LIMIT 20""",
                        user["id"]
                    )
                if records:
                    rec_parts = []
                    for r in records:
                        type_names = {"pressure": "Давление", "pulse": "Пульс", "sugar": "Сахар", "weight": "Вес"}
                        tname = type_names.get(r["type"], r["type"])
                        if r["value_2"]:
                            rec_parts.append(f"{tname}: {r['value_1']}/{r['value_2']}")
                        else:
                            rec_parts.append(f"{tname}: {r['value_1']}")
                    context_parts.append("Последние показатели: " + "; ".join(rec_parts))

        context = "\n".join(context_parts)
        system = HEALTH_PROMPT + ("\n\nКонтекст пользователя:\n" + context if context else "")
        messages = [
            {"role": "system", "content": system},
            {"role": "user", "content": text}
        ]
        reply = call_ai(messages)
        return {"reply": reply}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/v1/health/meds")
async def get_medications(user=Depends(get_current_user)):
    """Список активных лекарств пользователя"""
    async with db_pool.acquire() as conn:
        # Миграция: добавляем колонки если их нет
        for col, definition in [
            ("times_per_day", "INTEGER DEFAULT 1"),
            ("times_json",    "TEXT DEFAULT '[]'"),
            ("food_rule",     "VARCHAR(50) DEFAULT 'any'"),
            ("notify",        "BOOLEAN DEFAULT FALSE"),
            ("taken_today",   "TEXT DEFAULT '[]'"),
            ("taken_date",    "DATE"),
        ]:
            try:
                await conn.execute(f"ALTER TABLE medications ADD COLUMN IF NOT EXISTS {col} {definition}")
            except Exception:
                pass
        rows = await conn.fetch(
            "SELECT * FROM medications WHERE user_id=$1 AND is_active=TRUE ORDER BY id",
            user["id"]
        )
        result = []
        for r in rows:
            import json as _json
            d = dict(r)
            # Парсим JSON поля
            try: d["times"] = _json.loads(d.get("times_json") or "[]")
            except: d["times"] = []
            # Проверяем taken_today — сбрасываем если дата изменилась
            from datetime import date
            today = date.today()
            if d.get("taken_date") != today:
                d["taken_today_list"] = []
            else:
                try: d["taken_today_list"] = _json.loads(d.get("taken_today") or "[]")
                except: d["taken_today_list"] = []
            result.append(d)
        return result

class MedRequest(BaseModel):
    name: str
    dose: Optional[str] = None
    dose_unit: Optional[str] = "мг"
    times_per_day: Optional[int] = 1
    times: Optional[list] = []       # ["09:00", "21:00"]
    food_rule: Optional[str] = "any" # before/after/during/any
    food_minutes: Optional[int] = None
    notify: Optional[bool] = False
    start_date: Optional[str] = None
    end_date: Optional[str] = None

@app.post("/api/v1/health/meds")
async def add_medication(req: MedRequest, user=Depends(get_current_user)):
    """Добавить лекарство"""
    import json as _json
    from datetime import date
    async with db_pool.acquire() as conn:
        dose_str = f"{req.dose} {req.dose_unit}".strip() if req.dose else req.dose_unit or ""
        row_id = await conn.fetchval(
            """INSERT INTO medications
               (user_id, name, dose, schedule, times_per_day, times_json, food_rule, notify, start_date, is_active)
               VALUES ($1,$2,$3,$4,$5,$6,$7,$8,$9,TRUE) RETURNING id""",
            user["id"], req.name, dose_str,
            req.food_rule,
            req.times_per_day,
            _json.dumps(req.times or []),
            req.food_rule,
            req.notify,
            date.today()
        )
        return {"id": row_id, "status": "ok"}

@app.put("/api/v1/health/meds/{med_id}")
async def update_medication(med_id: int, req: MedRequest, user=Depends(get_current_user)):
    """Обновить лекарство"""
    import json as _json
    dose_str = f"{req.dose} {req.dose_unit}".strip() if req.dose else req.dose_unit or ""
    async with db_pool.acquire() as conn:
        await conn.execute("""
            UPDATE medications SET
                name=$3, dose=$4, schedule=$5, times_per_day=$6,
                times_json=$7, food_rule=$8, notify=$9
            WHERE id=$1 AND user_id=$2
        """, med_id, user["id"],
            req.name, dose_str, req.food_rule,
            req.times_per_day,
            _json.dumps(req.times or []),
            req.food_rule, req.notify
        )
    return {"status": "ok", "id": med_id}

@app.delete("/api/v1/health/meds/{med_id}")
async def delete_medication(med_id: int, user=Depends(get_current_user)):
    """Удалить лекарство (soft delete)"""
    async with db_pool.acquire() as conn:
        await conn.execute(
            "UPDATE medications SET is_active=FALSE WHERE id=$1 AND user_id=$2",
            med_id, user["id"]
        )
        return {"status": "ok"}

@app.post("/api/v1/health/meds/{med_id}/taken")
async def mark_taken(med_id: int, request: Request, user=Depends(get_current_user)):
    """Отметить приём лекарства"""
    import json as _json
    from datetime import date
    data = await request.json()
    time_slot = data.get("time", "")
    today = date.today()
    async with db_pool.acquire() as conn:
        row = await conn.fetchrow(
            "SELECT taken_today, taken_date FROM medications WHERE id=$1 AND user_id=$2",
            med_id, user["id"]
        )
        if not row:
            raise HTTPException(status_code=404, detail="Не найдено")
        if row["taken_date"] == today:
            try: taken = _json.loads(row["taken_today"] or "[]")
            except: taken = []
        else:
            taken = []
        if time_slot and time_slot not in taken:
            taken.append(time_slot)
        await conn.execute(
            "UPDATE medications SET taken_today=$1, taken_date=$2 WHERE id=$3",
            _json.dumps(taken), today, med_id
        )
        return {"status": "ok", "taken": taken}

# Finance
@app.get("/api/v1/finance/records")
async def get_finance_records(category: Optional[str] = None, user=Depends(get_current_user)):
    async with db_pool.acquire() as conn:
        if category:
            rows = await conn.fetch(
                "SELECT * FROM finance_records WHERE user_id=$1 AND category=$2 ORDER BY recorded_at DESC LIMIT 100",
                user["id"], category
            )
        else:
            rows = await conn.fetch(
                "SELECT * FROM finance_records WHERE user_id=$1 ORDER BY recorded_at DESC LIMIT 100",
                user["id"]
            )
        return [dict(r) for r in rows]

@app.post("/api/v1/finance/records")
async def add_finance_record(req: FinanceRecordRequest, user=Depends(get_current_user)):
    async with db_pool.acquire() as conn:
        row_id = await conn.fetchval(
            "INSERT INTO finance_records (user_id, category, amount, comment) VALUES ($1,$2,$3,$4) RETURNING id",
            user["id"], req.category, req.amount, req.comment
        )
        return {"id": row_id, "status": "ok"}

@app.get("/api/v1/finance/summary")
async def get_finance_summary(user=Depends(get_current_user)):
    async with db_pool.acquire() as conn:
        rows = await conn.fetch("""
            SELECT category, SUM(amount) as total, COUNT(*) as count
            FROM finance_records
            WHERE user_id=$1
            AND recorded_at >= date_trunc('month', NOW())
            GROUP BY category
        """, user["id"])
        return {"month": datetime.now().strftime("%Y-%m"), "categories": [dict(r) for r in rows]}

# ── Landing & App routes ─────────────────────────────────────────
@app.get("/")
async def root():
    """Лендинг — главная страница для поисковиков и новых посетителей"""
    return FileResponse("landing.html")

@app.get("/app")
async def app_route():
    """PWA приложение OkTolk"""
    return FileResponse("index.html")

@app.post("/chat")
async def chat(req: ChatRequest):
    sys_prompt = req.system or SYSTEM_PROMPTS.get(req.mode or "", SYSTEM_PROMPT)
    messages = [{"role": "system", "content": sys_prompt}]
    for h in req.history[-10:]:
        messages.append(h)
    messages.append({"role": "user", "content": req.message})
    reply = call_ai(messages)
    return {"reply": reply, "risk_level": 0}

@app.post("/analyze/text")
async def analyze_text(req: AnalyzeRequest):
    return await analyze_v1(req)

TAVILY_API_KEY = os.getenv("TAVILY_API_KEY", "")
_news_cache = {"data": [], "at": 0}

# ═══════════════════════════════════════════════════════════════
# Tavily helper-функции: Extract (чтение страниц) + улучшенный Search
# ═══════════════════════════════════════════════════════════════

def tavily_extract(urls, max_chars: int = 6000) -> str:
    """
    Читает реальное содержимое страниц через Tavily Extract API.
    Возвращает чистый текст (без HTML-мусора). urls — строка или список.
    """
    if not TAVILY_API_KEY:
        return ""
    if isinstance(urls, str):
        urls = [urls]
    try:
        r = requests.post(
            "https://api.tavily.com/extract",
            json={"api_key": TAVILY_API_KEY, "urls": urls[:5],
                  "include_raw_content": False},
            timeout=20
        )
        if r.status_code != 200:
            print(f"[tavily_extract] status={r.status_code}")
            return ""
        data = r.json()
        chunks = []
        for res in data.get("results", []):
            content = (res.get("raw_content") or res.get("content") or "").strip()
            if content:
                chunks.append(content)
        return "\n\n".join(chunks)[:max_chars]
    except Exception as e:
        print(f"[tavily_extract] error: {e}")
        return ""

def tavily_search_adv(query: str, topic: str = "general", time_range: str = None,
                      max_results: int = 5, include_answer: bool = False,
                      include_domains: list = None) -> dict:
    """
    Улучшенный поиск: topic (general/news), time_range (day/week/month),
    фильтр доменов. Возвращает полный JSON-ответ Tavily.
    """
    if not TAVILY_API_KEY:
        return {}
    payload = {
        "api_key": TAVILY_API_KEY,
        "query": query,
        "search_depth": "basic",
        "max_results": max_results,
        "include_answer": include_answer,
        "country": "ru",
        "topic": topic,
    }
    if time_range:
        payload["time_range"] = time_range
    if include_domains:
        payload["include_domains"] = include_domains
    try:
        r = requests.post("https://api.tavily.com/search", json=payload, timeout=15)
        if r.status_code == 200:
            return r.json()
        print(f"[tavily_search_adv] status={r.status_code}")
        return {}
    except Exception as e:
        print(f"[tavily_search_adv] error: {e}")
        return {}

async def simplify_news_ai(title: str, body: str) -> dict:
    """AI переводит на русский и упрощает новость для людей 40+"""
    prompt = (
        "Ты редактор новостей для обычных людей. Перепиши новость простым русским языком, как будто рассказываешь соседу.\n"
        "ПРАВИЛА:\n"
        "- ВСЕГДА на русском языке (если исходник на английском — переведи на русский)\n"
        "- Заголовок: коротко и понятно, максимум 8 слов\n"
        "- Текст: 2-3 простых предложения, без сложных слов и терминов\n"
        "- Только суть: что произошло и что это значит для обычного человека\n\n"
        f"Исходный заголовок: {title}\n"
        f"Исходный текст: {body}\n\n"
        "Ответь ТОЛЬКО валидным JSON без markdown:\n"
        '{"title": "простой заголовок", "body": "простой текст 2-3 предложения"}'
    )
    try:
        import asyncio, json as _json, re as _re
        loop = asyncio.get_event_loop()
        result = await loop.run_in_executor(None, lambda: call_ai([{"role": "user", "content": prompt}]))
        # Вырезаем JSON из ответа
        cleaned = result.strip()
        m = _re.search(r'\{.*\}', cleaned, _re.DOTALL)
        if m:
            cleaned = m.group(0)
        parsed = _json.loads(cleaned)
        return {"title": str(parsed.get("title", title))[:120], "body": str(parsed.get("body", body))[:300]}
    except Exception:
        return {"title": title[:120], "body": body[:300]}

async def topics_to_queries(topics_text: str) -> list:
    """AI превращает свободный текст тем пользователя в поисковые запросы.
    Возвращает список (query, cat, tag)."""
    prompt = (
        "Пользователь описал какие новости хочет получать. Преврати это в 3-5 поисковых запросов "
        "для новостного поиска на русском языке.\n"
        f"ЗАПРОС ПОЛЬЗОВАТЕЛЯ: {topics_text}\n\n"
        "Для каждого запроса определи категорию:\n"
        "- danger (опасности, мошенники, угрозы)\n"
        "- success (выплаты, льготы, хорошие новости)\n"
        "- warn (изменения, важные правила)\n"
        "- info (всё остальное)\n\n"
        "Ответь ТОЛЬКО валидным JSON-массивом без markdown:\n"
        '[{"query":"поисковый запрос", "cat":"info", "tag":"Тема"}]'
    )
    try:
        import asyncio, json as _json, re as _re
        loop = asyncio.get_event_loop()
        result = await loop.run_in_executor(None, lambda: call_ai([{"role": "user", "content": prompt}]))
        m = _re.search(r'\[.*\]', result.strip(), _re.DOTALL)
        if not m:
            return []
        arr = _json.loads(m.group(0))
        out = []
        for it in arr[:5]:
            q = str(it.get("query", "")).strip()
            if q:
                out.append((q, it.get("cat", "info"), str(it.get("tag", "Новость"))[:20]))
        return out
    except Exception as e:
        print(f"[topics_to_queries] {e}")
        return []

async def fetch_tavily_news(topics_text: str = None) -> list:
    """Получить новости через Tavily + AI упрощение.
    topics_text — персональные темы пользователя (если заданы)."""
    import asyncio, time
    now = time.time()
    personalized = bool(topics_text and topics_text.strip())

    # Кеш только для дефолтных новостей (персональные всегда свежие)
    if not personalized and _news_cache["data"] and now - _news_cache["at"] < 3600:
        return _news_cache["data"]

    if personalized:
        queries = await topics_to_queries(topics_text)
        if not queries:
            queries = [(topics_text.strip(), "info", "Новость")]
    else:
        # Дефолтные запросы с упором на официальные русские источники
        queries = [
            ("телефонные мошенники банки предупреждение МВД", "danger", "Опасно"),
            ("пенсии индексация выплаты пенсионерам СФР", "success", "Льготы"),
            ("новые правила ЖКХ коммунальные услуги", "warn", "Важно"),
            ("льготные лекарства бесплатные препараты Минздрав", "success", "Льготы"),
            ("мошенничество госуслуги фишинг безопасность", "danger", "Опасно"),
        ]

    def _fetch_tavily_sync(q, ru_only: bool = True):
        """Синхронный запрос к Tavily. ru_only=True фильтрует по стране."""
        payload = {
            "api_key": TAVILY_API_KEY,
            "query": q,
            "search_depth": "basic",
            "max_results": 2,
            "include_answer": True,
            "topic": "news",
            "time_range": "month",
            "country": "ru" if ru_only else None,
        }
        if not ru_only:
            payload.pop("country", None)
        try:
            resp = requests.post("https://api.tavily.com/search", json=payload, timeout=10)
            if resp.status_code != 200:
                print(f"[tavily_news] HTTP {resp.status_code} для запроса: {q}")
                return None
            data = resp.json()
            results = data.get("results", [])
            print(f"[tavily_news] '{q[:40]}' → {len(results)} результатов")
            return data
        except Exception as e:
            print(f"[tavily_news] error: {e}")
            return None

    raw = []
    try:
        for i, (q, cat, tag) in enumerate(queries):
            data = await asyncio.to_thread(_fetch_tavily_sync, q, True)
            if not data:
                continue
            items = data.get("results", [])
            if not items:
                continue
            item = items[0]
            title = item.get("title", "").strip()
            body = (data.get("answer") or item.get("content", "")).strip()
            url = item.get("url", "")
            source = url.split("/")[2].replace("www.", "").split(".")[0].upper() if url and "/" in url else "Источник"
            if title and body:
                raw.append({"id": i+1, "cat": cat, "tag": tag, "raw_title": title, "raw_body": body, "source": source, "url": url, "age": "сегодня"})
    except Exception:
        pass

    # AI упрощает и переводит каждую новость
    results = []
    for r in raw:
        simple = await simplify_news_ai(r["raw_title"], r["raw_body"])
        results.append({
            "id": r["id"], "cat": r["cat"], "tag": r["tag"],
            "title": simple["title"], "body": simple["body"],
            "source": r["source"], "url": r["url"], "age": r["age"]
        })

    # Fallback (только для дефолтных — у персональных свой пустой результат)
    if not results and not personalized:
        results = [
            {"id": 1, "cat": "danger", "tag": "Опасно", "title": "Мошенники звонят от имени банков", "body": "Участились случаи звонков мошенников, представляющихся сотрудниками банков. Никогда не сообщайте код из СМС.", "source": "МВД", "url": "https://mvd.ru", "age": "сегодня"},
            {"id": 2, "cat": "success", "tag": "Льготы", "title": "Перерасчёт пенсий с 1 июня 2026", "body": "Пенсионный фонд проведёт автоматический перерасчёт пенсий неработающим пенсионерам. Заявление не требуется.", "source": "СФР", "url": "https://sfr.gov.ru", "age": "сегодня"},
            {"id": 3, "cat": "warn", "tag": "Важно", "title": "Изменения в правилах ЖКХ", "body": "С 1 июля меняется порядок передачи показаний счётчиков. Срок — до 25 числа каждого месяца.", "source": "ГЖИ", "url": "", "age": "вчера"},
            {"id": 4, "cat": "success", "tag": "Льготы", "title": "Бесплатные лекарства расширили список", "body": "В перечень бесплатных препаратов для льготников включены 24 новых лекарства от давления и диабета.", "source": "Минздрав", "url": "https://minzdrav.gov.ru", "age": "вчера"},
            {"id": 5, "cat": "danger", "tag": "Опасно", "title": "Фальшивые госуслуги в Telegram", "body": "Появились боты, имитирующие портал Госуслуг. Запрашивают паспортные данные и СНИЛС.", "source": "Госуслуги", "url": "https://gosuslugi.ru", "age": "2 дн назад"},
        ]

    # Кеш и рассылка уведомлений — только для дефолтной ленты
    if personalized:
        return results

    _news_cache["data"] = results
    _news_cache["at"] = now

    # Создаём уведомления для опасных новостей всем пользователям
    danger_items = [r for r in results if r["cat"] == "danger"]
    if danger_items and db_pool:
        try:
            async with db_pool.acquire() as conn:
                user_ids = await conn.fetch("SELECT id FROM users WHERE is_demo=FALSE OR is_demo IS NULL")
                for uid in user_ids:
                    for item in danger_items[:2]:  # максимум 2 уведомления за обновление
                        await create_notification(
                            user_id=uid["id"],
                            type="scam",
                            title=item["title"],
                            body=item.get("body", "")[:200]
                        )
        except Exception as e:
            print(f"[news notif] {e}")

    return results

@app.get("/api/v1/news")
async def get_news_v1(authorization: Optional[str] = Header(None)):
    """Новости через Tavily API с AI-упрощением.
    Если у пользователя заданы персональные темы — лента под него, иначе общая (кеш 1 час)."""
    import traceback as _tb
    topics = None
    if authorization and authorization.startswith("Bearer ") and db_pool:
        try:
            token = authorization.split(" ", 1)[1]
            async with db_pool.acquire() as conn:
                session = await conn.fetchrow(
                    "SELECT user_id FROM sessions WHERE token=$1 AND expires_at > NOW()", token
                )
                if session:
                    row = await conn.fetchrow("SELECT news_topics FROM users WHERE id=$1", session["user_id"])
                    if row and row["news_topics"]:
                        topics = row["news_topics"]
        except Exception:
            pass
    try:
        return await fetch_tavily_news(topics)
    except Exception as e:
        print(f"[news] CRITICAL ERROR: {e}")
        _tb.print_exc()
        # Возвращаем хардкод-fallback вместо 500
        return [
            {"id": 1, "cat": "danger", "tag": "Опасно", "title": "Мошенники звонят от имени банков", "body": "Участились случаи звонков мошенников, представляющихся сотрудниками банков. Никогда не сообщайте код из СМС.", "source": "МВД", "url": "https://mvd.ru", "age": "сегодня"},
            {"id": 2, "cat": "success", "tag": "Льготы", "title": "Перерасчёт пенсий", "body": "Пенсионный фонд проведёт автоматический перерасчёт пенсий неработающим пенсионерам. Заявление не требуется.", "source": "СФР", "url": "https://sfr.gov.ru", "age": "сегодня"},
            {"id": 3, "cat": "warn", "tag": "Важно", "title": "Изменения в правилах ЖКХ", "body": "С 1 июля меняется порядок передачи показаний счётчиков. Срок — до 25 числа каждого месяца.", "source": "ГЖИ", "url": "", "age": "вчера"},
            {"id": 4, "cat": "success", "tag": "Льготы", "title": "Бесплатные лекарства расширили список", "body": "В перечень бесплатных препаратов для льготников включены новые лекарства от давления и диабета.", "source": "Минздрав", "url": "https://minzdrav.gov.ru", "age": "вчера"},
            {"id": 5, "cat": "danger", "tag": "Опасно", "title": "Фальшивые госуслуги в Telegram", "body": "Появились боты, имитирующие портал Госуслуг. Запрашивают паспортные данные и СНИЛС.", "source": "Госуслуги", "url": "https://gosuslugi.ru", "age": "2 дн назад"},
        ]

class NewsTopicsRequest(BaseModel):
    topics: str

@app.get("/api/v1/news/topics")
async def get_news_topics(user=Depends(get_current_user)):
    """Получить сохранённые темы новостей пользователя"""
    if not db_pool:
        return {"topics": ""}
    async with db_pool.acquire() as conn:
        row = await conn.fetchrow("SELECT news_topics FROM users WHERE id=$1", user["id"])
    return {"topics": (row["news_topics"] if row and row["news_topics"] else "")}

@app.post("/api/v1/news/topics")
async def set_news_topics(req: NewsTopicsRequest, user=Depends(get_current_user)):
    """Сохранить персональные темы новостей. Пустая строка = вернуться к общей ленте."""
    if not db_pool:
        raise HTTPException(status_code=503, detail="БД недоступна")
    topics = (req.topics or "").strip()[:500]
    async with db_pool.acquire() as conn:
        await conn.execute("UPDATE users SET news_topics=$1 WHERE id=$2",
                           topics if topics else None, user["id"])
    return {"status": "ok", "topics": topics}

# ═══════════════════════════════════════════════════════════════
# Ticketland: реферальные ссылки через Advcake
# ═══════════════════════════════════════════════════════════════

def build_ticketland_url(title: str, event_url: str = "", sub1: str = "") -> str:
    """
    Строит реферальную ссылку на Ticketland.
    - Если есть event_url от KudaGo и он ведёт на ticketland.ru — используем его напрямую
    - Иначе — поиск по названию: ticketland.ru/search/?q=TITLE
    - Если задан TICKETLAND_AFFILIATE_TEMPLATE — оборачиваем в Advcake deeplink
    """
    from urllib.parse import quote
    # Определяем целевой URL
    if event_url and "ticketland.ru" in event_url:
        target_url = event_url
    else:
        clean_title = title.replace("...", "").strip()
        target_url = f"https://www.ticketland.ru/search/?q={quote(clean_title)}"
    # Если есть шаблон от Advcake — оборачиваем
    if TICKETLAND_AFFILIATE_TEMPLATE:
        encoded_target = quote(target_url, safe="")
        link = TICKETLAND_AFFILIATE_TEMPLATE.replace("{url}", encoded_target)
        if sub1:
            link = link.replace("{sub1}", quote(sub1, safe=""))
        else:
            link = link.replace("&sub1={sub1}", "").replace("{sub1}", "")
        return link
    # Без шаблона — просто ссылка на Ticketland (без комиссии, но рабочая)
    return target_url

@app.get("/api/v1/events/ticket-link")
async def ticket_link_redirect(
    title: str = "",
    url: str = "",
    request: Request = None
):
    """
    Генерирует реф-ссылку на Ticketland + логирует клик.
    Frontend вызывает этот эндпоинт, получает URL и открывает его.
    """
    # Логируем клик в БД для аналитики
    try:
        user_id = None
        if request:
            auth = request.headers.get("Authorization", "")
            if auth.startswith("Bearer "):
                import jwt as _jwt
                try:
                    payload = _jwt.decode(auth[7:], SECRET_KEY, algorithms=["HS256"])
                    user_id = payload.get("user_id")
                except:
                    pass
        async with db_pool.acquire() as conn:
            await conn.execute("""
                INSERT INTO notifications (user_id, type, title, body, is_read)
                VALUES ($1, 'ticket_click', $2, $3, true)
            """, user_id, f"Переход: {title[:80]}", url[:200] or title[:200])
    except Exception as e:
        print(f"[ticket_link] log error (not critical): {e}")

    affiliate_active = bool(TICKETLAND_AFFILIATE_TEMPLATE)
    final_url = build_ticketland_url(title, url)
    print(f"[ticket_link] title='{title[:50]}' affiliate={affiliate_active} → {final_url[:80]}")
    return {"url": final_url, "affiliate_active": affiliate_active}

# ═══════════════════════════════════════════════════════════════
# KudaGo: умный поиск мероприятий (бесплатный API, без ключа)
# ═══════════════════════════════════════════════════════════════

_KUDAGO_CITIES = {
    "москва": "msk", "москве": "msk", "москву": "msk", "москвы": "msk",
    "питер": "spb", "санкт-петербург": "spb", "петербург": "spb", "спб": "spb",
    "казань": "kzn", "казани": "kzn",
    "новосибирск": "nsk", "новосибирске": "nsk",
    "екатеринбург": "ekb", "екате": "ekb", "екатеринбурге": "ekb",
    "нижний новгород": "nny", "нижнем": "nny", "нижний": "nny",
    "краснодар": "krd", "краснодаре": "krd",
    "сочи": "sochi", "уфа": "ufa", "уфе": "ufa",
    "ростов": "ras", "самара": "sam", "пермь": "per",
    "омск": "omsk", "воронеж": "vrn", "красноярск": "krasnoyarsk",
}

_KUDAGO_CATS = {
    "концерт": "concert", "концерты": "concert", "концерте": "concert",
    "театр": "theater", "театре": "theater",
    "спектакль": "theater", "спектакли": "theater",
    "выставк": "exhibition",
    "фестиваль": "festival", "фестивал": "festival",
    "квест": "quest", "квесты": "quest",
    "вечеринк": "party", "клуб": "party",
    "шоу": "show",
    "дети": "kids", "детей": "kids", "детское": "kids", "ребенок": "kids",
    "спорт": "sport",
    "мастер-класс": "master-class", "мастер класс": "master-class",
    "экскурс": "tour",
    "стендап": "stand_up", "stand-up": "stand_up",
    "ярмарк": "yarmarki",
    "перформанс": "show",
}

_EVENTS_KW = [
    "куда сходить", "куда пойти", "мероприятия", "мероприятие", "мероприятий",
    "афиша", "афише", "концерт", "выставк", "театр", "спектакль", "фестиваль",
    "квест", "вечеринк", "шоу", "экскурс", "мастер-класс", "мастер класс",
    "стендап", "ярмарк", "перформанс", "куда пойти", "что посетить",
    "что посмотреть", "событий", "событие", "развлечен",
]

_MONTHS_RU = {
    "January": "января", "February": "февраля", "March": "марта",
    "April": "апреля", "May": "мая", "June": "июня",
    "July": "июля", "August": "августа", "September": "сентября",
    "October": "октября", "November": "ноября", "December": "декабря",
}

def _is_events_query(q: str) -> bool:
    ql = q.lower()
    return any(kw in ql for kw in _EVENTS_KW)

def _is_info_query(q: str) -> bool:
    """Информационный вопрос (как/что/где/сколько/документы/правила),
    а не поиск товара. Такие запросы получают прямой AI-ответ + источники."""
    ql = q.lower().strip()
    if ql.endswith("?"):
        return True
    info_starts = (
        "как ", "что ", "где ", "когда ", "почему ", "зачем ", "сколько ",
        "какой ", "какая ", "какие ", "какое ", "можно ли", "нужно ли",
        "куда ", "кто ", "чем ", "за что", "из-за чего", "стоит ли",
        "положен", "положена", "положено", "имею ли", "обязан",
    )
    if ql.startswith(info_starts):
        return True
    info_words = (
        "как оформить", "как получить", "как сделать", "как написать",
        "как подать", "как вернуть", "что делать", "какие документы",
        "документы для", "документы на", "правила", "инструкция",
        "пошагово", "объясни", "расскажи про", "что такое",
    )
    return any(w in ql for w in info_words)

def _parse_kudago_params(query: str) -> dict:
    q = query.lower()
    # Город
    location = "msk"
    for word, code in _KUDAGO_CITIES.items():
        if word in q:
            location = code
            break
    # Категории
    cats = []
    for kw, cat in _KUDAGO_CATS.items():
        if kw in q and cat not in cats:
            cats.append(cat)
    # Бесплатно
    is_free = "бесплатн" in q
    # Теги — значимые слова для уточнения (место, тема, название)
    stop = {"в", "на", "по", "и", "или", "с", "к", "из", "для", "куда",
            "сходить", "пойти", "найди", "покажи", "хочу", "можно", "мне",
            "нам", "нас", "что", "где", "когда", "как", "это", "есть",
            "можно", "хочу", "хочется", "посетить", "посмотреть"}
    tags = []
    for w in q.split():
        w_clean = w.strip(".,!?")
        if (len(w_clean) > 3 and w_clean not in stop
                and not any(city in w_clean for city in _KUDAGO_CITIES)
                and not any(kw in w_clean for kw in ["концерт", "выставк", "театр",
                            "фестив", "квест", "мероприят", "афиша", "вечеринк"])):
            tags.append(w_clean)
        if len(tags) >= 3:
            break
    return {"location": location, "categories": cats, "tags": tags, "is_free": is_free}

def _kudago_fetch(parsed: dict) -> list:
    import time
    now = int(time.time())
    base_params = {
        "lang": "ru",
        "location": parsed["location"],
        "actual_since": now,
        "page_size": 15,
        "fields": "id,title,place,dates,site_url,price,is_free,description,slug,location",
        "expand": "place,dates",
        "text_format": "text",
        "order_by": "-publication_date",
    }
    if parsed["categories"]:
        base_params["categories"] = ",".join(parsed["categories"])
    if parsed["is_free"]:
        base_params["is_free"] = "true"

    results = []
    # Запрос 1: с тегами (максимально точный)
    try:
        p = dict(base_params)
        if parsed["tags"]:
            p["tags"] = ",".join(parsed["tags"])
        r = requests.get("https://kudago.com/public-api/v1.4/events/", params=p, timeout=10)
        if r.ok:
            results = r.json().get("results", [])
            print(f"[kudago] events_with_tags={len(results)}")
    except Exception as e:
        print(f"[kudago] events error: {e}")

    # Запрос 2: без тегов если мало (шире)
    if len(results) < 4 and parsed["tags"]:
        try:
            r2 = requests.get("https://kudago.com/public-api/v1.4/events/", params=base_params, timeout=10)
            if r2.ok:
                extra = r2.json().get("results", [])
                seen = {e.get("id") for e in results}
                for e in extra:
                    if e.get("id") not in seen:
                        results.append(e)
                print(f"[kudago] events_wide={len(extra)}")
        except Exception as e:
            print(f"[kudago] events_wide error: {e}")

    # Запрос 3: текстовый поиск если совсем мало
    if len(results) < 3 and parsed["tags"]:
        try:
            sq = " ".join(parsed["tags"])
            r3 = requests.get(
                "https://kudago.com/public-api/v1.4/search/",
                params={"q": sq, "type": "event", "location": parsed["location"],
                        "expand": "dates", "page_size": 10, "lang": "ru"},
                timeout=10
            )
            if r3.ok:
                extra2 = r3.json().get("results", [])
                seen = {e.get("id") for e in results}
                for e in extra2:
                    if e.get("id") not in seen:
                        results.append(e)
                print(f"[kudago] search_extra={len(extra2)}")
        except Exception as e:
            print(f"[kudago] search error: {e}")

    return results[:15]

def _kudago_format_date(start_ts) -> str:
    try:
        from datetime import datetime
        dt = datetime.fromtimestamp(int(start_ts))
        day = dt.strftime("%-d")
        month = _MONTHS_RU.get(dt.strftime("%B"), dt.strftime("%B"))
        year = dt.strftime("%Y")
        time_str = dt.strftime("%H:%M")
        return f"{day} {month} {year}, {time_str}"
    except:
        return ""

def _kudago_format(e: dict) -> dict:
    title = (e.get("title") or e.get("short_title") or "").strip()
    # Дата
    date_str = ""
    for d in (e.get("dates") or []):
        start = d.get("start")
        if start:
            date_str = _kudago_format_date(start)
            break
    # Место
    place = e.get("place") or {}
    place_name = place.get("title", "") if isinstance(place, dict) else ""
    address = place.get("address", "") if isinstance(place, dict) else ""
    # Цена
    price_raw = (e.get("price") or "").strip()
    if e.get("is_free"):
        price = "Бесплатно"
    elif price_raw:
        price = price_raw[:80]
    else:
        price = None
    # URL
    slug = e.get("slug", "")
    loc_info = e.get("location") or {}
    loc_slug = loc_info.get("slug", "msk") if isinstance(loc_info, dict) else "msk"
    url = e.get("site_url") or f"https://kudago.com/{loc_slug}/event/{slug}/"
    # Описание
    desc = (e.get("description") or "")[:200].strip()
    parts = []
    if date_str:
        parts.append(f"📅 {date_str}")
    if place_name:
        addr_part = f", {address}" if address else ""
        parts.append(f"📍 {place_name}{addr_part}")
    if desc:
        parts.append(desc)
    return {
        "title": title[:120],
        "description": "\n".join(parts),
        "price": price,
        "date": date_str,
        "place": place_name,
        "address": address,
        "source": "kudago.com",
        "url": url,
        "cat": "events",
    }

@app.post("/api/v1/search")
async def search_agent(request: Request):
    """Агент-поисковик: KudaGo (мероприятия) + Wildberries + Tavily + AI (синхронный, надёжный)"""
    import json as _json, re as _re, asyncio as _asyncio
    data = await request.json()
    query = (data.get("query") or "").strip()
    image_base64 = data.get("image_base64")
    mime_type = data.get("mime_type", "image/jpeg")

    print(f"[search] query='{query}', image={bool(image_base64)}")

    # ── 1. ФОТО → QUERY ──────────────────────────────────────────
    if image_base64 and not query:
        try:
            resp = requests.post(
                f"{AITUNNEL_BASE_URL}chat/completions",
                headers={"Authorization": f"Bearer {AITUNNEL_API_KEY}", "Content-Type": "application/json"},
                json={"model": "gemini-2.5-flash-lite", "max_tokens": 80, "messages": [{"role": "user", "content": [
                    {"type": "image_url", "image_url": {"url": f"data:{mime_type};base64,{image_base64}"}},
                    {"type": "text", "text": "Что на фото? Дай поисковый запрос 3-5 слов. Только запрос."}
                ]}]}, timeout=10
            )
            query = resp.json()["choices"][0]["message"]["content"].strip() or "товар"
        except Exception as e:
            print(f"[search] photo recognition failed: {e}")
            query = "товар"

    if not query:
        return {"results": [], "query": "", "error": "Пустой запрос"}

    # ── 2. КЛАССИФИКАЦИЯ — быстро без AI ─────────────────────────
    q_lower = query.lower()
    # KudaGo имеет приоритет — проверяем первым
    if _is_events_query(query):
        cat = "events"
    elif _is_info_query(query):
        cat = "info"
    elif any(w in q_lower for w in ["кино", "музе", "ресторан", "афиша", "билет"]):
        cat = "leisure"
    elif any(w in q_lower for w in ["мастер", "сантехник", "электрик", "ремонт", "клинер", "грузчик", "курьер", "услуг", "няня", "репетитор"]):
        cat = "services"
    else:
        cat = "goods"

    print(f"[search] category={cat}")

    raw_results = []

    # ── 3i. INFO — информационный вопрос: прямой AI-ответ + источники ──
    if cat == "info":
        data = tavily_search_adv(query, topic="general", max_results=5, include_answer=True)
        answer = (data.get("answer") or "").strip()
        items = data.get("results", []) or []
        # Если есть топ-источник — обогащаем ответ полным текстом через Extract
        if answer and items:
            top_url = items[0].get("url", "")
            if top_url:
                deep = tavily_extract(top_url, max_chars=3000)
                if deep:
                    try:
                        enrich = call_ai([{"role": "user", "content": (
                            "Ответь на вопрос простым русским языком для пожилого человека, "
                            "коротко и по делу (3-5 предложений), на основе ответа и источника.\n\n"
                            f"ВОПРОС: {query}\n\nКраткий ответ: {answer}\n\n"
                            f"Полный текст источника:\n{deep[:2500]}\n\n"
                            "Дай понятный практичный ответ без воды и без markdown."
                        )}])
                        if enrich and len(enrich.strip()) > 20:
                            answer = enrich.strip()
                    except Exception as e:
                        print(f"[search info] enrich error: {e}")
        if answer:
            raw_results.append({
                "title": "Ответ на ваш вопрос",
                "description": answer[:900],
                "is_answer": True,
                "source": "OkTolk", "url": "", "cat": "info"
            })
        for it in items[:4]:
            url = it.get("url", "")
            domain = url.split("/")[2].replace("www.", "") if url.startswith("http") else "источник"
            raw_results.append({
                "title": (it.get("title") or "")[:120],
                "description": (it.get("content") or "")[:200],
                "source": domain, "url": url, "cat": "info"
            })

    # ── 3a. KUDAGO — мероприятия и события ───────────────────────
    if cat == "events":
        parsed_kg = _parse_kudago_params(query)
        print(f"[kudago] params={parsed_kg}")
        kg_events = _kudago_fetch(parsed_kg)
        print(f"[kudago] total={len(kg_events)}")
        raw_results = [_kudago_format(e) for e in kg_events if (e.get("title") or "").strip()]
        if not raw_results:
            print("[kudago] empty — fallback to leisure/Tavily")
            cat = "leisure"  # fallback если KudaGo ничего не нашёл

    # ── 3. WILDBERRIES (для товаров) ─────────────────────────────
    if cat == "goods":
        wb_urls = [
            "https://search.wb.ru/exactmatch/ru/common/v5/search",
            "https://search.wb.ru/exactmatch/ru/common/v4/search",
        ]
        for wb_url in wb_urls:
            if raw_results:
                break
            try:
                r = requests.get(
                    wb_url,
                    params={
                        "query": query, "resultset": "catalog", "limit": 6,
                        "sort": "popular", "lang": "ru", "curr": "rub",
                        "dest": "-1257786",
                        "spp": "30",
                        "appType": "1",
                    },
                    headers={
                        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120 Safari/537.36",
                        "Accept": "*/*",
                        "Accept-Language": "ru-RU,ru;q=0.9",
                        "Origin": "https://www.wildberries.ru",
                        "Referer": "https://www.wildberries.ru/",
                    },
                    timeout=10
                )
                print(f"[search] WB status={r.status_code}")
                if r.status_code == 200:
                    wb_data = r.json()
                    # Диагностика — ключи ответа
                    print(f"[search] WB top keys: {list(wb_data.keys())[:5]}")
                    products = (wb_data.get("data") or {}).get("products", []) or wb_data.get("products", [])
                    if products:
                        # Логируем ключи первого товара
                        print(f"[search] WB product[0] keys: {list(products[0].keys())[:10]}")
                    print(f"[search] WB products: {len(products)}")
                    for p in products[:6]:
                        pid = p.get("id")
                        name = p.get("name", "")
                        brand = p.get("brand", "")
                        title = f"{brand} {name}".strip()
                        if not pid or not title:
                            print(f"[search] WB skip: pid={pid} title='{title}'")
                            continue
                        # Цены: проверяем все возможные форматы
                        price = p.get("salePriceU") or p.get("priceU") or p.get("sale") or 0
                        price_old = p.get("priceU") or 0
                        # v5: цены в sizes[0].price
                        if not price and p.get("sizes"):
                            try:
                                sz = p["sizes"][0]
                                pp = sz.get("price") or {}
                                price = pp.get("product") or pp.get("total") or 0
                                price_old = pp.get("basic") or 0
                            except: pass
                        rating = p.get("rating") or p.get("reviewRating") or 0
                        reviews = p.get("feedbacks") or p.get("nmReviewRating") or 0
                        raw_results.append({
                            "title": title[:120],
                            "description": f"Рейтинг {rating}/5 · {reviews} отзывов" if reviews else "",
                            "price": f"{int(price/100):,} ₽".replace(",", " ") if price > 100 else None,
                            "old_price": f"{int(price_old/100):,} ₽".replace(",", " ") if price_old > price else None,
                            "source": "wildberries.ru",
                            "url": f"https://www.wildberries.ru/catalog/{pid}/detail.aspx",
                            "cat": "goods"
                        })
                        print(f"[search] WB added: {title[:50]} price={price}")
            except Exception as e:
                print(f"[search] WB error: {e}")
                continue

        # Проверяем качество WB результатов — если все без title, сбрасываем
        if raw_results and all(not r.get("title") for r in raw_results):
            print("[search] WB results have empty titles, discarding")
            raw_results = []

    # ── 4. TAVILY — для досуга/услуг всегда, для товаров если WB пустой ──
    # events — не используем Tavily, только KudaGo
    if cat not in ("events", "info") and (len(raw_results) < 3 or cat in ["leisure", "services"]):
        # Для товаров — естественный запрос с упоминанием маркетплейса
        if cat == "goods":
            tq = f"{query} купить цена ozon wildberries"
        elif cat == "leisure":
            tq = f"{query} site:afisha.ru OR site:kinopoisk.ru OR site:kassir.ru OR site:kudago.com"
        else:  # services
            tq = f"{query} site:profi.ru OR site:youdo.ru"

        print(f"[search] Tavily query: {tq}")
        try:
            r = requests.post(
                "https://api.tavily.com/search",
                json={
                    "api_key": TAVILY_API_KEY,
                    "query": tq,
                    "search_depth": "basic",
                    "max_results": 6,
                    "include_answer": False,
                    "country": "ru",
                },
                timeout=15
            )
            print(f"[search] Tavily status={r.status_code}")
            if r.status_code == 200:
                td = r.json()
                items = td.get("results", [])
                print(f"[search] Tavily results: {len(items)}")
                for it in items:
                    url = it.get("url", "")
                    domain = url.split("/")[2].replace("www.", "") if url.startswith("http") else "источник"
                    description = (it.get("content") or "")[:250]
                    # Извлекаем цены из описания
                    price = None
                    old_price = None
                    # Ищем все цены в формате "1 234 ₽" или "1234₽" или "1234 руб"
                    prices = _re.findall(r'(\d[\d\s]{2,}\s*[₽рРP])', description)
                    if prices:
                        # Чистим: убираем не-цифры кроме разделителей
                        clean = [_re.sub(r'\s+', ' ', p).strip() for p in prices[:2]]
                        price = clean[0]
                        if len(clean) > 1:
                            # Меньшая - текущая, бОльшая - старая
                            n0 = int(_re.sub(r'\D', '', clean[0]) or 0)
                            n1 = int(_re.sub(r'\D', '', clean[1]) or 0)
                            if n0 < n1:
                                price, old_price = clean[0], clean[1]
                            elif n1 < n0:
                                price, old_price = clean[1], clean[0]
                    # Чистим title от хвостов типа " - Wildberries", " - Купить в интернет ..."
                    title = (it.get("title") or "").strip()
                    title = _re.sub(r'\s*[-–—|]\s*(Купить|Wildberries|Ozon|Купить в.*)$', '', title, flags=_re.IGNORECASE)
                    title = _re.sub(r'\s*\.\.\.\s*-?\s*\w+$', '', title)
                    raw_results.append({
                        "title": title[:120],
                        "description": description,
                        "price": price, "old_price": old_price,
                        "source": domain, "url": url, "cat": cat
                    })
            else:
                print(f"[search] Tavily error body: {r.text[:200]}")
        except Exception as e:
            print(f"[search] Tavily error: {e}")

    if not raw_results:
        return {"results": [], "query": query, "error": "Ничего не нашлось"}

    # Фильтруем пустые результаты
    raw_results = [r for r in raw_results if r.get("title") and len(r["title"].strip()) > 2]
    if not raw_results:
        return {"results": [], "query": query, "error": "Результаты некачественные"}

    print(f"[search] returning {len(raw_results)} valid results")
    return {"results": raw_results[:6], "query": query, "cat": cat}

@app.get("/manifest.json")
async def manifest():
    """PWA Manifest — fullscreen скрывает браузер и системную навигацию Android"""
    from fastapi.responses import JSONResponse
    return JSONResponse({
        "name": "OkTolk — AI Помощник",
        "short_name": "OkTolk",
        "description": "Простой AI-помощник: защита от мошенников, здоровье, финансы, поиск",
        "start_url": "/",
        "scope": "/",
        "display": "fullscreen",
        "display_override": ["fullscreen", "standalone", "minimal-ui"],
        "orientation": "portrait",
        "background_color": "#FAFAF7",
        "theme_color": "#3B4FE0",
        "lang": "ru",
        "icons": [
            {"src": "/icon-192.png", "sizes": "192x192", "type": "image/png", "purpose": "any maskable"},
            {"src": "/icon-512.png", "sizes": "512x512", "type": "image/png", "purpose": "any maskable"}
        ],
        "screenshots": [],
        "categories": ["utilities", "productivity"],
        "prefer_related_applications": False
    }, headers={"Content-Type": "application/manifest+json", "Cache-Control": "no-cache"})

@app.get("/icon-192.png")
@app.get("/icon-512.png")
async def icon(request: Request):
    """Временная SVG иконка как PNG — замените на реальные PNG файлы"""
    import base64
    # Простая SVG иконка OkTolk в виде PNG (base64 1x1 transparent)
    # Для продакшна — положить реальные PNG файлы в папку проекта
    path = request.url.path
    size = 192 if "192" in path else 512
    svg = f'''<svg width="{size}" height="{size}" xmlns="http://www.w3.org/2000/svg">
      <rect width="{size}" height="{size}" rx="{size//5}" fill="#3B4FE0"/>
      <text x="50%" y="56%" font-family="Arial" font-size="{size//3}" font-weight="bold"
        fill="white" text-anchor="middle" dominant-baseline="middle">Ok</text>
    </svg>'''
    from fastapi.responses import Response
    return Response(content=svg.encode(), media_type="image/svg+xml",
                   headers={"Cache-Control": "public, max-age=86400"})

async def get_sites():
    return [
        {"name": "Госуслуги", "url": "https://gosuslugi.ru", "description": "Государственные услуги онлайн"},
        {"name": "Пенсионный фонд", "url": "https://sfr.gov.ru", "description": "Социальный фонд России"},
        {"name": "МФЦ", "url": "https://mfc.ru", "description": "Мои документы — центр услуг"},
        {"name": "Банк России", "url": "https://cbr.ru", "description": "Центральный банк — проверка банков"},
        {"name": "МВД", "url": "https://mvd.ru", "description": "Сообщить о мошенниках"},
        {"name": "Минздрав", "url": "https://minzdrav.gov.ru", "description": "Здоровье и медицина"},
        {"name": "ФНС", "url": "https://nalog.gov.ru", "description": "Федеральная налоговая служба"}
    ]


import base64

# ── STT (Speech-to-Text) ─────────────────────────────────────────
class STTRequest(BaseModel):
    audio_base64: str
    mime_type: str = "audio/webm"

@app.post("/api/v1/stt")
async def speech_to_text(req: STTRequest):
    if not AITUNNEL_API_KEY:
        raise HTTPException(status_code=503, detail="AI недоступен")
    try:
        headers = {
            "Authorization": f"Bearer {AITUNNEL_API_KEY}",
            "Content-Type": "application/json"
        }
        data = {
            "model": "gemini-2.5-flash-lite",
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "input_audio",
                            "input_audio": {
                                "data": req.audio_base64,
                                "format": "webm"
                            }
                        },
                        {
                            "type": "text",
                            "text": "Транскрибируй точно что сказано. Верни только текст, без комментариев."
                        }
                    ]
                }
            ],
            "max_tokens": 500
        }
        response = requests.post(
            f"{AITUNNEL_BASE_URL}chat/completions",
            headers=headers,
            json=data,
            timeout=30
        )
        result = response.json()
        text = result["choices"][0]["message"]["content"].strip()
        return {"text": text, "status": "ok"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка STT: {str(e)}")


# ── TTS (Text-to-Speech) Yandex SpeechKit ───────────────────────
from fastapi.responses import StreamingResponse
import io

class TTSRequest(BaseModel):
    text: str
    voice: str = "alena"

@app.post("/api/v1/tts")
async def text_to_speech(req: TTSRequest):
    if not YANDEX_API_KEY:
        raise HTTPException(status_code=503, detail="Яндекс TTS недоступен")
    try:
        params = {
            "text": req.text[:5000],
            "lang": "ru-RU",
            "voice": req.voice,
            "speed": "1.0",
            "format": "mp3",
            "sampleRateHertz": "48000"
        }
        headers = {
            "Authorization": f"Api-Key {YANDEX_API_KEY}"
        }
        response = requests.post(
            "https://tts.api.cloud.yandex.net/speech/v1/tts:synthesize",
            headers=headers,
            data=params,
            timeout=30
        )
        if response.status_code != 200:
            raise HTTPException(status_code=500, detail=f"Ошибка Яндекс TTS: {response.text}")
        return StreamingResponse(
            io.BytesIO(response.content),
            media_type="audio/mpeg",
            headers={"Content-Disposition": "inline; filename=speech.mp3"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка TTS: {str(e)}")

# ── Chat with Image ───────────────────────────────────────────────
class ImageChatRequest(BaseModel):
    image_base64: str
    mime_type: str = "image/jpeg"
    context: Optional[str] = None

@app.post("/api/v1/chat/image")
async def chat_with_image(req: ImageChatRequest):
    if not AITUNNEL_API_KEY:
        raise HTTPException(status_code=503, detail="AI недоступен")
    try:
        headers = {
            "Authorization": f"Bearer {AITUNNEL_API_KEY}",
            "Content-Type": "application/json"
        }
        data = {
            "model": "gemini-2.5-flash-lite",
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:{req.mime_type};base64,{req.image_base64}"
                            }
                        },
                        {
                            "type": "text",
                            "text": "Извлеки ТОЛЬКО факты из переписки или сообщения на скриншоте. Игнорируй интерфейс приложений и телефона. Анализируй только текст переписки: что предлагают, какие суммы, что просят, какие обещания. Только факты списком."
                        }
                    ]
                }
            ],
            "max_tokens": 500
        }
        response = requests.post(
            f"{AITUNNEL_BASE_URL}chat/completions",
            headers=headers,
            json=data,
            timeout=30
        )
        result = response.json()
        facts = result["choices"][0]["message"]["content"].strip()

        # Шаг 2 — антискам анализ фактов
        antiscam_data = {
            "model": "gemini-2.5-flash-lite",
            "messages": [{"role": "user", "content": f"""Ты антискам система. Проанализируй факты и дай краткий вердикт.

ФАКТЫ:
{facts}

Ответь ТОЛЬКО JSON:
{{"risk_level": 0, "summary": "1 предложение - есть угроза или нет", "how_they_manipulate": "как пытаются обмануть или null", "consequences": "к чему может привести или null", "recommendation": "одно конкретное действие"}}

ШКАЛА: 0=безопасно 1=минимальный 2=подозрения 3=высокий 4=очень высокий 5=явное мошенничество
Только JSON."""}],
            "max_tokens": 300
        }
        antiscam_response = requests.post(
            f"{AITUNNEL_BASE_URL}chat/completions",
            headers=headers,
            json=antiscam_data,
            timeout=30
        )
        antiscam_result = antiscam_response.json()
        reply_text = antiscam_result["choices"][0]["message"]["content"].strip()

        import re as re2
        json_match = re2.search(r'\{.*\}', reply_text, re2.DOTALL)
        if json_match:
            parsed = json.loads(json_match.group())
            return parsed
        return {"reply": reply_text}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка: {str(e)}")

# ══════════════════════════════════════════════════════
# СТОП МОШЕННИК — антискам модуль
# ══════════════════════════════════════════════════════

class AntiscamExtractRequest(BaseModel):
    """Извлечение фактов из материала"""
    image_base64: Optional[str] = None
    mime_type: str = "image/jpeg"
    audio_base64: Optional[str] = None
    audio_mime: str = "audio/webm"
    url: Optional[str] = None
    text: Optional[str] = None

class AntiscamAnalyzeRequest(BaseModel):
    """Анализ накопленного контекста"""
    facts: str  # все извлечённые факты через \n---\n
    question: Optional[str] = None  # уточняющий вопрос пользователя

def call_gemini_vision(image_base64: str, mime_type: str, prompt: str) -> str:
    """Вызов Gemini с изображением"""
    headers = {"Authorization": f"Bearer {AITUNNEL_API_KEY}", "Content-Type": "application/json"}
    data = {
        "model": "gemini-2.5-flash-lite",
        "messages": [{"role": "user", "content": [
            {"type": "image_url", "image_url": {"url": f"data:{mime_type};base64,{image_base64}"}},
            {"type": "text", "text": prompt}
        ]}],
        "max_tokens": 800
    }
    r = requests.post(f"{AITUNNEL_BASE_URL}chat/completions", headers=headers, json=data, timeout=30)
    return r.json()["choices"][0]["message"]["content"].strip()

def call_gemini_audio(audio_base64: str, mime_type: str) -> str:
    """Транскрипция аудио через Gemini"""
    headers = {"Authorization": f"Bearer {AITUNNEL_API_KEY}", "Content-Type": "application/json"}
    data = {
        "model": "gemini-2.5-flash-lite",
        "messages": [{"role": "user", "content": [
            {"type": "input_audio", "input_audio": {"data": audio_base64, "format": mime_type.split("/")[-1]}},
            {"type": "text", "text": "Транскрибируй полностью. Раздели реплики по говорящим если их несколько. Верни только текст разговора."}
        ]}],
        "max_tokens": 1000
    }
    r = requests.post(f"{AITUNNEL_BASE_URL}chat/completions", headers=headers, json=data, timeout=60)
    return r.json()["choices"][0]["message"]["content"].strip()

def call_gemini_text(prompt: str, max_tokens: int = 600) -> str:
    """Вызов Gemini только с текстом"""
    headers = {"Authorization": f"Bearer {AITUNNEL_API_KEY}", "Content-Type": "application/json"}
    data = {
        "model": "gemini-2.5-flash-lite",
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": max_tokens
    }
    r = requests.post(f"{AITUNNEL_BASE_URL}chat/completions", headers=headers, json=data, timeout=30)
    return r.json()["choices"][0]["message"]["content"].strip()

EXTRACT_PROMPT = """Ты извлекаешь факты из скриншота переписки или сообщения для анализа на мошенничество.

ВАЖНО: Игнорируй интерфейс приложений, браузеров, телефона — анализируй ТОЛЬКО содержимое переписки или сообщения.
Если на скриншоте видно приложение OkTolk, мессенджер, браузер — анализируй только текст переписки внутри, не сам интерфейс.

Извлеки из переписки/сообщения:
- Что предлагают или просят сделать
- Какие суммы или ценности упоминаются  
- Какие обещания дают
- Какие данные запрашивают (карта, пароль, код из СМС)
- Признаки давления или срочности
- Имена, ники, контакты отправителя
- Ссылки если есть

Если на изображении нет переписки или сообщения — напиши: "Переписка не обнаружена"
Формат: короткий список фактов. Только содержимое переписки."""

@app.post("/api/v1/antiscam/extract")
async def antiscam_extract(req: AntiscamExtractRequest):
    """Этап 1 — извлечение фактов из любого материала"""
    if not AITUNNEL_API_KEY:
        raise HTTPException(status_code=503, detail="AI недоступен")
    try:
        facts = ""

        if req.image_base64:
            facts = call_gemini_vision(req.image_base64, req.mime_type, EXTRACT_PROMPT)

        elif req.audio_base64:
            transcript = call_gemini_audio(req.audio_base64, req.audio_mime)
            facts = call_gemini_text(f"Из этой транскрипции извлеки факты для анализа на мошенничество:\n{transcript}\n\n{EXTRACT_PROMPT}")

        elif req.url:
            page_text = tavily_extract(req.url)
            if page_text:
                facts = call_gemini_text(
                    f"Из содержимого этой веб-страницы извлеки факты для анализа на мошенничество:\n{page_text}\n\n{EXTRACT_PROMPT}"
                )
            else:
                # Fallback — прямая загрузка если Tavily недоступен
                try:
                    resp = requests.get(req.url, timeout=10, headers={"User-Agent": "Mozilla/5.0"})
                    html = resp.text[:5000]
                    facts = call_gemini_text(f"Из этого HTML извлеки факты для анализа на мошенничество:\n{html}\n\n{EXTRACT_PROMPT}")
                except:
                    facts = f"URL: {req.url} — не удалось загрузить страницу. Сама ссылка выглядит так: {req.url}"

        elif req.text:
            facts = req.text

        return {"facts": facts, "status": "ok"}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка извлечения: {str(e)}")


ANALYZE_PROMPT = """Ты — антискам система OkTolk. Проанализируй собранные факты и дай структурированный вердикт.

НАКОПЛЕННЫЕ ФАКТЫ ИЗ ВСЕХ МАТЕРИАЛОВ:
{facts}

{question_block}

Найди связи между фактами. Определи схему мошенничества если она есть.

Ответь ТОЛЬКО валидным JSON:
{{
  "risk_level": 0,
  "what_doing": "Коротко что делает/пытается сделать мошенник. 1-2 предложения.",
  "risk_description": "Описание рисков. 1 предложение.",
  "consequences": "Что произойдёт если продолжить общение. 1-2 предложения.",
  "recommendation": "Конкретное действие прямо сейчас. 1 предложение.",
  "scheme": "название схемы мошенничества или null",
  "confidence": "высокая или средняя или низкая"
}}

ШКАЛА РИСКА: 0=безопасно 1=минимальный 2=подозрения 3=высокий 4=очень высокий 5=явное мошенничество
Только JSON."""

@app.post("/api/v1/antiscam/analyze")
async def antiscam_analyze(req: AntiscamAnalyzeRequest):
    """Этап 2 — анализ всего накопленного контекста"""
    if not AITUNNEL_API_KEY:
        raise HTTPException(status_code=503, detail="AI недоступен")
    try:
        question_block = f"ВОПРОС ПОЛЬЗОВАТЕЛЯ: {req.question}" if req.question else ""
        prompt = ANALYZE_PROMPT.format(facts=req.facts, question_block=question_block)
        ds_key = DEEPSEEK_API_KEY or AITUNNEL_API_KEY
        ds_base = DEEPSEEK_BASE_URL if DEEPSEEK_API_KEY else AITUNNEL_BASE_URL
        ds_model = "deepseek-reasoner" if DEEPSEEK_API_KEY else "deepseek-r1-0528"
        ds_headers = {"Authorization": f"Bearer {ds_key}", "Content-Type": "application/json"}
        ds_data = {"model": ds_model, "messages": [{"role": "user", "content": prompt}], "max_tokens": 1000}
        ds_resp = requests.post(f"{ds_base}chat/completions", headers=ds_headers, json=ds_data, timeout=60)
        result_text = ds_resp.json()["choices"][0]["message"]["content"].strip()

        import re as re_mod
        json_match = re_mod.search(r"\{.*\}", result_text, re_mod.DOTALL)
        if json_match:
            return json.loads(json_match.group())
        return {"error": "parse_error", "raw": result_text}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка анализа: {str(e)}")


class AntiscamChatRequest(BaseModel):
    facts: str
    question: str
    history: Optional[list] = []

@app.post("/api/v1/antiscam/chat")
async def antiscam_chat(req: AntiscamChatRequest):
    """Диалог в контексте антискам анализа — DeepSeek прямой"""
    try:
        system = """Ты консультант по защите от мошенников OkTolk.
Ты уже проанализировал материалы пользователя и знаешь контекст ситуации.
Действуй как опытный следователь и помощник одновременно.

ПРАВИЛА:
- Если вопрос простой — отвечай коротко и по делу
- Если ситуация сложная — задай уточняющий вопрос чтобы дать точный совет
- Если человек уже перевёл деньги или передал данные — дай пошаговый план действий прямо сейчас
- Если человек спрашивает что делать — давай конкретные действия, не общие слова
- Говори простым языком без терминов
- НЕ используй markdown: никаких **, *, ##, ###, ---, нумерованных списков с точками
- Если нужен список — пиши каждый пункт с новой строки через цифру и точку: "1. 2. 3."
- Никогда не говори "я не могу помочь" — всегда давай конкретный совет или задай уточняющий вопрос

ЕСЛИ ЧЕЛОВЕК УЖЕ ПОСТРАДАЛ:
1. Сначала скажи что нужно сделать ПРЯМО СЕЙЧАС (заблокировать карту, сменить пароль и т.д.)
2. Потом объясни следующие шаги (куда обратиться, как вернуть деньги)
3. Предупреди о возможных дальнейших попытках мошенников"""

        messages = [{"role": "system", "content": system}]
        messages.append({"role": "user", "content": f"Контекст проверки:\n{req.facts}"})
        messages.append({"role": "assistant", "content": "Понял, я изучил материалы. Задавайте вопросы."})
        for h in (req.history or [])[-6:]:
            if h.get("role") and h.get("content"):
                messages.append({"role": h["role"], "content": h["content"]})
        messages.append({"role": "user", "content": req.question})

        reply = _ds_post(messages[1:], system=system, max_tokens=400)
        return {"reply": reply}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка: {str(e)}")


# ── Finance AI Parser ────────────────────────────────────────────
class FinanceParseRequest(BaseModel):
    text: Optional[str] = None
    image_base64: Optional[str] = None
    mime_type: str = "image/jpeg"
    user_id: Optional[int] = None  # берётся из токена

class DocumentAnalyzeRequest(BaseModel):
    file_base64: str
    mime_type: str = "application/pdf"
    file_name: str = ""
    context: str = ""
    user_question: str = ""


@app.post("/api/v1/document/analyze")
async def analyze_document(req: DocumentAnalyzeRequest, user=Depends(get_current_user)):
    """Универсальный анализ документа (квитанция ЖКУ, договор кредита, медкарта).
    Извлекает данные → пропускает через специализированного эксперта → понятный ответ."""
    if not AITUNNEL_API_KEY:
        raise HTTPException(status_code=503, detail="AI недоступен")
    if not req.file_base64:
        raise HTTPException(status_code=400, detail="Файл не передан")

    context = req.context or "general"
    extract_prompts = {
        "utility": (
            "Это квитанция за ЖКУ. Извлеки и аккуратно перепиши все ключевые данные:\n"
            "— период; адрес; общая сумма к оплате;\n"
            "— все статьи начислений отдельно с тарифом и суммой "
            "(содержание, ХВС, ГВС, водоотведение, отопление, электроэнергия по тарифам, "
            "газ, ОДН/КР на СОИ, капремонт, вывоз ТКО, домофон и др.);\n"
            "— показания счётчиков если есть;\n"
            "— долг/переплата за прошлые периоды;\n"
            "— пени если есть.\n"
            "Верни структурированный текст по-русски, БЕЗ markdown."
        ),
        "credit": (
            "Это финансовый документ (кредитный договор, график платежей или справка). Извлеки:\n"
            "— тип документа; банк/кредитор; ФИО заёмщика;\n"
            "— сумма кредита; ставка; срок; ежемесячный платёж; полная стоимость кредита (ПСК);\n"
            "— дата выдачи и дата окончания;\n"
            "— страховка (название, сумма, обязательная или нет);\n"
            "— скрытые комиссии или дополнительные платежи;\n"
            "— штрафы за просрочку.\n"
            "Верни структурированный текст по-русски, БЕЗ markdown."
        ),
        "pharmacy": (
            "Это медицинский документ (рецепт, выписка, чек из аптеки или направление). Извлеки:\n"
            "— тип документа;\n"
            "— названия лекарств с дозировкой и количеством;\n"
            "— цены если указаны;\n"
            "— назначения врача;\n"
            "— по рецепту ли (для налогового вычета).\n"
            "НЕ интерпретируй диагнозы. Верни структурированный текст по-русски, БЕЗ markdown."
        ),
    }
    extract_prompt = extract_prompts.get(context, (
        "Это документ. Извлеки и структурированно перепиши все значимые данные: "
        "тип документа, основные цифры, даты, суммы, стороны, условия. Без markdown."
    ))

    try:
        fname = (req.file_name or "").lower()
        extracted = ""

        # ── Текстовые форматы извлекаем локально (без Gemini) ──
        if fname.endswith(".txt") or req.mime_type == "text/plain":
            try:
                raw = base64.b64decode(req.file_base64)
                extracted = raw.decode("utf-8", errors="ignore")[:8000]
            except Exception as e:
                print(f"[document/analyze] txt error: {e}")

        elif fname.endswith(".docx") or "wordprocessingml" in req.mime_type:
            try:
                import docx
                import io as _io
                d = docx.Document(_io.BytesIO(base64.b64decode(req.file_base64)))
                parts = [p.text for p in d.paragraphs if p.text.strip()]
                # таблицы тоже
                for tbl in d.tables:
                    for row in tbl.rows:
                        cells = [c.text.strip() for c in row.cells if c.text.strip()]
                        if cells:
                            parts.append(" | ".join(cells))
                extracted = "\n".join(parts)[:8000]
            except Exception as e:
                print(f"[document/analyze] docx error: {e}")

        elif fname.endswith((".xlsx", ".xls")) or "spreadsheetml" in req.mime_type:
            try:
                import openpyxl
                import io as _io
                wb = openpyxl.load_workbook(_io.BytesIO(base64.b64decode(req.file_base64)), read_only=True, data_only=True)
                rows = []
                for ws in wb.worksheets:
                    rows.append(f"Лист: {ws.title}")
                    for row in ws.iter_rows(values_only=True):
                        cells = [str(c) for c in row if c is not None]
                        if cells:
                            rows.append("\t".join(cells))
                        if len(rows) > 200:
                            break
                extracted = "\n".join(rows)[:8000]
            except Exception as e:
                print(f"[document/analyze] xlsx error: {e}")

        # ── PDF и изображения читаем через Gemini Vision ──
        if not extracted:
            is_pdf = (req.mime_type == "application/pdf"
                      or fname.endswith(".pdf"))
            if is_pdf:
                content = [
                    {"type": "text", "text": extract_prompt},
                    {"type": "file", "file": {"file_data": f"data:application/pdf;base64,{req.file_base64}"}},
                ]
            else:
                content = [
                    {"type": "text", "text": extract_prompt},
                    {"type": "image_url", "image_url": {"url": f"data:{req.mime_type};base64,{req.file_base64}"}},
                ]

            def _call_extract():
                headers = {"Authorization": f"Bearer {AITUNNEL_API_KEY}", "Content-Type": "application/json"}
                data = {
                    "model": "gemini-2.5-flash-lite",
                    "messages": [{"role": "user", "content": content}],
                    "max_tokens": 1500,
                }
                r = requests.post(f"{AITUNNEL_BASE_URL}chat/completions", headers=headers, json=data, timeout=60)
                return r.json()["choices"][0]["message"]["content"]

            extracted = await asyncio.to_thread(_call_extract)

        if not extracted or len(extracted.strip()) < 20:
            return {"status": "error", "message": "Не удалось прочитать документ. Попробуйте более чёткое фото или другой формат файла."}

        print(f"[document/analyze] context={context}, формат={fname.split('.')[-1] if '.' in fname else req.mime_type}, extracted={len(extracted)} символов")

        expert_prompt = FINANCE_CATEGORY_PROMPTS.get(context, FINANCE_CHAT_PROMPT)
        profile_ctx = await build_profile_context(user["id"]) if user else ""
        user_question = req.user_question.strip() or "Разбери документ: что важного, есть ли подозрительные суммы или возможность оптимизировать."

        messages = [
            {"role": "system", "content": expert_prompt + profile_ctx},
            {"role": "user", "content": f"Документ:\n{extracted}\n\nВопрос пользователя: {user_question}"},
        ]
        reply = await asyncio.to_thread(call_ai, messages)
        await add_tokens(user["id"], extracted + user_question, reply)
        await save_chat_message(user["id"], "finance", "user", f"[документ: {req.file_name or 'файл'}] {user_question}")
        await save_chat_message(user["id"], "finance", "ai", reply)

        return {"status": "ok", "extracted": extracted, "reply": reply}
    except Exception as e:
        print(f"[document/analyze] error: {e}")
        raise HTTPException(status_code=500, detail=f"Ошибка анализа документа: {str(e)}")


@app.post("/api/v1/finance/parse")
async def finance_parse(req: FinanceParseRequest, user=Depends(get_current_user)):
    """AI парсинг — поддерживает несколько расходов в одном сообщении"""
    if not AITUNNEL_API_KEY:
        raise HTTPException(status_code=503, detail="AI недоступен")
    try:
        headers = {"Authorization": f"Bearer {AITUNNEL_API_KEY}", "Content-Type": "application/json"}
        FINANCE_PROMPT = """Ты финансовый помощник. Извлеки из текста или фото чека информацию о расходах.
ВАЖНО: если в тексте несколько расходов — верни массив, каждый расход отдельным объектом.

Категории (category): shop, pharmacy, utility, credit, transport, leisure, other.
Подкатегории: transport(taxi/public/fuel/parking/service/other_transport), pharmacy(drugs/doctor/tests/hospital), shop(groceries/clothes/household), leisure(cafe/cinema/subscription/sport). Остальным subcategory=null.
item_name — название лекарства (только для pharmacy/drugs), иначе null.
comment — краткое русское описание 1-3 слова.
merchant — название места если есть, иначе null.

Верни ТОЛЬКО валидный JSON-массив:
[{"amount":число,"category":"код","subcategory":"код или null","item_name":"строка или null","comment":"описание","merchant":"строка или null"}]
Один расход — тоже массив из одного элемента.
Сумма не найдена — верни [{"error":"no_amount"}]."""

        if req.image_base64:
            # Фото чека — Gemini Vision через AItunnel (мультимодальный)
            headers = {"Authorization": f"Bearer {AITUNNEL_API_KEY}", "Content-Type": "application/json"}
            messages = [{"role": "user", "content": [
                {"type": "image_url", "image_url": {"url": f"data:{req.mime_type};base64,{req.image_base64}"}},
                {"type": "text", "text": FINANCE_PROMPT}
            ]}]
            data = {"model": "gemini-2.5-flash-lite", "messages": messages, "max_tokens": 600}
            response = requests.post(f"{AITUNNEL_BASE_URL}chat/completions", headers=headers, json=data, timeout=30)
            result_text = response.json()["choices"][0]["message"]["content"].strip()
        else:
            # Текст — DeepSeek прямой (дешевле)
            result_text = _ds_post(
                [{"role": "user", "content": f"{FINANCE_PROMPT}\n\nТекст: {req.text}"}],
                max_tokens=600
            )

        import re as re_mod
        json_match = re_mod.search(r"(\[.*?\]|\{.*?\})", result_text, re_mod.DOTALL)
        if not json_match:
            raise HTTPException(status_code=400, detail="Не удалось распознать")

        parsed_raw = json.loads(json_match.group())
        if isinstance(parsed_raw, dict):
            parsed_raw = [parsed_raw]
        if not parsed_raw or parsed_raw[0].get("error"):
            raise HTTPException(status_code=400, detail="Сумма не найдена")

        saved_items = []
        if db_pool:
            async with db_pool.acquire() as conn:
                for item in parsed_raw:
                    if not item.get("amount"):
                        continue
                    sub = item.get("subcategory")
                    if sub in (None, "null", ""): sub = None
                    iname = item.get("item_name")
                    if iname in (None, "null", ""): iname = None
                    merchant = item.get("merchant")
                    if merchant in (None, "null", ""): merchant = None
                    comment = item.get("comment", "")
                    full_comment = comment
                    if merchant and merchant.lower() not in comment.lower():
                        full_comment = comment + " · " + merchant if comment else merchant
                    row_id = await conn.fetchval(
                        "INSERT INTO finance_records (user_id, category, subcategory, item_name, amount, comment) VALUES ($1,$2,$3,$4,$5,$6) RETURNING id",
                        user["id"], item.get("category", "other"), sub, iname, float(item["amount"]), full_comment
                    )
                    saved_items.append({"id": row_id, "amount": float(item["amount"]), "category": item.get("category","other"), "comment": full_comment, "status": "saved"})

        return {"status": "saved", "count": len(saved_items), "items": saved_items}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка: {str(e)}")

@app.delete("/api/v1/finance/records/{record_id}")
async def delete_finance_record(record_id: int, user=Depends(get_current_user)):
    if not db_pool:
        raise HTTPException(status_code=503, detail="БД недоступна")
    async with db_pool.acquire() as conn:
        result = await conn.execute("DELETE FROM finance_records WHERE id=$1 AND user_id=$2", record_id, user["id"])
    return {"status": "deleted", "id": record_id}

class FinanceRecordUpdateRequest(BaseModel):
    amount: Optional[float] = None
    comment: Optional[str] = None

@app.patch("/api/v1/finance/records/{record_id}")
async def update_finance_record(record_id: int, req: FinanceRecordUpdateRequest, user=Depends(get_current_user)):
    if not db_pool:
        raise HTTPException(status_code=503, detail="БД недоступна")
    async with db_pool.acquire() as conn:
        row = await conn.fetchrow("SELECT * FROM finance_records WHERE id=$1 AND user_id=$2", record_id, user["id"])
        if not row:
            raise HTTPException(status_code=404, detail="Запись не найдена")
        new_amount = req.amount if req.amount is not None else row["amount"]
        new_comment = req.comment if req.comment is not None else row["comment"]
        await conn.execute("UPDATE finance_records SET amount=$1, comment=$2 WHERE id=$3 AND user_id=$4", new_amount, new_comment, record_id, user["id"])
    return {"status": "updated", "id": record_id, "amount": new_amount}


@app.get("/api/v1/finance/summary/me")
async def get_finance_summary_me(user=Depends(get_current_user)):
    """Summary расходов текущего пользователя (user_id из токена)"""
    if not db_pool:
        return {"month": "", "categories": []}
    async with db_pool.acquire() as conn:
        rows = await conn.fetch("""
            SELECT category, SUM(amount) as total, COUNT(*) as count
            FROM finance_records
            WHERE user_id=$1
            AND recorded_at >= date_trunc('month', NOW())
            GROUP BY category
        """, user["id"])
        return {"month": datetime.now().strftime("%Y-%m"), "categories": [dict(r) for r in rows]}

@app.get("/api/v1/finance/category/{category}")
async def get_finance_category_detail(category: str, user=Depends(get_current_user)):
    """Детальная статистика по категории: суммы, подкатегории, топ товаров, помесячно"""
    if not db_pool:
        return {}
    async with db_pool.acquire() as conn:
        uid = user["id"]
        # Сумма за текущий месяц
        month_total = await conn.fetchval(
            "SELECT COALESCE(SUM(amount),0) FROM finance_records WHERE user_id=$1 AND category=$2 AND recorded_at >= date_trunc('month', NOW())",
            uid, category) or 0
        # Сумма за прошлый месяц
        prev_total = await conn.fetchval(
            "SELECT COALESCE(SUM(amount),0) FROM finance_records WHERE user_id=$1 AND category=$2 AND recorded_at >= date_trunc('month', NOW()) - interval '1 month' AND recorded_at < date_trunc('month', NOW())",
            uid, category) or 0
        # Количество операций за месяц
        count = await conn.fetchval(
            "SELECT COUNT(*) FROM finance_records WHERE user_id=$1 AND category=$2 AND recorded_at >= date_trunc('month', NOW())",
            uid, category) or 0
        # Разбивка по подкатегориям (текущий месяц)
        subcats = await conn.fetch(
            "SELECT subcategory, SUM(amount) as total, COUNT(*) as cnt FROM finance_records WHERE user_id=$1 AND category=$2 AND recorded_at >= date_trunc('month', NOW()) AND subcategory IS NOT NULL GROUP BY subcategory ORDER BY total DESC",
            uid, category)
        # Топ товаров/лекарств (item_name)
        items = await conn.fetch(
            "SELECT item_name, COUNT(*) as cnt, SUM(amount) as total FROM finance_records WHERE user_id=$1 AND category=$2 AND item_name IS NOT NULL AND recorded_at >= date_trunc('month', NOW()) GROUP BY item_name ORDER BY cnt DESC, total DESC LIMIT 5",
            uid, category)
        # Помесячная динамика (6 месяцев)
        monthly = await conn.fetch(
            "SELECT to_char(date_trunc('month', recorded_at), 'YYYY-MM') as month, SUM(amount) as total FROM finance_records WHERE user_id=$1 AND category=$2 AND recorded_at >= date_trunc('month', NOW()) - interval '5 months' GROUP BY 1 ORDER BY 1",
            uid, category)
        # Профиль для расчёта % от дохода
        prof = await conn.fetchrow("SELECT income FROM user_profile WHERE user_id=$1", uid)
        income = (prof["income"] if prof and prof["income"] else 0) or 0

        return {
            "month_total": float(month_total),
            "prev_total": float(prev_total),
            "count": count,
            "avg_check": round(float(month_total) / count) if count > 0 else 0,
            "income": float(income),
            "income_pct": round(float(month_total) / income * 100) if income > 0 else None,
            "subcategories": [{"key": r["subcategory"], "total": float(r["total"]), "count": r["cnt"]} for r in subcats],
            "top_items": [{"name": r["item_name"], "count": r["cnt"], "total": float(r["total"])} for r in items],
            "monthly": [{"month": r["month"], "total": float(r["total"])} for r in monthly],
        }

@app.get("/api/v1/finance/credits-overview")
async def get_credits_overview(user=Depends(get_current_user)):
    """Сводка по кредитам и займам для дашборда нагрузки"""
    if not db_pool:
        return {}
    async with db_pool.acquire() as conn:
        uid = user["id"]
        credits = await conn.fetch("SELECT name, amount, monthly_payment, rate, term_months FROM credits WHERE user_id=$1", uid)
        loans = await conn.fetch("SELECT name, amount, monthly_payment, due_date FROM loans WHERE user_id=$1", uid)
        prof = await conn.fetchrow("SELECT income FROM user_profile WHERE user_id=$1", uid)
        income = (prof["income"] if prof and prof["income"] else 0) or 0
        # Расходы за текущий месяц (для расчёта свободных денег)
        expenses = await conn.fetchval(
            "SELECT COALESCE(SUM(amount),0) FROM finance_records WHERE user_id=$1 AND recorded_at >= date_trunc('month', NOW())", uid) or 0

        total_payment = sum((c["monthly_payment"] or 0) for c in credits) + sum((l["monthly_payment"] or 0) for l in loans)
        total_debt = sum((c["amount"] or 0) for c in credits) + sum((l["amount"] or 0) for l in loans)
        load_pct = round(total_payment / income * 100) if income > 0 else None
        free_money = income - float(expenses) - total_payment if income > 0 else None

        return {
            "income": float(income),
            "expenses": float(expenses),
            "total_payment": float(total_payment),
            "total_debt": float(total_debt),
            "load_pct": load_pct,
            "free_money": free_money,
            "credits": [{"name": c["name"], "amount": float(c["amount"] or 0), "payment": float(c["monthly_payment"] or 0), "rate": c["rate"], "term": c["term_months"]} for c in credits],
            "loans": [{"name": l["name"], "amount": float(l["amount"] or 0), "payment": float(l["monthly_payment"] or 0), "due": l["due_date"]} for l in loans],
        }

# ═════════════════════════════════════════════════════════════════
# Web Push — уведомления о приёме лекарств
# Размещено в конце файла, после get_current_user, db_pool и всех зависимостей
# ═════════════════════════════════════════════════════════════════

def _normalize_time(t: str) -> str:
    """Нормализует строку времени: '9:00' → '09:00'"""
    if not t or ':' not in t:
        return t
    parts = t.split(':')
    return f"{int(parts[0]):02d}:{int(parts[1]):02d}"

def send_web_push(subscription_info: dict, title: str, body: str, icon: str = "/icon-192.png") -> bool:
    """Отправить Web Push уведомление через pywebpush"""
    if not VAPID_PRIVATE_PEM:
        print("[push] VAPID_PRIVATE_PEM не задан — пропуск")
        return False
    try:
        from pywebpush import webpush
        webpush(
            subscription_info=subscription_info,
            data=json.dumps({"title": title, "body": body, "icon": icon}),
            vapid_private_key=VAPID_PRIVATE_PEM,
            vapid_claims=dict(VAPID_CLAIMS)  # копируем, pywebpush мутирует
        )
        return True
    except Exception as e:
        print(f"[push] Ошибка отправки: {e}")
        return False

# Состояние scheduler для диагностики (in-memory, перезаписывается из БД)
_push_scheduler_state = {
    "started_at": None,
    "last_check": None,
    "checks_count": 0,
    "last_error": None,
    "pushes_sent": 0,
    "is_leader": False,  # этот воркер захватил advisory lock
}

# Уникальный ID для advisory lock (любое 64-битное число)
PUSH_SCHEDULER_LOCK_ID = 8675309

async def push_scheduler():
    """Фоновой планировщик — каждую минуту проверяет расписание лекарств.
    Использует pg_try_advisory_xact_lock — lock берётся и автоматически
    освобождается вместе с транзакцией. Это надёжнее чем session lock
    при rolling restart: старый воркер умер → транзакция закрылась → lock освободился.
    """
    from datetime import date
    import pytz
    global _push_scheduler_state
    print(f"[push_scheduler] Поток запущен (PID {os.getpid()})")
    _push_scheduler_state["started_at"] = datetime.now().isoformat()
    await asyncio.sleep(5)

    # Ждём готовность БД
    while not db_pool:
        await asyncio.sleep(2)

    # Создаём таблицу для heartbeat если её нет
    try:
        async with db_pool.acquire() as conn:
            await conn.execute("""
                CREATE TABLE IF NOT EXISTS scheduler_state (
                    name VARCHAR(50) PRIMARY KEY,
                    started_at TIMESTAMP,
                    last_check TIMESTAMP,
                    checks_count INTEGER DEFAULT 0,
                    pushes_sent INTEGER DEFAULT 0,
                    last_error TEXT,
                    leader_pid INTEGER
                );
            """)
    except Exception as e:
        print(f"[push_scheduler] Ошибка создания таблицы: {e}")

    print(f"[push_scheduler] Начинаю цикл — буду пробовать стать лидером каждую минуту")

    while True:
        try:
            await asyncio.sleep(60)
            now = datetime.now()

            # Пробуем захватить транзакционный advisory lock
            # pg_try_advisory_xact_lock: освобождается АВТОМАТИЧЕСКИ при конце транзакции
            # Даже если процесс умер — PostgreSQL сам освобождает lock
            async with db_pool.acquire() as conn:
                async with conn.transaction():
                    got_lock = await conn.fetchval(
                        "SELECT pg_try_advisory_xact_lock($1)", PUSH_SCHEDULER_LOCK_ID
                    )
                    if not got_lock:
                        # Другой воркер сейчас работает — пропускаем итерацию
                        _push_scheduler_state["is_leader"] = False
                        continue

                    # Мы лидер в этой итерации
                    _push_scheduler_state["is_leader"] = True
                    _push_scheduler_state["last_check"] = now.isoformat()
                    _push_scheduler_state["checks_count"] += 1

                    last_err = None
                    if not VAPID_PRIVATE_PEM:
                        last_err = "VAPID_PRIVATE_PEM не задан"
                        _push_scheduler_state["last_error"] = last_err

                    # Heartbeat в БД
                    await conn.execute("""
                        INSERT INTO scheduler_state
                            (name, started_at, last_check, checks_count, pushes_sent, last_error, leader_pid)
                        VALUES ('push', $1, $2, $3, $4, $5, $6)
                        ON CONFLICT (name) DO UPDATE SET
                            last_check=$2, checks_count=$3, pushes_sent=$4,
                            last_error=$5, leader_pid=$6,
                            started_at=CASE WHEN scheduler_state.leader_pid!=$6
                                           THEN $1 ELSE scheduler_state.started_at END
                    """, now, now,
                        _push_scheduler_state["checks_count"],
                        _push_scheduler_state["pushes_sent"],
                        last_err, os.getpid())

                    if not VAPID_PRIVATE_PEM:
                        # lock освободится при выходе из transaction
                        continue

                    # Текущее время UTC
                    now_utc = datetime.now(pytz.utc)

                    # Берём лекарства вместе с timezone пользователя
                    meds = await conn.fetch("""
                        SELECT m.id, m.user_id, m.name, m.dose, m.times_json,
                               m.taken_today, m.taken_date, m.notify,
                               COALESCE(u.timezone, 'Europe/Moscow') as user_timezone
                        FROM medications m
                        JOIN users u ON u.id = m.user_id
                        WHERE m.is_active=TRUE AND m.notify=TRUE
                    """)

                    for med in meds:
                        try:
                            times = json.loads(med["times_json"] or "[]")
                            times = [_normalize_time(t) for t in times]
                        except:
                            times = []
                        if not times:
                            continue
                        # Конвертируем UTC в timezone пользователя
                        try:
                            user_tz = pytz.timezone(med["user_timezone"])
                        except Exception:
                            user_tz = pytz.timezone("Europe/Moscow")
                        now_user = now_utc.astimezone(user_tz)
                        target_user = now_user + timedelta(minutes=10)
                        target_time = f"{target_user.hour:02d}:{target_user.minute:02d}"

                        if target_time not in times:
                            continue
                        # Проверяем что не принято сегодня
                        today = date.today()
                        if med["taken_date"] == today:
                            try: taken = json.loads(med["taken_today"] or "[]")
                            except: taken = []
                            if target_time in taken:
                                continue
                        # Отправляем push всем подпискам пользователя
                        subs = await conn.fetch(
                            "SELECT endpoint, p256dh, auth FROM push_subscriptions WHERE user_id=$1",
                            med["user_id"]
                        )
                        if not subs:
                            continue
                        body = f"Приём в {target_time}"
                        if med.get("dose"):
                            body += f" · {med['dose']}"
                        for sub in subs:
                            sent = send_web_push(
                                {"endpoint": sub["endpoint"], "keys": {"p256dh": sub["p256dh"], "auth": sub["auth"]}},
                                title=f"Пора принять {med['name']}",
                                body=body
                            )
                            if sent:
                                _push_scheduler_state["pushes_sent"] += 1
                                print(f"[push] ✅ {med['name']} → user {med['user_id']} @ {target_time}")
                                await create_notification(
                                    user_id=med["user_id"],
                                    type="med",
                                    title=f"Время принять {med['name']}",
                                    body=body
                                )
                    # Транзакция завершается → lock автоматически освобождается

        except Exception as e:
            print(f"[push_scheduler] Ошибка итерации: {e}")
            _push_scheduler_state["last_error"] = str(e)

# ── Service Worker (с push-обработчиком и pushsubscriptionchange) ──
SW_CONTENT = """
const CACHE = 'oktolk-v5';
const PRECACHE = ['/', '/manifest.json'];

self.addEventListener('install', e => {
  e.waitUntil(caches.open(CACHE).then(c => c.addAll(PRECACHE)).then(() => self.skipWaiting()));
});
self.addEventListener('activate', e => {
  e.waitUntil(caches.keys().then(keys =>
    Promise.all(keys.filter(k => k !== CACHE).map(k => caches.delete(k)))
  ).then(() => self.clients.claim()));
});
self.addEventListener('fetch', e => {
  if (e.request.method !== 'GET') return;
  e.respondWith(fetch(e.request).catch(() => caches.match(e.request)));
});

/* ── Web Push: получение и показ уведомления ── */
self.addEventListener('push', e => {
  let data = { title: 'OkTolk', body: 'Напоминание' };
  try { data = e.data ? e.data.json() : data; } catch(err) {}
  e.waitUntil(self.registration.showNotification(data.title || 'OkTolk', {
    body: data.body || '',
    icon: data.icon || '/icon-192.png',
    badge: '/icon-192.png',
    vibrate: [200, 100, 200],
    tag: 'meds-reminder',
    requireInteraction: true,
    data: { url: '/' }
  }));
});

self.addEventListener('notificationclick', e => {
  e.notification.close();
  e.waitUntil(clients.matchAll({ type: 'window', includeUncontrolled: true }).then(cs => {
    const c = cs.find(x => x.url.includes(self.location.host));
    return c ? c.focus() : clients.openWindow('/');
  }));
});

/* ── Ротация подписки — пересохраняем на сервере ── */
self.addEventListener('pushsubscriptionchange', e => {
  e.waitUntil(
    self.registration.pushManager.subscribe({ userVisibleOnly: true })
      .then(sub => {
        const json = sub.toJSON();
        return fetch('/api/v1/push/subscribe', {
          method: 'POST',
          headers: { 'Content-Type': 'application/json' },
          body: JSON.stringify({ endpoint: json.endpoint, p256dh: json.keys.p256dh, auth: json.keys.auth })
        });
      })
  );
});
"""

@app.get("/sw.js")
async def service_worker():
    from fastapi.responses import Response
    return Response(content=SW_CONTENT, media_type="application/javascript",
                   headers={"Service-Worker-Allowed": "/", "Cache-Control": "no-cache"})

# ── Push API endpoints ────────────────────────────────────────────

@app.get("/api/v1/push/vapid-key")
async def get_vapid_key():
    """Возвращает VAPID public key для подписки на push"""
    return {"public_key": VAPID_PUBLIC_KEY}

class PushSubscriptionRequest(BaseModel):
    endpoint: str
    p256dh: str
    auth: str

@app.post("/api/v1/push/subscribe")
async def subscribe_push(req: PushSubscriptionRequest, user=Depends(get_current_user)):
    """Сохранить push-подписку пользователя"""
    if not db_pool:
        raise HTTPException(status_code=503, detail="БД недоступна")
    async with db_pool.acquire() as conn:
        await conn.execute("""
            INSERT INTO push_subscriptions (user_id, endpoint, p256dh, auth)
            VALUES ($1, $2, $3, $4)
            ON CONFLICT (user_id, endpoint) DO UPDATE SET p256dh=EXCLUDED.p256dh, auth=EXCLUDED.auth
        """, user["id"], req.endpoint, req.p256dh, req.auth)
    return {"status": "subscribed"}

class PushUnsubRequest(BaseModel):
    endpoint: str

@app.delete("/api/v1/push/subscribe")
async def unsubscribe_push(req: PushUnsubRequest, user=Depends(get_current_user)):
    """Удалить push-подписку"""
    if not db_pool:
        raise HTTPException(status_code=503, detail="БД недоступна")
    async with db_pool.acquire() as conn:
        await conn.execute(
            "DELETE FROM push_subscriptions WHERE user_id=$1 AND endpoint=$2",
            user["id"], req.endpoint
        )
    return {"status": "unsubscribed"}

@app.post("/api/v1/push/test")
async def test_push(user=Depends(get_current_user)):
    """Тестовый push — для проверки что подписка работает"""
    if not db_pool:
        raise HTTPException(status_code=503, detail="БД недоступна")
    async with db_pool.acquire() as conn:
        subs = await conn.fetch("SELECT * FROM push_subscriptions WHERE user_id=$1", user["id"])
    if not subs:
        return {"status": "no_subscriptions", "message": "Нет активных подписок"}
    sent = 0
    for sub in subs:
        ok = send_web_push(
            {"endpoint": sub["endpoint"], "keys": {"p256dh": sub["p256dh"], "auth": sub["auth"]}},
            "OkTolk — тест уведомлений",
            "Push работает! Лекарства будут напоминать вовремя."
        )
        if ok:
            sent += 1
    return {"status": "ok", "sent": sent, "total": len(subs)}

# ═════════════════════════════════════════════════════════════════


# ── Notifications ──────────────────────────────────────────────────────────

async def create_notification(user_id: int, type: str, title: str, body: str = None):
    """Создать уведомление в БД. type: med | scam | system | tokens"""
    if not db_pool:
        return
    try:
        async with db_pool.acquire() as conn:
            # Чистим старые > 7 дней
            await conn.execute(
                "DELETE FROM notifications WHERE created_at < NOW() - INTERVAL '7 days'"
            )
            # Не дублируем одинаковые уведомления в течение 5 минут
            recent = await conn.fetchval("""
                SELECT id FROM notifications
                WHERE user_id=$1 AND title=$2
                  AND created_at > NOW() - INTERVAL '5 minutes'
                LIMIT 1
            """, user_id, title)
            if recent:
                return
            await conn.execute(
                "INSERT INTO notifications (user_id, type, title, body) VALUES ($1,$2,$3,$4)",
                user_id, type, title, body
            )
    except Exception as e:
        print(f"[notif] Ошибка: {e}")

@app.get("/api/v1/notifications")
async def get_notifications(user=Depends(get_current_user)):
    """Получить уведомления пользователя за последние 7 дней"""
    if not db_pool:
        return {"items": [], "unread": 0}
    async with db_pool.acquire() as conn:
        rows = await conn.fetch("""
            SELECT id, type, title, body, is_read,
                   to_char(created_at AT TIME ZONE 'UTC', 'YYYY-MM-DD"T"HH24:MI:SS"Z"') as created_at
            FROM notifications
            WHERE user_id=$1
              AND created_at > NOW() - INTERVAL '7 days'
            ORDER BY created_at DESC
            LIMIT 50
        """, user["id"])
        items = [dict(r) for r in rows]
        unread = sum(1 for r in items if not r["is_read"])
        return {"items": items, "unread": unread}

@app.post("/api/v1/notifications/read-all")
async def read_all_notifications(user=Depends(get_current_user)):
    """Отметить все уведомления как прочитанные"""
    if not db_pool:
        return {"status": "ok"}
    async with db_pool.acquire() as conn:
        await conn.execute(
            "UPDATE notifications SET is_read=TRUE WHERE user_id=$1 AND is_read=FALSE",
            user["id"]
        )
    return {"status": "ok"}

@app.post("/api/v1/notifications/{notif_id}/read")
async def read_notification(notif_id: int, user=Depends(get_current_user)):
    """Отметить одно уведомление как прочитанное"""
    if not db_pool:
        return {"status": "ok"}
    async with db_pool.acquire() as conn:
        await conn.execute(
            "UPDATE notifications SET is_read=TRUE WHERE id=$1 AND user_id=$2",
            notif_id, user["id"]
        )
    return {"status": "ok"}

@app.get("/api/v1/push/diagnostic")
async def push_diagnostic():
    """Диагностика push-уведомлений + состояние лекарств в БД"""
    diag = {
        "vapid_public_key_set": bool(VAPID_PUBLIC_KEY),
        "vapid_private_pem_set": bool(VAPID_PRIVATE_PEM),
        "vapid_private_length": len(VAPID_PRIVATE_PEM) if VAPID_PRIVATE_PEM else 0,
        "db_pool_ready": db_pool is not None,
        "scheduler": dict(_push_scheduler_state),  # копия
        "worker_pid": os.getpid(),
        "pywebpush_available": False,
    }
    try:
        from pywebpush import webpush
        diag["pywebpush_available"] = True
    except ImportError as e:
        diag["pywebpush_error"] = str(e)
    # Общее количество подписок + СОСТОЯНИЕ SCHEDULER ИЗ БД
    if db_pool:
        try:
            async with db_pool.acquire() as conn:
                diag["total_subscriptions"] = await conn.fetchval(
                    "SELECT COUNT(*) FROM push_subscriptions"
                ) or 0
                # Читаем реальное состояние scheduler из БД
                row = await conn.fetchrow("SELECT * FROM scheduler_state WHERE name='push'")
                if row:
                    diag["scheduler_db"] = {
                        "started_at":  row["started_at"].isoformat() if row["started_at"] else None,
                        "last_check":  row["last_check"].isoformat() if row["last_check"] else None,
                        "checks_count": row["checks_count"],
                        "pushes_sent": row["pushes_sent"],
                        "last_error":  row["last_error"],
                        "leader_pid":  row["leader_pid"],
                    }
                    # Считаем секунды с последней проверки
                    if row["last_check"]:
                        diff = (datetime.now() - row["last_check"]).total_seconds()
                        diag["scheduler_db"]["seconds_since_last_check"] = int(diff)
                        diag["scheduler_db"]["alive"] = diff < 180  # жив если проверял <3 мин назад
                else:
                    diag["scheduler_db"] = None
                    diag["scheduler_db_note"] = "Запись scheduler в БД ещё не создана. Подождите 1-2 минуты после деплоя."
        except Exception as e:
            diag["subscriptions_error"] = str(e)
    # Лекарства с уведомлениями — показываем что видит scheduler
    if db_pool:
        try:
            import pytz as _pytz
            from datetime import timedelta as _td
            _now_utc = datetime.now(_pytz.utc)
            async with db_pool.acquire() as conn:
                meds = await conn.fetch("""
                    SELECT m.id, m.name, m.times_per_day, m.times_json,
                           m.notify, m.is_active,
                           COALESCE(u.timezone, 'Europe/Moscow') as user_timezone
                    FROM medications m
                    JOIN users u ON u.id = m.user_id
                    WHERE m.notify=TRUE
                    ORDER BY m.id
                """)
                diag["meds_with_notify"] = []
                for m in meds:
                    try:
                        _tz = _pytz.timezone(m["user_timezone"])
                        _now_user = _now_utc.astimezone(_tz)
                        _target = _now_user + _td(minutes=10)
                        _target_str = f"{_target.hour:02d}:{_target.minute:02d}"
                        _now_str = f"{_now_user.hour:02d}:{_now_user.minute:02d}"
                    except:
                        _target_str = "error"
                        _now_str = "error"
                    diag["meds_with_notify"].append({
                        "id": m["id"],
                        "name": m["name"],
                        "times_json": m["times_json"],
                        "is_active": m["is_active"],
                        "user_timezone": m["user_timezone"],
                        "now_in_user_tz": _now_str,
                        "scheduler_checks_at": _target_str,
                    })
                diag["meds_with_notify_count"] = len(diag["meds_with_notify"])
                diag["server_time_utc"] = _now_utc.strftime("%H:%M")
        except Exception as e:
            diag["meds_error"] = str(e)
    # Итоговый вердикт
    scheduler_alive = diag.get("scheduler_db", {}) and diag["scheduler_db"].get("alive", False)
    diag["status"] = "ok" if (
        diag["vapid_public_key_set"] and
        diag["vapid_private_pem_set"] and
        diag["db_pool_ready"] and
        diag["pywebpush_available"] and
        scheduler_alive
    ) else "broken"
    diag["issues"] = []
    if not diag["vapid_public_key_set"]: diag["issues"].append("VAPID_PUBLIC_KEY не задан")
    if not diag["vapid_private_pem_set"]: diag["issues"].append("VAPID_PRIVATE_PEM не задан в Amvera env")
    if not diag["db_pool_ready"]: diag["issues"].append("База данных недоступна")
    if not diag["pywebpush_available"]: diag["issues"].append("pywebpush не установлен")
    # Проверка scheduler по БД (реальное состояние)
    sdb = diag.get("scheduler_db")
    if not sdb:
        diag["issues"].append("Scheduler ещё не запустился (подождите 1-2 минуты)")
    elif sdb.get("last_check") is None:
        diag["issues"].append("Scheduler запустился, но ни разу не проверил расписание")
    elif not sdb.get("alive"):
        diag["issues"].append(f"Scheduler не отвечает {sdb.get('seconds_since_last_check', 0)} сек")
    return diag

@app.post("/api/v1/auth/sync-timezone")
async def sync_timezone(request: Request):
    """Обновить timezone пользователя. Не требует авторизации если токен невалидный."""
    data = await request.json()
    tz = data.get("timezone", "Europe/Moscow")
    # Валидируем timezone
    try:
        import pytz
        pytz.timezone(tz)
    except Exception:
        return {"status": "ignored", "reason": "invalid_timezone"}
    # Сохраняем в localStorage на фронте (это делает фронт сам)
    # Обновляем в БД только если авторизован
    token = request.headers.get("authorization", "").replace("Bearer ", "").strip()
    if not token or token == "local_demo" or not db_pool:
        return {"status": "ok", "timezone": tz, "saved_to_db": False}
    try:
        user = await get_current_user(authorization=request.headers.get("authorization"))
        async with db_pool.acquire() as conn:
            await conn.execute("UPDATE users SET timezone=$1 WHERE id=$2", tz, user["id"])
        return {"status": "ok", "timezone": tz, "saved_to_db": True}
    except Exception:
        return {"status": "ok", "timezone": tz, "saved_to_db": False}

@app.post("/api/v1/auth/logout-all")
async def logout_all(user=Depends(get_current_user)):
    """Удалить все сессии пользователя — выход со всех устройств"""
    if not db_pool:
        raise HTTPException(status_code=503, detail="БД недоступна")
    async with db_pool.acquire() as conn:
        deleted = await conn.fetchval(
            "DELETE FROM sessions WHERE user_id=$1 RETURNING COUNT(*)",
            user["id"]
        )
    return {"status": "ok", "sessions_deleted": deleted or 0}

@app.get("/robots.txt")
async def robots_txt():
    from fastapi.responses import PlainTextResponse
    content = """User-agent: *
Allow: /
Allow: /app
Allow: /app/finance
Allow: /app/health
Allow: /app/news
Allow: /app/search
Allow: /finance
Allow: /health
Allow: /news
Allow: /search
Disallow: /api/
Disallow: /sw.js

Sitemap: https://oktolk.ru/sitemap.xml
"""
    return PlainTextResponse(content, media_type="text/plain")

@app.get("/sitemap.xml")
async def sitemap_xml():
    from fastapi.responses import Response
    from datetime import date
    today = date.today().isoformat()
    content = f"""<?xml version="1.0" encoding="UTF-8"?>
<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">
  <url>
    <loc>https://oktolk.ru/</loc>
    <lastmod>{today}</lastmod>
    <changefreq>weekly</changefreq>
    <priority>1.0</priority>
  </url>
  <url>
    <loc>https://oktolk.ru/app</loc>
    <lastmod>{today}</lastmod>
    <changefreq>weekly</changefreq>
    <priority>0.9</priority>
  </url>
  <url>
    <loc>https://oktolk.ru/app/finance</loc>
    <lastmod>{today}</lastmod>
    <changefreq>weekly</changefreq>
    <priority>0.8</priority>
  </url>
  <url>
    <loc>https://oktolk.ru/app/health</loc>
    <lastmod>{today}</lastmod>
    <changefreq>weekly</changefreq>
    <priority>0.8</priority>
  </url>
  <url>
    <loc>https://oktolk.ru/app/news</loc>
    <lastmod>{today}</lastmod>
    <changefreq>daily</changefreq>
    <priority>0.8</priority>
  </url>
  <url>
    <loc>https://oktolk.ru/app/search</loc>
    <lastmod>{today}</lastmod>
    <changefreq>weekly</changefreq>
    <priority>0.7</priority>
  </url>
</urlset>"""
    return Response(content=content, media_type="application/xml")

@app.get("/app/finance")
@app.get("/app/health")
@app.get("/app/news")
@app.get("/app/search")
async def spa_app_routes():
    """SPA fallback для маршрутов /app/* — отдают index.html"""
    return FileResponse("index.html")

@app.get("/finance")
@app.get("/health")
@app.get("/news")
@app.get("/search")
async def spa_routes_legacy():
    """Обратная совместимость — старые маршруты тоже работают"""
    return FileResponse("index.html")

app.mount("/", StaticFiles(directory=".", html=True), name="static")
